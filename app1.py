from __future__ import annotations

import base64
import copy
import html
import json
import re
import time
from dataclasses import asdict
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from textwrap import dedent
from typing import Any

import networkx as nx
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
import streamlit.components.v1 as components
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.ai.task_generator import generate_task_plan
from src.analytics.metrics import compute_metrics
from src.analytics.risk_drivers import rank_delay_drivers
from src.analytics.scenarios import scenario_comparison
from src.config import settings
from src.ml.predictor import RiskModelStatus, load_risk_model
from src.ml.service import build_default_feature_df, score_tasks
from src.modeling.critical_path import critical_path_by_mean
from src.modeling.graph_builder import build_project_graph
from src.simulation.monte_carlo import run_monte_carlo
from src.storage.db import init_db
from src.storage.repository import get_run_details, list_recent_runs, save_session_run
from src.utils.errors import GraphValidationError, TaskGenerationError
from src.utils.emailing import (
    build_subscription_confirmation_link,
    send_registration_confirmation_email,
    send_subscription_confirmation_email,
)
from src.visualization.charts import (
    completion_histogram,
    dependency_graph_figure,
    risk_drivers_bar,
    scenario_comparison_chart,
)


st.set_page_config(
    page_title="certAIn | SaaS Experience",
    page_icon="A",
    layout="wide",
    initial_sidebar_state="expanded",
)


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
LEADERSHIP_IMAGES = ASSETS / "leadership_images"
NAV_LOGO_PNG_PATH = ASSETS / "logo-mark.png"
LOGO_PNG_PATH = ASSETS / "logo.png"
LEGACY_LOGO_PNG_PATH = ASSETS / "cetAIn logo.png"
LOGO_SVG_PATH = ASSETS / "logo.svg"


NAV_ITEMS = [
    ("home", "Home"),
    ("product", "Product"),
    ("forecasting", "AI Forecasting"),
    ("monte-carlo", "Monte Carlo"),
    ("pricing", "Pricing"),
    ("dashboard", "Dashboard"),
    ("about", "About Us"),
    ("contact", "Contact Us"),
    ("faq", "FAQ"),
    ("login", "Login"),
]

PAGE_TITLES = {key: label for key, label in NAV_ITEMS}
PAGE_TITLES["subscription-confirmation"] = "Subscription Confirmation"
PAGE_TITLES["registration-confirmation"] = "Registration Confirmation"
PAGE_TITLES["privacy-policy"] = "Privacy Policy"
PAGE_TITLES["terms-of-service"] = "Terms of Service"
PAGE_TITLES["cookie-settings"] = "Cookie Settings"
PAGE_TITLES["security-data-protection"] = "Security & Data Protection"

DASHBOARD_ITEMS = [
    ("guide", "User Guide"),
    ("decision-hub", "Decision Hub"),
    ("task-architect", "AI Task Architect"),
    ("risk-intelligence", "Risk Intelligence"),
    ("simulator", "Monte Carlo Simulator"),
    ("what-if", "What-If Scenarios"),
    ("history", "Version History"),
    ("summary", "Executive Summary"),
    ("integrations", "Tools Integration"),
    ("settings-page", "Settings"),
]

for key, label in DASHBOARD_ITEMS:
    PAGE_TITLES[key] = label

COMMAND_PALETTE_ITEMS = [
    ("Platform", "home", "Home", "Return to the landing page"),
    ("Platform", "product", "Product", "Explore the full product story"),
    ("Platform", "forecasting", "AI Forecasting", "Review advisory ML and forecast evidence"),
    ("Platform", "monte-carlo", "Monte Carlo", "See probabilistic schedule simulation"),
    ("Platform", "pricing", "Pricing", "Compare plan options and trial entry points"),
    ("Platform", "dashboard", "Dashboard", "Open the forecast console"),
    ("Platform", "about", "About Us", "Read the capstone and team story"),
    ("Platform", "contact", "Contact Us", "Reach the certAIn team"),
    ("Platform", "faq", "FAQ", "Browse common product and demo questions"),
    ("Platform", "login", "Login", "Open the authentication page"),
    ("Workspace", "guide", "User Guide", "Learn the workspace flow"),
    ("Workspace", "decision-hub", "Decision Hub", "Review executive decision signals"),
    ("Workspace", "task-architect", "AI Task Architect", "Generate and inspect task plans"),
    ("Workspace", "risk-intelligence", "Risk Intelligence", "Inspect generated task structures and dependency risk"),
    ("Workspace", "simulator", "Monte Carlo Simulator", "Run and compare schedule simulations"),
    ("Workspace", "what-if", "What-If Scenarios", "Test mitigation and acceleration options"),
    ("Workspace", "history", "Version History", "Open recent workspace runs"),
    ("Workspace", "summary", "Executive Summary", "Open the stakeholder-ready summary"),
    ("Workspace", "integrations", "Tools Integration", "Review connected delivery tools"),
    ("Workspace", "settings-page", "Settings", "Adjust workspace preferences"),
]

DASHBOARD_PAGE_KEYS = {key for key, _ in DASHBOARD_ITEMS} | {"dashboard"}

DEFAULT_SAMPLE = "Digital Product Launch"

SAMPLE_SCENARIOS: dict[str, dict[str, Any]] = {
    "Digital Product Launch": {
        "description": (
            "Launch a B2B analytics SaaS with SSO, subscription billing, AI forecast pages, "
            "and an executive dashboard for pilot customers."
        ),
        "deadline": 96.0,
        "iterations": 1200,
        "max_tasks": 10,
        "tasks": [
            {"id": "T1", "name": "Product scope and requirements", "mean_duration": 6, "std_dev": 1.4, "dependencies": [], "risk_factor": 0.24},
            {"id": "T2", "name": "Platform architecture and security design", "mean_duration": 8, "std_dev": 2.0, "dependencies": ["T1"], "risk_factor": 0.33},
            {"id": "T3", "name": "Core application build", "mean_duration": 18, "std_dev": 4.5, "dependencies": ["T2"], "risk_factor": 0.44},
            {"id": "T4", "name": "Billing and subscription workflow", "mean_duration": 10, "std_dev": 3.0, "dependencies": ["T2"], "risk_factor": 0.48},
            {"id": "T5", "name": "SSO and identity controls", "mean_duration": 9, "std_dev": 2.6, "dependencies": ["T2"], "risk_factor": 0.4},
            {"id": "T6", "name": "AI forecasting experience", "mean_duration": 12, "std_dev": 3.1, "dependencies": ["T3"], "risk_factor": 0.46},
            {"id": "T7", "name": "Integration and data contracts", "mean_duration": 11, "std_dev": 3.7, "dependencies": ["T3", "T4", "T5"], "risk_factor": 0.56},
            {"id": "T8", "name": "Pilot QA and hardening", "mean_duration": 13, "std_dev": 4.0, "dependencies": ["T6", "T7"], "risk_factor": 0.59},
            {"id": "T9", "name": "Executive reporting and rollout", "mean_duration": 6, "std_dev": 1.8, "dependencies": ["T8"], "risk_factor": 0.28},
        ],
    },
    "ERP Transformation": {
        "description": (
            "Roll out an ERP transformation across finance, procurement, and operations with "
            "data migration, vendor coordination, training, and cutover readiness."
        ),
        "deadline": 142.0,
        "iterations": 1400,
        "max_tasks": 11,
        "tasks": [
            {"id": "T1", "name": "Program charter and scope alignment", "mean_duration": 9, "std_dev": 1.7, "dependencies": [], "risk_factor": 0.21},
            {"id": "T2", "name": "Current-state process mapping", "mean_duration": 12, "std_dev": 2.5, "dependencies": ["T1"], "risk_factor": 0.31},
            {"id": "T3", "name": "ERP configuration blueprint", "mean_duration": 16, "std_dev": 3.2, "dependencies": ["T2"], "risk_factor": 0.39},
            {"id": "T4", "name": "Data migration design", "mean_duration": 14, "std_dev": 3.8, "dependencies": ["T2"], "risk_factor": 0.47},
            {"id": "T5", "name": "Vendor and integration build", "mean_duration": 18, "std_dev": 4.3, "dependencies": ["T3"], "risk_factor": 0.52},
            {"id": "T6", "name": "Master data cleansing", "mean_duration": 15, "std_dev": 4.1, "dependencies": ["T4"], "risk_factor": 0.55},
            {"id": "T7", "name": "User training and change management", "mean_duration": 11, "std_dev": 2.4, "dependencies": ["T3"], "risk_factor": 0.34},
            {"id": "T8", "name": "System integration testing", "mean_duration": 13, "std_dev": 3.6, "dependencies": ["T5", "T6"], "risk_factor": 0.61},
            {"id": "T9", "name": "Cutover rehearsal", "mean_duration": 8, "std_dev": 2.2, "dependencies": ["T7", "T8"], "risk_factor": 0.43},
            {"id": "T10", "name": "Go-live and stabilization", "mean_duration": 10, "std_dev": 2.8, "dependencies": ["T9"], "risk_factor": 0.37},
        ],
    },
    "Infrastructure Modernization": {
        "description": (
            "Modernize a smart operations platform with edge sensors, data pipelines, a command "
            "dashboard, and a phased deployment across multiple sites."
        ),
        "deadline": 128.0,
        "iterations": 1100,
        "max_tasks": 10,
        "tasks": [
            {"id": "T1", "name": "Site discovery and rollout sequencing", "mean_duration": 8, "std_dev": 1.8, "dependencies": [], "risk_factor": 0.26},
            {"id": "T2", "name": "Network and platform architecture", "mean_duration": 11, "std_dev": 2.3, "dependencies": ["T1"], "risk_factor": 0.34},
            {"id": "T3", "name": "Sensor procurement and logistics", "mean_duration": 13, "std_dev": 3.4, "dependencies": ["T1"], "risk_factor": 0.49},
            {"id": "T4", "name": "Data ingestion pipeline build", "mean_duration": 15, "std_dev": 3.1, "dependencies": ["T2"], "risk_factor": 0.42},
            {"id": "T5", "name": "Control dashboard implementation", "mean_duration": 14, "std_dev": 2.9, "dependencies": ["T2"], "risk_factor": 0.37},
            {"id": "T6", "name": "Site installation wave one", "mean_duration": 16, "std_dev": 4.6, "dependencies": ["T3"], "risk_factor": 0.57},
            {"id": "T7", "name": "Telemetry integration and validation", "mean_duration": 12, "std_dev": 3.0, "dependencies": ["T4", "T6"], "risk_factor": 0.51},
            {"id": "T8", "name": "Alerting and forecasting layer", "mean_duration": 11, "std_dev": 2.4, "dependencies": ["T5", "T7"], "risk_factor": 0.45},
            {"id": "T9", "name": "Site deployment wave two", "mean_duration": 14, "std_dev": 4.3, "dependencies": ["T7"], "risk_factor": 0.54},
            {"id": "T10", "name": "Operational handover", "mean_duration": 7, "std_dev": 1.5, "dependencies": ["T8", "T9"], "risk_factor": 0.25},
        ],
    },
}

HOME_TICKER_ITEMS = [
    "Risk model updated",
    "Schedule optimized",
    "3 anomalies detected",
    "Portfolio drift reviewed",
    "Executive summary refreshed",
    "Resource plan rebalanced",
]

HOME_METRICS = [
    ("Projects Optimized", "2,847", "+18% portfolio throughput"),
    ("Time Saved", "34%", "Planning cycle reduction"),
    ("Risk Accuracy", "94.7%", "Monitoring-backed signal quality"),
    ("Cost Reduction", "$2.1M", "Modeled efficiency upside"),
]

HOME_FEATURES = [
    ("AI Schedule Optimization", "Turn a raw brief into structured timelines, dependencies, and delivery logic."),
    ("Monte Carlo Simulation", "Show P50, P80, and delay probability with a presentation-ready forecast story."),
    ("Risk Intelligence Engine", "Blend advisory ML, portfolio history, and monitoring drift into one view."),
    ("Resource Optimization", "Spot overloaded workstreams and rebalance before the schedule slips."),
    ("Decision Intelligence", "Compare mitigation, baseline, and acceleration options side by side."),
    ("Predictive Analytics", "Connect historical portfolio evidence to your current plan and executive narrative."),
]

HOW_IT_WORKS_STEPS = [
    ("01", "Upload", "Import a brief, spreadsheet, or PM dataset and let certAIn structure the work."),
    ("02", "AI analyzes", "Generate tasks, score risk, simulate completion ranges, and identify fragile assumptions."),
    ("03", "Get insights", "Present dashboards, executive summaries, and scenario recommendations in one workflow."),
]

TRUSTED_BRANDS = [
    "Airbus",
    "Boeing",
    "Siemens",
    "Bosch",
    "Deloitte",
    "Accenture",
    "JPMorgan",
    "PwC",
]

FOOTER_LINK_COLUMNS = [
    (
        "Product",
        [
            ("AI Forecasting", "forecasting", None),
            ("Monte Carlo Simulation", "monte-carlo", None),
            ("Dashboard", "dashboard", None),
            ("Pricing", "pricing", None),
        ],
    ),
    (
        "Company",
        [
            ("About Us", "about", None),
            ("Contact Us", "contact", None),
            ("About the Product", "product", None),
            ("Careers", None, "Coming Soon"),
        ],
    ),
    (
        "Resources",
        [
            ("User Guide", "guide", None),
            ("Tools Integration", "integrations", None),
            ("Executive Summary", "summary", None),
            ("FAQ", "faq", None),
        ],
    ),
    (
        "Legal",
        [
            ("Privacy Policy", "privacy-policy", None),
            ("Terms of Service", "terms-of-service", None),
            ("Cookie Settings", "cookie-settings", None),
            ("Security & Data Protection", "security-data-protection", None),
        ],
    ),
]

FOOTER_SOCIALS = [
    ("linkedin", "LinkedIn"),
    ("x", "X"),
    ("github", "GitHub"),
    ("youtube", "YouTube"),
    ("instagram", "Instagram"),
    ("facebook", "Facebook"),
]

FOOTER_SOCIAL_LINKS = {
    "linkedin": "https://www.linkedin.com/",
    "x": "https://x.com/",
    "github": "https://github.com/",
    "youtube": "https://www.youtube.com/@meysamehshojaei1296",
    "instagram": "https://www.instagram.com/",
    "facebook": "https://www.facebook.com/",
}

LANGUAGE_OPTIONS = [
    ("en", "English", "en"),
    ("fa", "فارسی", "fa"),
    ("iw", "עברית", "iw"),
    ("de", "Deutsch", "de"),
    ("it", "Italiano", "it"),
    ("es", "Español", "es"),
    ("fr", "Français", "fr"),
    ("pt", "Português", "pt"),
    ("zh-CN", "中文", "zh-CN"),
    ("ja", "日本語", "ja"),
    ("ar", "العربية", "ar"),
]

LANGUAGE_LABELS = {code: label for code, label, _ in LANGUAGE_OPTIONS}
LANGUAGE_WIDGET_CODES = {code: widget_code for code, _, widget_code in LANGUAGE_OPTIONS}
SUPPORTED_LANGUAGE_CODES = set(LANGUAGE_LABELS)

FOOTER_COMPLIANCE = [
    ("ISO", "ISO 27001", "International"),
    ("S2", "SOC2", "USA"),
    ("CA", "CCPA", "California"),
    ("EU", "GDPR", "EU"),
    ("PI", "PIPEDA", "Canada"),
    ("LG", "LGPD", "Brazil"),
    ("AP", "APPI", "Japan"),
]

GLOBAL_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Manrope:wght@400;500;600;700;800&family=Sora:wght@400;500;600;700;800&display=swap');

:root {
  --bg: #020711;
  --bg-soft: #071321;
  --bg-card: rgba(7, 18, 31, 0.74);
  --bg-card-strong: rgba(10, 20, 34, 0.92);
  --purple: #9b5cff;
  --purple-strong: #7a43ff;
  --line: rgba(109, 165, 255, 0.17);
  --line-strong: rgba(109, 165, 255, 0.28);
  --text: #f4f8ff;
  --muted: #9eb0cb;
  --blue: #64a2ff;
  --blue-strong: #4c8af4;
  --brand-blue-base: #0b6fd6;
  --brand-blue-accent: #5ca9ff;
  --cyan: #76e7ff;
  --teal: #18d1c7;
  --gold: #f5c56b;
  --rose: #ff7f8c;
  --shadow: 0 28px 90px rgba(0, 0, 0, 0.46);
  --radius-xl: 30px;
  --radius-lg: 24px;
  --radius-md: 18px;
  --radius-sm: 14px;
  --content-width: 1240px;
  --nav-width: 1760px;
  --max-width: var(--content-width);
}

::selection {
  background: rgba(100, 162, 255, 0.32);
  color: #ffffff;
}

html, body, [class*="css"] {
  font-family: "Manrope", sans-serif;
}

body {
  color: var(--text);
}

body {
  top: 0 !important;
}

.goog-te-banner-frame.skiptranslate,
.goog-te-gadget,
.goog-logo-link,
.goog-te-gadget span,
#google_translate_element,
#certain_google_translate_element {
  display: none !important;
  visibility: hidden !important;
}

a {
  color: inherit;
}

@keyframes floatDrift {
  0% { transform: translate3d(0, 0, 0) scale(1); }
  50% { transform: translate3d(0, -14px, 0) scale(1.04); }
  100% { transform: translate3d(0, 0, 0) scale(1); }
}

@keyframes gradientShift {
  0% { background-position: 0% 50%; }
  50% { background-position: 100% 50%; }
  100% { background-position: 0% 50%; }
}

@keyframes tickerSlide {
  0% { transform: translateX(0); }
  100% { transform: translateX(-50%); }
}

@keyframes scanDrift {
  0% { transform: translateY(-125%); opacity: 0; }
  20% { opacity: 0.55; }
  80% { opacity: 0.18; }
  100% { transform: translateY(130%); opacity: 0; }
}

@keyframes pulseGlow {
  0%, 100% { box-shadow: 0 0 0 rgba(122, 67, 255, 0); }
  50% { box-shadow: 0 0 0 1px rgba(155, 92, 255, 0.18), 0 24px 64px rgba(76, 138, 244, 0.22); }
}

@keyframes previewShellSweep {
  0% { transform: translateX(-120%); opacity: 0; }
  18% { opacity: 0.32; }
  55% { opacity: 0.14; }
  100% { transform: translateX(135%); opacity: 0; }
}

@keyframes previewCardFloat {
  0%, 100% { transform: translate3d(0, 0, 0); }
  50% { transform: translate3d(0, -8px, 0); }
}

@keyframes chartBarLift {
  0%, 100% { transform: translateY(0) scaleY(1); filter: brightness(1); }
  50% { transform: translateY(-4px) scaleY(1.05); filter: brightness(1.12); }
}

@keyframes alertSweep {
  0% { transform: translateX(-130%); opacity: 0; }
  20% { opacity: 0.2; }
  100% { transform: translateX(135%); opacity: 0; }
}

@keyframes alertPulse {
  0%, 100% { box-shadow: none; }
  50% { box-shadow: 0 0 0 1px rgba(255, 255, 255, 0.04), 0 10px 24px rgba(0, 0, 0, 0.16); }
}

@keyframes alertDotBlink {
  0%, 100% { transform: scale(1); opacity: 0.95; }
  50% { transform: scale(1.18); opacity: 1; }
}

@keyframes timelineFlow {
  0% { background-position: 0% 50%; filter: brightness(0.98); }
  50% { background-position: 100% 50%; filter: brightness(1.1); }
  100% { background-position: 0% 50%; filter: brightness(0.98); }
}

@keyframes constellationDrift {
  0%, 100% { transform: translate3d(0, 0, 0) scale(var(--group-scale, 1)); }
  50% { transform: translate3d(var(--drift-x, 12px), var(--drift-y, -10px), 0) scale(var(--group-scale, 1)); }
}

@keyframes nodePulse {
  0%, 100% { opacity: 0.72; filter: brightness(0.96); }
  50% { opacity: 1; filter: brightness(1.18); }
}

@keyframes linkPulse {
  0%, 100% { opacity: 0.14; }
  50% { opacity: 0.34; }
}

[data-testid="stHeader"],
[data-testid="stToolbar"],
[data-testid="stDecoration"],
[data-testid="stStatusWidget"],
#MainMenu,
footer {
  display: none !important;
  visibility: hidden !important;
}

[data-testid="stSidebar"] {
  background: linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) !important;
  border-right: 1px solid rgba(109, 165, 255, 0.14);
}

[data-testid="stSidebar"] > div:first-child {
  background: transparent !important;
  height: 100vh !important;
  overflow-y: auto !important;
  overflow-x: hidden !important;
  padding-bottom: 1.25rem;
  scrollbar-width: thin;
  scrollbar-color: rgba(117, 231, 255, 0.34) rgba(255, 255, 255, 0.04);
}

[data-testid="stSidebar"] [data-testid="stSidebarContent"],
[data-testid="stSidebar"] [data-testid="stSidebarUserContent"] {
  height: auto !important;
  min-height: 100vh;
  overflow: visible !important;
  padding-bottom: 1.25rem;
}

[data-testid="stSidebar"] > div:first-child::-webkit-scrollbar {
  width: 10px;
}

[data-testid="stSidebar"] > div:first-child::-webkit-scrollbar-track {
  background: rgba(255, 255, 255, 0.04);
}

[data-testid="stSidebar"] > div:first-child::-webkit-scrollbar-thumb {
  border-radius: 999px;
  background: linear-gradient(180deg, rgba(122, 67, 255, 0.75), rgba(24, 209, 199, 0.72));
}

.stApp,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > .main {
  background:
    radial-gradient(circle at 14% 10%, rgba(155, 92, 255, 0.16), transparent 26%),
    radial-gradient(circle at 32% 12%, rgba(24, 110, 255, 0.12), transparent 24%),
    radial-gradient(circle at 80% 24%, rgba(25, 210, 201, 0.12), transparent 20%),
    linear-gradient(180deg, #030811 0%, #02060d 100%) !important;
  color: var(--text) !important;
}

.block-container {
  max-width: var(--nav-width) !important;
  padding-top: 1.1rem !important;
  padding-bottom: 4rem !important;
}

.saas-shell {
  position: fixed;
  inset: 0;
  width: 100vw;
  height: 100vh;
  pointer-events: none;
  overflow: hidden;
  z-index: 0;
}

.particle-network {
  position: fixed;
  inset: 0;
  z-index: 0;
  pointer-events: none;
  overflow: hidden;
  opacity: 0.9;
}

.particle-network::before {
  content: "";
  position: absolute;
  inset: 0;
  background:
    radial-gradient(circle, rgba(196, 221, 255, 0.22) 0 1px, transparent 1.4px) 8% 18% / 320px 220px,
    radial-gradient(circle, rgba(120, 176, 255, 0.18) 0 1px, transparent 1.5px) 72% 32% / 360px 260px,
    radial-gradient(circle, rgba(117, 239, 230, 0.18) 0 1px, transparent 1.4px) 34% 74% / 420px 280px;
  opacity: 0.45;
  animation: gradientShift 38s linear infinite;
}

.particle-network::after {
  content: "";
  position: absolute;
  inset: -12%;
  background:
    radial-gradient(circle at 65% 40%, rgba(102, 198, 255, 0.14), transparent 20%),
    radial-gradient(circle at 36% 62%, rgba(25, 210, 201, 0.08), transparent 18%),
    radial-gradient(circle at 82% 18%, rgba(155, 92, 255, 0.1), transparent 16%);
  filter: blur(48px);
  opacity: 0.42;
}

.particle-group {
  position: absolute;
  width: var(--group-w);
  height: var(--group-h);
  animation: constellationDrift var(--duration, 18s) ease-in-out infinite;
  animation-delay: var(--delay, 0s);
}

.particle-group.group-a {
  top: 16%;
  right: 7%;
  --group-w: 280px;
  --group-h: 156px;
  --drift-x: 16px;
  --drift-y: -10px;
  --duration: 19s;
}

.particle-group.group-b {
  top: 43%;
  left: 31%;
  --group-w: 212px;
  --group-h: 172px;
  --drift-x: -12px;
  --drift-y: 12px;
  --duration: 17s;
  --delay: -7s;
}

.particle-group.group-c {
  bottom: 15%;
  right: 9%;
  --group-w: 184px;
  --group-h: 118px;
  --drift-x: 10px;
  --drift-y: -8px;
  --duration: 14s;
  --delay: -11s;
}

.particle-link {
  position: absolute;
  height: 1px;
  border-radius: 999px;
  transform-origin: left center;
  background: linear-gradient(
    90deg,
    rgba(115, 171, 255, 0),
    rgba(158, 214, 255, 0.55),
    rgba(115, 171, 255, 0.08)
  );
  box-shadow: 0 0 14px rgba(92, 169, 255, 0.18);
  animation: linkPulse 6.5s ease-in-out infinite;
}

.particle-node {
  position: absolute;
  width: var(--size, 8px);
  height: var(--size, 8px);
  border-radius: 999px;
  transform: translate(-50%, -50%);
  background: radial-gradient(
    circle,
    rgba(244, 250, 255, 0.98) 0 24%,
    rgba(var(--particle-rgb, 114, 178, 255), 0.94) 42%,
    rgba(var(--particle-rgb, 114, 178, 255), 0) 78%
  );
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.06),
    0 0 18px rgba(var(--particle-rgb, 114, 178, 255), 0.3),
    0 0 38px rgba(var(--particle-rgb, 114, 178, 255), 0.2);
  animation: nodePulse var(--pulse, 5.5s) ease-in-out infinite;
  animation-delay: var(--delay, 0s);
}

.particle-node.blue {
  --particle-rgb: 111, 170, 255;
}

.particle-node.teal {
  --particle-rgb: 104, 236, 227;
}

.particle-node.violet {
  --particle-rgb: 174, 127, 255;
}

.particle-node.soft {
  --particle-rgb: 210, 230, 255;
}

.particle-node::after {
  content: "";
  position: absolute;
  inset: -7px;
  border-radius: inherit;
  background: radial-gradient(circle, rgba(var(--particle-rgb, 114, 178, 255), 0.18), transparent 72%);
  filter: blur(7px);
}

.ambient-orb {
  position: fixed;
  border-radius: 999px;
  filter: blur(110px);
  opacity: 0.28;
  z-index: 0;
  pointer-events: none;
  animation: floatDrift 18s ease-in-out infinite;
}

.orb-a {
  width: 16rem;
  height: 16rem;
  top: 7rem;
  left: 7%;
  background: rgba(155, 92, 255, 0.18);
}

.orb-b {
  width: 20rem;
  height: 20rem;
  top: 18rem;
  right: 7%;
  background: rgba(25, 210, 201, 0.12);
  animation-delay: -6s;
}

.orb-c {
  width: 18rem;
  height: 18rem;
  top: 26rem;
  left: 34%;
  background: rgba(100, 162, 255, 0.13);
  animation-delay: -11s;
}

@media (max-width: 900px) {
  .particle-group.group-a {
    right: 4%;
    --group-scale: 0.86;
  }

  .particle-group.group-b {
    left: 14%;
    top: 48%;
    --group-scale: 0.82;
  }

  .particle-group.group-c {
    right: 4%;
    bottom: 18%;
    --group-scale: 0.78;
  }
}

.hero-shell,
.page-shell {
  width: min(100%, var(--content-width));
  margin-inline: auto;
}

[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > .main,
.block-container,
.nav-shell,
[data-testid="stSidebar"] {
  position: relative;
}

.block-container,
.nav-shell,
[data-testid="stSidebar"] {
  z-index: 1;
}

.nav-shell {
  width: min(100%, var(--nav-width));
  margin-inline: auto;
  position: sticky;
  top: 0;
  z-index: 20;
  display: flex;
  justify-content: flex-start;
  align-items: center;
  gap: 0.75rem;
  padding: 0.65rem 1rem;
  backdrop-filter: blur(18px);
  border: 1px solid transparent;
  background:
    linear-gradient(rgba(8, 16, 29, 0.88), rgba(8, 16, 29, 0.82)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.34), rgba(100, 162, 255, 0.22), rgba(24, 209, 199, 0.26)) border-box;
  border-radius: 24px;
  overflow: hidden;
}

.nav-brand-zone,
.top-nav-links,
.top-nav-meta,
.sidebar-row-link,
.sidebar-preview-wrap {
  display: flex;
}

.nav-brand-zone {
  align-items: center;
  gap: 0.55rem;
  flex: 0 0 auto;
  text-decoration: none;
}

.top-nav-links {
  align-items: center;
  gap: 0.18rem;
  flex: 1 1 auto;
  min-width: 0;
  justify-content: flex-start;
  overflow-x: visible;
  scrollbar-width: none;
}

.top-nav-links::-webkit-scrollbar {
  display: none;
}

.top-nav-meta {
  align-items: center;
  gap: 0.5rem;
  flex: 0 0 auto;
  margin-left: auto;
}

.top-nav-pill {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  gap: 0.5rem;
  min-height: 38px;
  padding: 0 0.9rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.04);
  color: var(--muted);
  font-size: 0.82rem;
  font-weight: 700;
  white-space: nowrap;
}

.top-nav-pill-dot {
  width: 10px;
  height: 10px;
  border-radius: 999px;
  background: #19d97c;
  box-shadow: 0 0 0 6px rgba(25, 217, 124, 0.12);
  flex: 0 0 auto;
}

.nav-brand-box {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 62px;
  height: 58px;
  padding: 0;
  border-radius: 0.35rem;
  border: 1px solid rgba(142, 97, 255, 0.16);
  background:
    linear-gradient(180deg, rgba(10, 14, 28, 0.98), rgba(6, 9, 18, 0.98));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.015) inset,
    0 0 34px rgba(122, 67, 255, 0.34),
    0 18px 42px rgba(0, 0, 0, 0.28);
  position: relative;
  overflow: visible;
}

.nav-brand-box::before {
  content: "";
  position: absolute;
  inset: -10px;
  border-radius: 0.7rem;
  background:
    radial-gradient(circle, rgba(122, 67, 255, 0.38) 0%, rgba(91, 157, 255, 0.2) 36%, rgba(122, 67, 255, 0) 76%);
  filter: blur(14px);
  z-index: -1;
}

.nav-brand-icon {
  display: block;
  width: 100%;
  height: 100%;
  object-fit: cover;
  border-radius: 0.28rem;
  filter: drop-shadow(0 10px 24px rgba(91, 157, 255, 0.2));
}

.nav-brand-word {
  font-family: "Sora", sans-serif;
  font-size: 1.18rem;
  font-weight: 700;
  letter-spacing: -0.045em;
  white-space: nowrap;
  color: #eef5ff;
  text-shadow: 0 2px 10px rgba(0, 0, 0, 0.22);
  text-decoration: none;
}

.nav-brand-word .nav-brand-accent {
  color: var(--brand-blue-accent);
  text-shadow: 0 0 18px rgba(92, 169, 255, 0.22);
}

.brand {
  display: inline-flex;
  align-items: center;
  gap: 12px;
  text-decoration: none;
  color: var(--text);
}

.brand.brand-lockup {
  gap: 0.9rem;
}

.brand-media {
  display: inline-flex;
  align-items: center;
}

.brand-media.with-wordmark {
  gap: 0.8rem;
  width: auto;
  max-width: 100%;
}

.brand-logo-shell {
  width: 58px;
  flex: 0 0 58px;
  display: inline-flex;
  align-items: center;
}

.brand-lockup.nav .brand-media {
  width: clamp(118px, 16vw, 168px);
}

.brand-lockup.sidebar {
  width: 100%;
  flex-direction: column;
  align-items: flex-start;
  gap: 0.7rem;
}

.brand-lockup.sidebar .brand-media {
  width: auto;
  max-width: 100%;
}

.brand-lockup.footer .brand-media {
  width: min(100%, 180px);
}

.brand-logo {
  display: block;
  width: 100%;
  height: auto;
  filter: drop-shadow(0 14px 30px rgba(76, 138, 244, 0.2));
}

.brand-word-inline {
  font-family: "Sora", sans-serif;
  font-size: 1.28rem;
  font-weight: 700;
  letter-spacing: -0.04em;
  line-height: 1.05;
  white-space: nowrap;
  color: #eef5ff;
}

.brand-copy {
  display: flex;
  flex-direction: column;
  gap: 0.12rem;
}

.brand-caption {
  color: var(--muted);
  font-size: 0.82rem;
  line-height: 1.5;
}

.brand-mark {
  width: 46px;
  height: 46px;
  border-radius: 14px;
  position: relative;
  display: grid;
  place-items: center;
  background: linear-gradient(180deg, rgba(5, 10, 19, 0.96), rgba(2, 7, 17, 0.88));
  border: 1px solid rgba(112, 165, 255, 0.16);
  box-shadow: inset 0 1px 0 rgba(255, 255, 255, 0.06);
}

.brand-mark.small {
  width: 40px;
  height: 40px;
}

.brand-triangle {
  position: absolute;
  width: 18px;
  height: 21px;
  background: linear-gradient(180deg, rgba(155, 92, 255, 0.98), rgba(73, 134, 255, 0.9), rgba(72, 223, 255, 0.78));
  clip-path: polygon(50% 0%, 100% 100%, 0% 100%);
  top: 8px;
}

.brand-core {
  position: absolute;
  bottom: 6px;
  font-family: "Sora", sans-serif;
  font-size: 0.8rem;
  font-weight: 700;
  color: rgba(230, 241, 255, 0.84);
}

.brand-word {
  font-family: "Sora", sans-serif;
  font-weight: 700;
  font-size: 1.8rem;
  letter-spacing: -0.04em;
}

.ai-gradient {
  background: linear-gradient(90deg, var(--purple) 0%, var(--blue) 52%, var(--teal) 100%);
  background-size: 200% 200%;
  -webkit-background-clip: text;
  background-clip: text;
  color: transparent;
  animation: gradientShift 8s linear infinite;
}

.topbar-stack {
  display: flex;
  flex-direction: column;
  gap: 0.25rem;
}

.topbar-label {
  color: var(--cyan);
  font-size: 0.74rem;
  font-weight: 700;
  letter-spacing: 0.1em;
  text-transform: uppercase;
}

.topbar-page-title {
  font-family: "Sora", sans-serif;
  font-size: 1.15rem;
  letter-spacing: -0.03em;
}

.topbar-sub {
  color: var(--muted);
  font-size: 0.9rem;
}

.nav-pills,
.cta-row,
.stats-strip,
.trust-row,
.footer-meta,
.footer-grid,
.feature-grid,
.split-grid,
.quote-grid,
.pricing-grid,
.scenario-card-grid {
  display: flex;
}

.nav-pills {
  flex-wrap: wrap;
  justify-content: center;
  gap: 10px;
  padding: 10px 12px;
  border: 1px solid rgba(255, 255, 255, 0.05);
  border-radius: 999px;
  background: rgba(5, 12, 24, 0.54);
}

.nav-pills a,
.footer-grid a,
.footer-grid span,
.text-link {
  color: var(--muted);
  text-decoration: none;
}

.nav-link {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  padding: 0.72rem 0.8rem;
  border-radius: 14px;
  transition: background 160ms ease, color 160ms ease, transform 160ms ease;
  white-space: nowrap;
  font-size: 0.94rem;
}

.nav-link:hover,
.nav-link.active {
  color: var(--text);
  background: rgba(255, 255, 255, 0.08);
  transform: translateY(-1px);
}

.nav-actions {
  display: flex;
  gap: 10px;
  align-items: center;
  flex-wrap: wrap;
}

.text-link {
  font-weight: 600;
}

.btn {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-height: 48px;
  padding: 0 20px;
  border-radius: 16px;
  text-decoration: none;
  font-weight: 700;
  transition: transform 160ms ease, box-shadow 160ms ease, background 160ms ease;
  position: relative;
  overflow: hidden;
}

.btn:hover {
  transform: translateY(-1px);
}

.btn::after {
  content: "";
  position: absolute;
  inset: 0;
  background: linear-gradient(110deg, transparent 0%, rgba(255, 255, 255, 0.16) 46%, transparent 100%);
  transform: translateX(-140%);
  transition: transform 320ms ease;
}

.btn:hover::after {
  transform: translateX(140%);
}

.btn-primary {
  color: #fff;
  background: linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%);
  box-shadow: 0 16px 40px rgba(122, 67, 255, 0.28);
}

.btn-outline-cta {
  border: 1.5px solid transparent;
  color: #f5fbff;
  font-weight: 800;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box;
  box-shadow: none;
}

.btn-outline-cta:hover {
  box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.22);
}

.btn-secondary {
  color: var(--text);
  background: rgba(255, 255, 255, 0.04);
  border: 1px solid rgba(255, 255, 255, 0.08);
}

.hero-shell {
  padding-top: 1.6rem;
}

.hero-grid,
.page-grid,
.dashboard-hero-grid {
  display: grid;
  grid-template-columns: 1.2fr 0.8fr;
  gap: 22px;
  align-items: stretch;
}

.glass {
  border: 1px solid transparent;
  border-radius: var(--radius-xl);
  background:
    linear-gradient(180deg, rgba(10, 20, 34, 0.92), rgba(5, 12, 24, 0.74)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.34), rgba(100, 162, 255, 0.22), rgba(24, 209, 199, 0.24)) border-box;
  box-shadow: var(--shadow);
  position: relative;
  overflow: hidden;
}

.glass::before {
  content: "";
  position: absolute;
  inset: 0;
  background: linear-gradient(180deg, rgba(255, 255, 255, 0.045), transparent 24%);
  pointer-events: none;
}

.hero-card {
  padding: 2.4rem;
  animation: pulseGlow 9s ease-in-out infinite;
}

.kicker {
  display: inline-flex;
  align-items: center;
  padding: 0.34rem 0.7rem;
  border-radius: 999px;
  border: 1px solid rgba(118, 231, 255, 0.22);
  color: var(--cyan);
  font-size: 0.78rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

.home-hero-stage {
  width: min(100%, 1120px);
  margin-inline: auto;
  padding: clamp(2.8rem, 5vw, 5rem) 0 2.4rem;
  display: grid;
  justify-items: center;
  text-align: center;
}

.home-logo-stack {
  display: grid;
  justify-items: center;
  gap: 1.2rem;
}

.home-hero-ticker {
  width: min(100%, 1220px);
  margin: 0 auto 2rem;
}

.home-logo-frame {
  position: relative;
  width: min(100%, 268px);
  padding: 1.1rem 1.35rem;
  border-radius: 24px;
  background: linear-gradient(180deg, rgba(4, 5, 10, 0.99), rgba(1, 3, 8, 0.99));
  border: 1px solid rgba(122, 67, 255, 0.18);
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.03) inset,
    0 0 44px rgba(122, 67, 255, 0.38),
    0 28px 64px rgba(0, 0, 0, 0.34);
}

.home-logo-frame::before {
  content: "";
  position: absolute;
  inset: -18px;
  border-radius: 30px;
  background: radial-gradient(circle, rgba(122, 67, 255, 0.42) 0%, rgba(91, 157, 255, 0.16) 36%, rgba(122, 67, 255, 0) 74%);
  filter: blur(18px);
  z-index: -1;
}

.home-logo-image {
  display: block;
  width: 100%;
  height: auto;
}

.home-powered-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.6rem;
  min-height: 44px;
  padding: 0 1rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.035);
  color: var(--muted);
  font-size: 0.96rem;
}

.home-powered-glyph {
  color: #cfdcff;
  font-size: 0.95rem;
}

.home-powered-dot,
.home-live-dot {
  width: 10px;
  height: 10px;
  border-radius: 999px;
  background: #19d97c;
  box-shadow: 0 0 0 6px rgba(25, 217, 124, 0.12);
}

.home-hero-title {
  margin: 1.9rem 0 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(3.1rem, 7vw, 5.8rem);
  line-height: 0.98;
  letter-spacing: -0.075em;
}

.home-hero-title span {
  display: block;
}

.home-hero-emphasis {
  background: linear-gradient(90deg, #a7c0ff 0%, #72a0ff 24%, #4f88ff 56%, #2f76ff 100%);
  -webkit-background-clip: text;
  background-clip: text;
  color: transparent;
}

.home-hero-copy {
  max-width: 760px;
  margin: 1.45rem auto 0;
  color: var(--muted);
  font-size: clamp(1.12rem, 1.8vw, 1.55rem);
  line-height: 1.5;
}

.home-live-pill {
  margin-top: 1.7rem;
  display: inline-flex;
  align-items: center;
  gap: 0.65rem;
  min-height: 42px;
  padding: 0 1rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: var(--muted);
  font-size: 0.95rem;
}

.home-hero-actions {
  margin-top: 2.2rem;
  display: flex;
  justify-content: center;
  gap: 1rem;
  flex-wrap: wrap;
}

.home-hero-actions .btn {
  min-width: 190px;
  min-height: 54px;
  padding: 0 1.5rem;
}

.home-hero-actions .btn-secondary {
  background: rgba(255, 255, 255, 0.015);
}

.home-preview-shell {
  width: min(100%, 1040px);
  margin-top: 2.3rem;
  padding: 1rem;
  border-radius: 24px;
  border: 1px solid rgba(88, 226, 255, 0.55);
  background: rgba(7, 13, 24, 0.62);
  box-shadow: 0 0 0 1px rgba(122, 67, 255, 0.16), 0 18px 42px rgba(0, 0, 0, 0.22);
  position: relative;
  overflow: hidden;
}

.home-preview-shell::after {
  content: "";
  position: absolute;
  inset: 0;
  background: linear-gradient(90deg, transparent 0%, rgba(112, 220, 255, 0.16) 48%, transparent 100%);
  transform: translateX(-120%);
  animation: previewShellSweep 6.2s ease-in-out infinite;
  pointer-events: none;
}

.home-proof-strip {
  width: min(100%, 1120px);
  margin-top: 2.15rem;
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 1.4rem;
}

.home-proof-item {
  text-align: center;
}

.home-proof-value {
  font-family: "Sora", sans-serif;
  font-size: clamp(2.4rem, 4vw, 3.25rem);
  font-weight: 700;
  line-height: 1;
  letter-spacing: -0.06em;
  color: #6e7cff;
}

.home-proof-label {
  margin-top: 0.65rem;
  color: #eef5ff;
  font-size: 1.08rem;
  font-weight: 700;
}

.home-proof-copy {
  margin-top: 0.35rem;
  color: var(--muted);
  font-size: 0.98rem;
  line-height: 1.45;
}

.home-intelligence-shell {
  width: min(100%, 1220px);
  margin-top: 3.25rem;
  display: grid;
  gap: 1.8rem;
  justify-items: center;
}

.home-intelligence-head {
  max-width: 760px;
  display: grid;
  gap: 1rem;
  justify-items: center;
  text-align: center;
}

.home-intelligence-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
  min-height: 42px;
  padding: 0 1rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.035);
  color: #b9c6dd;
  font-size: 0.95rem;
}

.home-intelligence-pill-glyph {
  color: #dbe7ff;
  font-size: 0.92rem;
}

.home-intelligence-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(2.3rem, 4.8vw, 3.55rem);
  line-height: 1.08;
  letter-spacing: -0.06em;
}

.home-intelligence-copy {
  margin: 0;
  color: var(--muted);
  font-size: clamp(1.02rem, 1.7vw, 1.22rem);
  line-height: 1.6;
}

.home-intelligence-grid {
  width: 100%;
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 1.45rem;
}

.home-intelligence-card {
  min-height: 214px;
  padding: 1.55rem;
  border-radius: 22px;
  border: 1px solid rgba(82, 218, 255, 0.65);
  background: rgba(5, 9, 15, 0.68);
  box-shadow:
    0 0 0 1px rgba(122, 67, 255, 0.22),
    0 18px 36px rgba(0, 0, 0, 0.24);
  text-align: left;
  transition: transform 180ms ease, box-shadow 180ms ease, border-color 180ms ease;
}

.home-intelligence-card:hover {
  transform: translateY(-4px);
  box-shadow:
    0 0 0 1px rgba(122, 67, 255, 0.28),
    0 24px 46px rgba(0, 0, 0, 0.28);
}

.home-intelligence-icon-shell {
  width: 42px;
  height: 42px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 14px;
  background: radial-gradient(circle at 30% 25%, rgba(108, 71, 255, 0.28), rgba(35, 56, 110, 0.2) 70%);
  box-shadow: 0 0 18px rgba(104, 88, 255, 0.18);
  color: #eef5ff;
  font-size: 1.12rem;
}

.home-intelligence-card h3 {
  margin: 1.15rem 0 0.8rem;
  font-family: "Sora", sans-serif;
  font-size: 1.05rem;
  line-height: 1.35;
}

.home-intelligence-card p {
  margin: 0;
  color: var(--muted);
  font-size: 0.98rem;
  line-height: 1.6;
}

.home-testimonial-shell {
  width: min(100%, 1220px);
  margin: 0 auto 2.4rem;
}

.home-testimonial-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 1.35rem;
}

.home-testimonial-card {
  min-height: 176px;
  padding: 1.55rem 1.6rem;
  border-radius: 22px;
  border: 1px solid rgba(82, 218, 255, 0.62);
  background: rgba(6, 10, 16, 0.72);
  box-shadow:
    0 0 0 1px rgba(122, 67, 255, 0.18),
    0 18px 36px rgba(0, 0, 0, 0.22);
  text-align: left;
}

.home-testimonial-quote {
  margin: 0;
  color: #b2bccd;
  font-size: 0.98rem;
  line-height: 1.65;
  font-style: italic;
}

.home-testimonial-author {
  margin-top: 1.15rem;
  display: grid;
  gap: 0.15rem;
}

.home-testimonial-author strong {
  font-family: "Sora", sans-serif;
  font-size: 1.02rem;
  color: #eef5ff;
}

.home-testimonial-author span {
  color: var(--muted);
  font-size: 0.98rem;
}

.home-cta-shell {
  width: min(100%, 980px);
  margin: 0 auto 2.8rem;
  display: grid;
  gap: 1.2rem;
  justify-items: center;
  text-align: center;
}

.home-cta-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.5rem;
  min-height: 42px;
  padding: 0 1rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.035);
  color: #b9c6dd;
  font-size: 0.95rem;
}

.home-cta-pill-glyph {
  color: #dbe7ff;
  font-size: 0.92rem;
}

.home-cta-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(2.1rem, 4.8vw, 3.55rem);
  line-height: 1.08;
  letter-spacing: -0.06em;
}

.home-cta-copy {
  margin: 0;
  color: var(--muted);
  font-size: clamp(1.02rem, 1.7vw, 1.22rem);
  line-height: 1.55;
}

.home-cta-shell .btn {
  min-width: 190px;
}

.home-preview-grid,
.home-alert-item,
.home-timeline-row,
.home-chart-bars {
  display: flex;
}

.home-preview-grid {
  gap: 1rem;
}

.home-preview-card {
  flex: 1 1 0;
  min-height: 210px;
  padding: 1rem 1rem 0.95rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(10, 16, 28, 0.7);
  text-align: left;
  position: relative;
  overflow: hidden;
  animation: previewCardFloat 6.4s ease-in-out infinite;
}

.home-preview-card:nth-child(2) {
  animation-delay: 0.75s;
}

.home-preview-card:nth-child(3) {
  animation-delay: 1.5s;
}

.home-preview-head {
  display: flex;
  align-items: center;
  gap: 0.55rem;
  color: var(--muted);
  font-size: 0.97rem;
}

.home-preview-icon {
  color: #d8e6fb;
  font-size: 0.92rem;
}

.home-chart-wrap {
  margin-top: 1.1rem;
  padding: 0.9rem 0.9rem 0.8rem;
  border-radius: 16px;
  background: rgba(6, 11, 21, 0.72);
  border: 1px solid rgba(255, 255, 255, 0.05);
}

.home-chart-bars {
  align-items: flex-end;
  gap: 0.38rem;
  min-height: 112px;
}

.home-chart-bar {
  flex: 1 1 0;
  min-height: 12px;
  border-radius: 8px 8px 4px 4px;
  background: linear-gradient(180deg, #2fe8ff 0%, #376dff 100%);
  box-shadow: 0 0 14px rgba(56, 170, 255, 0.22);
  transform-origin: center bottom;
  animation: chartBarLift 2.8s ease-in-out infinite;
}

.home-chart-bar:nth-child(1) { animation-delay: 0.05s; }
.home-chart-bar:nth-child(2) { animation-delay: 0.12s; }
.home-chart-bar:nth-child(3) { animation-delay: 0.19s; }
.home-chart-bar:nth-child(4) { animation-delay: 0.26s; }
.home-chart-bar:nth-child(5) { animation-delay: 0.33s; }
.home-chart-bar:nth-child(6) { animation-delay: 0.40s; }
.home-chart-bar:nth-child(7) { animation-delay: 0.47s; }
.home-chart-bar:nth-child(8) { animation-delay: 0.54s; }
.home-chart-bar:nth-child(9) { animation-delay: 0.61s; }
.home-chart-bar:nth-child(10) { animation-delay: 0.68s; }
.home-chart-bar:nth-child(odd) {
  animation-duration: 3.15s;
}

.home-chart-axis {
  margin-top: 0.8rem;
  height: 1px;
  background: rgba(102, 208, 255, 0.45);
}

.home-chart-caption {
  margin-top: 0.7rem;
  color: var(--muted);
  font-size: 0.96rem;
}

.home-alert-stack {
  display: grid;
  gap: 0.7rem;
  margin-top: 1rem;
}

.home-alert-item {
  align-items: center;
  gap: 0.75rem;
  padding: 0.75rem 0.85rem;
  border-radius: 14px;
  border: 1px solid transparent;
  position: relative;
  overflow: hidden;
  animation: alertPulse 4.8s ease-in-out infinite;
}

.home-alert-item::after {
  content: "";
  position: absolute;
  inset: 0;
  background: linear-gradient(100deg, transparent 0%, rgba(255, 255, 255, 0.12) 48%, transparent 100%);
  transform: translateX(-130%);
  animation: alertSweep 5.4s ease-in-out infinite;
  pointer-events: none;
}

.home-alert-item.high {
  background: rgba(111, 24, 24, 0.42);
  border-color: rgba(239, 68, 68, 0.25);
}

.home-alert-item.medium {
  background: rgba(94, 58, 9, 0.38);
  border-color: rgba(245, 158, 11, 0.22);
}

.home-alert-item.low {
  background: rgba(8, 80, 50, 0.34);
  border-color: rgba(52, 211, 153, 0.22);
}

.home-alert-item.medium {
  animation-delay: 0.8s;
}

.home-alert-item.low {
  animation-delay: 1.6s;
}

.home-alert-dot {
  width: 9px;
  height: 9px;
  border-radius: 999px;
  flex: 0 0 auto;
  animation: alertDotBlink 2.4s ease-in-out infinite;
}

.home-alert-item.high .home-alert-dot {
  background: #ef4444;
}

.home-alert-item.medium .home-alert-dot {
  background: #f59e0b;
}

.home-alert-item.low .home-alert-dot {
  background: #34d399;
}

.home-alert-copy {
  color: #e8f1ff;
  font-size: 0.98rem;
  line-height: 1.32;
}

.home-timeline-stack {
  display: grid;
  gap: 0.9rem;
  margin-top: 1rem;
}

.home-timeline-row {
  flex-direction: column;
  gap: 0.38rem;
}

.home-timeline-top {
  display: flex;
  justify-content: space-between;
  gap: 0.8rem;
  color: #e8f1ff;
  font-size: 0.97rem;
}

.home-timeline-bar {
  height: 8px;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.09);
  overflow: hidden;
}

.home-timeline-bar span {
  display: block;
  height: 100%;
  border-radius: inherit;
  background: linear-gradient(90deg, #3e7bff 0%, #2fe8ff 100%);
  background-size: 200% 100%;
  box-shadow: 0 0 16px rgba(47, 232, 255, 0.18);
  animation: timelineFlow 4.6s linear infinite;
}

.hero-title,
.page-title {
  margin: 1rem 0 0.8rem;
  font-family: "Sora", sans-serif;
  font-size: clamp(2.6rem, 5vw, 4.2rem);
  line-height: 1.02;
  letter-spacing: -0.06em;
}

.page-title {
  font-size: clamp(2rem, 4vw, 3rem);
}

.hero-subtitle,
.page-copy,
.section-copy,
.eyebrow-copy {
  color: var(--muted);
  font-size: 1.04rem;
  line-height: 1.75;
}

.cta-row {
  flex-wrap: wrap;
  gap: 14px;
  margin-top: 1.6rem;
}

.stats-strip {
  flex-wrap: wrap;
  gap: 12px;
  margin-top: 1.6rem;
}

.ticker-shell {
  margin-top: 1rem;
  padding: 0.95rem 0;
}

.ticker-mask {
  overflow: hidden;
  white-space: nowrap;
}

.ticker-track {
  display: inline-flex;
  min-width: max-content;
  gap: 1rem;
  animation: tickerSlide 20s linear infinite;
}

.ticker-item {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
  padding: 0.55rem 0.95rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.035);
  color: var(--muted);
  font-size: 0.92rem;
}

.ticker-item::before {
  content: "";
  width: 0.55rem;
  height: 0.55rem;
  border-radius: 999px;
  background: linear-gradient(180deg, var(--purple), var(--teal));
  box-shadow: 0 0 0 6px rgba(24, 209, 199, 0.08);
}

.stat-pill,
.mini-pill,
.status-chip,
.plan-chip {
  display: inline-flex;
  align-items: center;
  gap: 8px;
  padding: 0.72rem 0.95rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.07);
  background: rgba(255, 255, 255, 0.03);
}

.stat-pill strong,
.metric-card strong,
.workspace-card strong,
.timeline-card strong,
.scenario-mini-card strong {
  display: block;
  font-family: "Sora", sans-serif;
}

.insight-card {
  padding: 2rem;
}

.insight-stack {
  display: grid;
  gap: 14px;
}

.hero-visual {
  padding: 2rem;
  display: grid;
  gap: 1.1rem;
}

.hero-visual::after {
  content: "";
  position: absolute;
  inset: -20% 0 auto;
  height: 32%;
  background: linear-gradient(180deg, rgba(118, 231, 255, 0.22), rgba(118, 231, 255, 0));
  mix-blend-mode: screen;
  animation: scanDrift 8.5s ease-in-out infinite;
  pointer-events: none;
}

.hero-visual-head,
.hero-health-grid,
.hero-driver-item,
.brand-marquee,
.compliance-band,
.feature-card-top,
.process-card-head {
  display: flex;
}

.hero-visual-head {
  justify-content: space-between;
  gap: 1rem;
  align-items: center;
  flex-wrap: wrap;
}

.hero-health-grid {
  gap: 12px;
  flex-wrap: wrap;
}

.hero-health-card {
  flex: 1 1 145px;
  padding: 1rem;
  border-radius: 20px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.03);
}

.hero-health-card span,
.metric-showcase-card span,
.process-copy,
.logo-chip,
.feature-caption {
  color: var(--muted);
}

.hero-health-card strong {
  display: block;
  margin-top: 0.45rem;
  font-family: "Sora", sans-serif;
  font-size: 1.55rem;
  letter-spacing: -0.04em;
}

.hero-progress {
  display: grid;
  gap: 0.8rem;
}

.hero-progress-row {
  display: grid;
  gap: 0.4rem;
}

.hero-progress-top {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 0.75rem;
  color: var(--muted);
  font-size: 0.92rem;
}

.hero-progress-bar {
  height: 12px;
  border-radius: 999px;
  background: rgba(255, 255, 255, 0.06);
  overflow: hidden;
}

.hero-progress-bar span {
  display: block;
  height: 100%;
  border-radius: inherit;
  background: linear-gradient(90deg, rgba(155, 92, 255, 0.92), rgba(100, 162, 255, 0.92), rgba(24, 209, 199, 0.92));
  box-shadow: 0 10px 28px rgba(76, 138, 244, 0.2);
}

.hero-driver-list {
  display: grid;
  gap: 0.75rem;
}

.hero-driver-item {
  justify-content: space-between;
  gap: 1rem;
  align-items: center;
  padding: 0.8rem 0.95rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.025);
}

.hero-driver-copy strong,
.process-title,
.metric-showcase-card strong {
  display: block;
  font-family: "Sora", sans-serif;
}

.hero-driver-copy p {
  margin: 0.25rem 0 0;
  color: var(--muted);
  font-size: 0.9rem;
}

.info-block,
.feature-card,
.quote-card,
.pricing-card,
.detail-card,
.workspace-card,
.timeline-card,
.scenario-mini-card,
.history-card {
  border: 1px solid transparent;
  border-radius: 22px;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.03), rgba(255, 255, 255, 0.025)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.22), rgba(100, 162, 255, 0.12), rgba(24, 209, 199, 0.16)) border-box;
}

.info-block,
.workspace-card,
.timeline-card,
.scenario-mini-card,
.history-card,
.detail-card {
  padding: 1.2rem 1.25rem;
}

.section-label {
  margin: 2.4rem 0 0.8rem;
  color: var(--cyan);
  font-size: 0.82rem;
  font-weight: 700;
  letter-spacing: 0.1em;
  text-transform: uppercase;
}

.section-heading {
  margin: 0 0 0.55rem;
  font-family: "Sora", sans-serif;
  font-size: clamp(1.6rem, 3vw, 2.25rem);
  letter-spacing: -0.04em;
}

.section-copy {
  margin-top: 0;
}

.feature-grid,
.quote-grid,
.pricing-grid,
.scenario-card-grid {
  flex-wrap: wrap;
  gap: 16px;
}

.feature-card,
.quote-card,
.pricing-card {
  flex: 1 1 220px;
  padding: 1.4rem;
  transition: transform 180ms ease, box-shadow 180ms ease, border-color 180ms ease;
}

.feature-card:hover,
.quote-card:hover,
.pricing-card:hover,
.detail-card:hover,
.history-card:hover,
.workspace-card:hover {
  transform: translateY(-4px);
  box-shadow: 0 28px 60px rgba(0, 0, 0, 0.24);
}

.feature-card h3,
.quote-card h3,
.pricing-card h3,
.detail-card h3 {
  margin: 0.5rem 0;
  font-family: "Sora", sans-serif;
  font-size: 1.05rem;
}

.feature-card p,
.quote-card p,
.pricing-card p,
.detail-card p,
.detail-card li,
.info-block p {
  color: var(--muted);
  line-height: 1.7;
}

.badge-dot {
  width: 12px;
  height: 12px;
  border-radius: 999px;
  background: linear-gradient(180deg, var(--teal), var(--blue));
  box-shadow: 0 0 0 8px rgba(25, 210, 201, 0.09);
}

.quote-card span,
.pricing-tag,
.eyebrow {
  color: var(--cyan);
  font-weight: 700;
  font-size: 0.86rem;
}

.pricing-hero-card {
  margin-bottom: 1.2rem;
}

.pricing-showcase-shell {
  padding-bottom: 3.2rem;
}

.pricing-showcase {
  align-items: stretch;
  justify-content: center;
  gap: 22px;
}

.pricing-showcase .pricing-plan-card {
  position: relative;
  display: flex;
  flex-direction: column;
  flex: 1 1 300px;
  max-width: 358px;
  min-height: 540px;
  padding: 2rem 2rem 1.9rem;
  border-radius: 24px;
  border: 1px solid rgba(104, 117, 161, 0.22);
  background: linear-gradient(180deg, rgba(10, 13, 19, 0.98), rgba(8, 11, 16, 0.98));
  box-shadow:
    0 14px 32px rgba(0, 0, 0, 0.34),
    inset 0 0 0 1px rgba(255, 255, 255, 0.02);
  gap: 1.2rem;
}

.pricing-showcase .pricing-plan-card:hover {
  transform: translateY(-4px);
  box-shadow:
    0 22px 44px rgba(0, 0, 0, 0.42),
    inset 0 0 0 1px rgba(255, 255, 255, 0.03);
}

.pricing-showcase .pricing-plan-card-featured {
  border-color: rgba(62, 133, 255, 0.78);
  box-shadow:
    0 0 0 1px rgba(59, 131, 255, 0.18),
    0 18px 42px rgba(10, 26, 61, 0.52),
    0 0 34px rgba(46, 130, 255, 0.18);
}

.pricing-showcase .pricing-plan-card-featured:hover {
  box-shadow:
    0 0 0 1px rgba(73, 145, 255, 0.24),
    0 22px 48px rgba(10, 26, 61, 0.58),
    0 0 42px rgba(46, 130, 255, 0.24);
}

.pricing-plan-badge {
  position: absolute;
  top: -14px;
  left: 50%;
  transform: translateX(-50%);
  padding: 0.45rem 1rem;
  border-radius: 999px;
  background: linear-gradient(90deg, rgba(133, 105, 255, 0.98), rgba(42, 191, 234, 0.98));
  color: #ffffff;
  font-size: 0.88rem;
  font-weight: 700;
  letter-spacing: -0.01em;
  box-shadow: 0 10px 24px rgba(50, 112, 255, 0.28);
}

.pricing-plan-head {
  display: flex;
  flex-direction: column;
  gap: 0.95rem;
}

.pricing-showcase .pricing-plan-name {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: 1.05rem;
  color: rgba(246, 249, 255, 0.96);
}

.pricing-price-row {
  display: flex;
  align-items: flex-end;
  gap: 0.38rem;
}

.pricing-price-row-custom {
  align-items: center;
}

.pricing-price-value {
  font-family: "Sora", sans-serif;
  font-size: clamp(3rem, 4vw, 3.75rem);
  line-height: 0.95;
  font-weight: 700;
  letter-spacing: -0.06em;
  color: rgba(247, 249, 255, 0.98);
}

.pricing-price-value-custom {
  font-size: clamp(2.85rem, 3.8vw, 3.4rem);
}

.pricing-price-suffix {
  margin-bottom: 0.38rem;
  color: rgba(183, 192, 211, 0.88);
  font-size: 1rem;
  font-weight: 500;
}

.pricing-plan-copy {
  margin: 0;
  color: rgba(180, 188, 205, 0.92) !important;
  line-height: 1.5;
}

.pricing-feature-list {
  display: flex;
  flex-direction: column;
  gap: 0.9rem;
  margin: 0;
  padding: 0.4rem 0 0;
  list-style: none;
}

.pricing-feature-list li {
  position: relative;
  padding-left: 1.5rem;
  color: rgba(236, 241, 249, 0.96);
  line-height: 1.45;
}

.pricing-feature-list li::before {
  content: "✓";
  position: absolute;
  left: 0;
  top: 0;
  color: rgba(232, 224, 255, 0.98);
  font-weight: 800;
}

.pricing-plan-cta {
  margin-top: auto;
  display: flex;
  min-height: 52px;
  align-items: center;
  justify-content: center;
  border-radius: 14px;
  text-decoration: none !important;
  font-weight: 700;
  font-size: 1rem;
  transition: transform 180ms ease, box-shadow 180ms ease, border-color 180ms ease;
}

.pricing-plan-cta:hover {
  transform: translateY(-1px);
}

.pricing-plan-cta-muted {
  border: 1px solid transparent;
  background:
    linear-gradient(180deg, rgba(18, 21, 31, 0.98), rgba(11, 14, 22, 0.98)) padding-box,
    linear-gradient(90deg, rgba(129, 89, 255, 0.98), rgba(42, 191, 234, 0.98)) border-box;
  color: rgba(242, 246, 255, 0.96);
  box-shadow:
    0 0 0 1px rgba(92, 82, 255, 0.14),
    0 10px 22px rgba(8, 12, 26, 0.32),
    0 0 18px rgba(55, 138, 255, 0.12);
}

.pricing-plan-cta-muted:hover {
  box-shadow:
    0 0 0 1px rgba(92, 82, 255, 0.22),
    0 14px 26px rgba(8, 12, 26, 0.38),
    0 0 24px rgba(55, 138, 255, 0.18);
}

.pricing-plan-cta-primary {
  background: linear-gradient(90deg, rgba(70, 117, 255, 0.98), rgba(46, 191, 233, 0.98));
  color: #ffffff;
  box-shadow: 0 16px 34px rgba(30, 118, 255, 0.28);
}

.pricing-plan-cta-primary:hover {
  box-shadow: 0 20px 38px rgba(30, 118, 255, 0.34);
}

.metric-showcase-grid,
.process-grid {
  display: grid;
  gap: 16px;
}

.metric-showcase-grid {
  margin-top: 1.2rem;
  grid-template-columns: repeat(4, minmax(0, 1fr));
}

.metric-showcase-card {
  padding: 1.15rem 1.2rem;
  border-radius: 22px;
  border: 1px solid transparent;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.035), rgba(255, 255, 255, 0.025)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.24), rgba(100, 162, 255, 0.18), rgba(24, 209, 199, 0.2)) border-box;
}

.metric-showcase-card strong {
  margin-top: 0.5rem;
  font-size: clamp(1.7rem, 3vw, 2.3rem);
  letter-spacing: -0.06em;
}

.metric-showcase-card small {
  display: block;
  margin-top: 0.65rem;
  color: var(--cyan);
  font-size: 0.88rem;
}

.brand-marquee,
.compliance-band {
  flex-wrap: wrap;
  gap: 12px;
}

.logo-chip {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  min-height: 54px;
  padding: 0 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.03);
  font-family: "Sora", sans-serif;
  font-size: 0.95rem;
}

.process-grid {
  grid-template-columns: repeat(3, minmax(0, 1fr));
}

.process-card {
  padding: 1.45rem;
  border-radius: 24px;
  border: 1px solid transparent;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.035), rgba(255, 255, 255, 0.025)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.26), rgba(100, 162, 255, 0.16), rgba(24, 209, 199, 0.18)) border-box;
}

.process-card-head {
  justify-content: space-between;
  align-items: center;
  gap: 0.8rem;
}

.step-badge {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 3rem;
  height: 3rem;
  border-radius: 1rem;
  background: linear-gradient(135deg, rgba(155, 92, 255, 0.18), rgba(24, 209, 199, 0.16));
  border: 1px solid rgba(118, 231, 255, 0.18);
  color: var(--text);
  font-family: "Sora", sans-serif;
  font-weight: 700;
}

.process-title {
  font-size: 1.1rem;
}

.process-copy {
  margin-top: 0.85rem;
  line-height: 1.7;
}

.feature-card-top {
  align-items: center;
  gap: 0.8rem;
}

.feature-icon-shell {
  width: 2.8rem;
  height: 2.8rem;
  border-radius: 16px;
  background: linear-gradient(135deg, rgba(155, 92, 255, 0.24), rgba(24, 209, 199, 0.2));
  border: 1px solid rgba(118, 231, 255, 0.15);
  display: inline-flex;
  align-items: center;
  justify-content: center;
  font-size: 1.15rem;
}

.feature-caption {
  font-size: 0.86rem;
}

.page-shell {
  padding-top: 1.2rem;
}

.page-hero-card {
  padding: 2rem 2.1rem;
}

.about-proto-shell {
  padding-top: 1.55rem;
}

.about-top-stage {
  width: min(100%, 1218px);
  margin-inline: auto;
}

.about-proto-hero,
.about-journey-section,
.about-final-cta {
  text-align: center;
}

.about-proto-hero {
  display: grid;
  justify-items: center;
  gap: 1.45rem;
  margin-bottom: 4.05rem;
}

.about-proto-logo-shell {
  width: clamp(258px, 25.4vw, 296px);
  height: clamp(160px, 17.1vw, 184px);
  display: grid;
  place-items: center;
  border-radius: 6px;
  border: 1px solid rgba(137, 109, 255, 0.16);
  background:
    linear-gradient(180deg, rgba(2, 5, 11, 0.995), rgba(2, 5, 11, 0.995));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 0 64px rgba(122, 67, 255, 0.42),
    0 30px 72px rgba(0, 0, 0, 0.34);
}

.about-proto-logo-shell img {
  width: 75%;
  height: auto;
  display: block;
}

.about-proto-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(3.3rem, 4.9vw, 4.6rem);
  line-height: 0.97;
  letter-spacing: -0.06em;
}

.about-proto-subtitle {
  max-width: 900px;
  margin: 0;
  color: rgba(232, 241, 255, 0.76);
  font-size: 1.08rem;
  line-height: 1.72;
}

.about-mv-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 1.95rem;
  margin-bottom: 4.25rem;
}

.about-mv-card {
  border: 1px solid transparent;
  border-radius: 18px;
  min-height: 280px;
  padding: 2rem 2rem 1.9rem;
  text-align: left;
  background:
    linear-gradient(180deg, rgba(7, 10, 16, 0.985), rgba(6, 9, 15, 0.995)) padding-box,
    linear-gradient(135deg, rgba(129, 90, 255, 0.62), rgba(70, 155, 255, 0.48), rgba(55, 214, 255, 0.5)) border-box;
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 0 0 1px rgba(51, 211, 255, 0.08),
    0 18px 34px rgba(0, 0, 0, 0.2);
}

.about-card-icon,
.about-value-icon {
  width: 3.05rem;
  height: 3.05rem;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 15px;
  border: 1px solid rgba(116, 231, 255, 0.14);
  background:
    linear-gradient(180deg, rgba(36, 47, 105, 0.92), rgba(18, 30, 68, 0.97));
  color: #eef5ff;
  font-size: 1.05rem;
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.01) inset,
    0 0 26px rgba(92, 169, 255, 0.24);
}

.about-card-icon svg,
.about-value-icon svg {
  width: 1.24rem;
  height: 1.24rem;
  display: block;
  stroke: currentColor;
  fill: none;
  stroke-width: 1.8;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.about-mv-card h3,
.about-timeline-card h3 {
  margin: 1.32rem 0 0.88rem;
  font-size: 1.2rem;
  letter-spacing: -0.04em;
}

.about-mv-card p,
.about-timeline-card p,
.about-member-card p,
.about-value-card p {
  margin: 0;
  color: rgba(232, 241, 255, 0.78);
  line-height: 1.76;
}

.about-values-section {
  width: min(100%, 1220px);
  margin-inline: auto;
  text-align: center;
  margin-bottom: 5.15rem;
}

.about-values-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
  min-height: 40px;
  padding: 0 0.95rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: rgba(232, 241, 255, 0.72);
  font-size: 0.84rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

.about-values-pill svg {
  width: 0.9rem;
  height: 0.9rem;
  display: block;
  stroke: currentColor;
  fill: none;
  stroke-width: 1.8;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.about-section-heading {
  margin: 1rem 0 0.7rem;
  font-family: "Sora", sans-serif;
  font-size: clamp(2.35rem, 3.4vw, 3.45rem);
  line-height: 1.02;
  letter-spacing: -0.06em;
}

.about-section-copy {
  max-width: 640px;
  margin: 0 auto;
  color: rgba(232, 241, 255, 0.72);
}

.about-values-grid {
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 1.2rem;
  margin-top: 2.15rem;
}

.about-value-card {
  min-height: 210px;
  padding: 1.7rem 1.25rem 1.5rem;
  text-align: center;
  border: 1px solid rgba(255, 255, 255, 0.08);
  border-radius: 18px;
  background: linear-gradient(180deg, rgba(9, 14, 24, 0.9), rgba(6, 10, 18, 0.96));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.015) inset,
    0 18px 34px rgba(0, 0, 0, 0.16);
}

.about-value-card h3 {
  margin: 1.15rem 0 0.55rem;
  font-size: 1.18rem;
  letter-spacing: -0.04em;
}

.about-journey-section {
  width: min(100%, 1220px);
  margin: 0 auto 5.1rem;
}

.about-timeline {
  position: relative;
  width: min(100%, 1000px);
  margin: 2.55rem auto 0;
}

.about-timeline::before {
  content: "";
  position: absolute;
  top: 0;
  bottom: 0;
  left: 50%;
  width: 1px;
  transform: translateX(-50%);
  background: linear-gradient(180deg, rgba(255, 255, 255, 0.04), rgba(92, 169, 255, 0.46), rgba(255, 255, 255, 0.04));
  opacity: 0.7;
}

.about-timeline-row {
  position: relative;
  display: grid;
  grid-template-columns: minmax(0, 1fr) 64px minmax(0, 1fr);
  gap: 1.35rem;
  align-items: center;
  margin-bottom: 1.5rem;
}

.about-timeline-row:last-child {
  margin-bottom: 0;
}

.about-timeline-side {
  min-width: 0;
}

.about-timeline-axis {
  position: relative;
  min-height: 100%;
  display: grid;
  place-items: center;
}

.about-timeline-dot {
  width: 12px;
  height: 12px;
  border-radius: 999px;
  background: linear-gradient(180deg, #67d8ff, #3b8cff);
  box-shadow:
    0 0 0 7px rgba(47, 162, 255, 0.08),
    0 0 24px rgba(92, 169, 255, 0.42);
}

.about-timeline-card {
  position: relative;
  width: min(100%, 360px);
  min-height: 132px;
  padding: 1.4rem 1.45rem 1.3rem;
  text-align: left;
  border: 1px solid transparent;
  border-radius: 20px;
  overflow: hidden;
  isolation: isolate;
  background:
    linear-gradient(180deg, rgba(8, 13, 24, 0.98), rgba(5, 9, 18, 0.96)) padding-box,
    linear-gradient(135deg, rgba(138, 91, 255, 0.62), rgba(70, 155, 255, 0.34), rgba(52, 215, 255, 0.18)) border-box;
  box-shadow:
    0 24px 52px rgba(0, 0, 0, 0.26),
    0 0 0 1px rgba(255, 255, 255, 0.03) inset,
    0 0 34px rgba(73, 132, 255, 0.08);
  transition: transform 220ms ease, box-shadow 220ms ease;
}

.about-timeline-card::before,
.about-timeline-card::after {
  content: "";
  position: absolute;
  border-radius: 999px;
  filter: blur(10px);
  z-index: 0;
}

.about-timeline-card::before {
  width: 10rem;
  height: 10rem;
  top: -4rem;
  right: -3.5rem;
  background: radial-gradient(circle, rgba(132, 96, 255, 0.24), transparent 70%);
}

.about-timeline-card::after {
  width: 8rem;
  height: 8rem;
  left: -2.6rem;
  bottom: -3.2rem;
  background: radial-gradient(circle, rgba(74, 156, 255, 0.16), transparent 72%);
}

.about-timeline-card > * {
  position: relative;
  z-index: 1;
}

.about-timeline-card:hover {
  transform: translateY(-4px) scale(1.02);
  box-shadow:
    0 30px 68px rgba(0, 0, 0, 0.32),
    0 0 0 1px rgba(173, 198, 255, 0.07) inset,
    0 0 48px rgba(96, 121, 255, 0.18);
}

.about-timeline-row.left .about-timeline-card {
  margin-left: auto;
  text-align: right;
}

.about-timeline-row.right .about-timeline-card {
  margin-right: auto;
}

.about-timeline-year {
  display: block;
  margin-bottom: 0.56rem;
  color: #8cc3ff;
  font-size: 0.76rem;
  font-weight: 800;
  letter-spacing: 0.16em;
  text-transform: uppercase;
}

.about-journey-step-title {
  margin: 0 0 0.58rem;
  font-size: 1.18rem;
  line-height: 1.2;
  letter-spacing: -0.04em;
  font-weight: 800;
}

.about-team-section {
  width: min(100%, 1220px);
  margin-inline: auto;
  text-align: left;
  margin-bottom: 5rem;
}

.about-team-head {
  display: flex;
  align-items: center;
  gap: 0.8rem;
  margin-bottom: 1.9rem;
}

.about-team-icon {
  width: 1.55rem;
  height: 1.55rem;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  color: rgba(238, 245, 255, 0.92);
}

.about-team-icon svg {
  width: 100%;
  height: 100%;
  display: block;
  stroke: currentColor;
  fill: none;
  stroke-width: 1.8;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.about-team-head .about-section-heading {
  margin: 0;
  font-size: clamp(2.05rem, 2.9vw, 2.85rem);
}

.about-team-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 1.15rem;
  margin-top: 0;
  align-items: stretch;
}

.about-member-card {
  position: relative;
  display: flex;
  flex-direction: column;
  justify-content: flex-start;
  align-items: center;
  min-height: 100%;
  padding: 1.55rem 1.45rem 1.5rem;
  text-align: center;
  border: 1px solid transparent;
  border-radius: 20px;
  overflow: hidden;
  isolation: isolate;
  background:
    linear-gradient(180deg, rgba(8, 13, 24, 0.98), rgba(5, 9, 18, 0.96)) padding-box,
    linear-gradient(135deg, rgba(138, 91, 255, 0.62), rgba(70, 155, 255, 0.34), rgba(52, 215, 255, 0.18)) border-box;
  box-shadow:
    0 26px 52px rgba(0, 0, 0, 0.28),
    0 0 0 1px rgba(255, 255, 255, 0.03) inset,
    0 0 34px rgba(73, 132, 255, 0.08);
  transition: transform 220ms ease, box-shadow 220ms ease;
}

.about-member-card::before,
.about-member-card::after {
  content: "";
  position: absolute;
  border-radius: 999px;
  filter: blur(10px);
  z-index: 0;
}

.about-member-card::before {
  width: 11rem;
  height: 11rem;
  top: -4.5rem;
  right: -4rem;
  background: radial-gradient(circle, rgba(132, 96, 255, 0.24), transparent 70%);
}

.about-member-card::after {
  width: 9rem;
  height: 9rem;
  left: -3rem;
  bottom: -4rem;
  background: radial-gradient(circle, rgba(74, 156, 255, 0.18), transparent 72%);
}

.about-member-card > * {
  position: relative;
  z-index: 1;
}

.about-member-card:hover {
  transform: translateY(-4px) scale(1.02);
  box-shadow:
    0 32px 68px rgba(0, 0, 0, 0.34),
    0 0 0 1px rgba(173, 198, 255, 0.07) inset,
    0 0 46px rgba(96, 121, 255, 0.18);
}

.avatar-wrapper {
  width: 100%;
  display: flex;
  justify-content: center;
  margin-bottom: 16px;
}

.team-avatar {
  width: 160px;
  height: 160px;
  border-radius: 50%;
  overflow: hidden;
  position: relative;
  background: #111;
}

.team-avatar img {
  width: 100%;
  height: 100%;
  object-fit: cover;
  object-position: center 30%;
  display: block;
  filter: brightness(0.85);
}

.team-avatar::after {
  content: "";
  position: absolute;
  inset: 0;
  background: rgba(0, 0, 0, 0.25);
  pointer-events: none;
}

.about-member-card h3 {
  margin: 0 0 0.5rem;
  font-size: 1.17rem;
  letter-spacing: -0.04em;
  background: linear-gradient(90deg, var(--purple) 0%, var(--blue) 52%, var(--teal) 100%);
  background-size: 200% 200%;
  -webkit-background-clip: text;
  background-clip: text;
  color: transparent;
  animation: gradientShift 8s linear infinite;
}

.about-member-role {
  margin-top: 0;
  margin-bottom: 0.95rem;
  color: #91c4ff;
  font-size: 0.76rem;
  font-weight: 700;
  letter-spacing: 0.16em;
  text-transform: uppercase;
}

.about-final-cta {
  width: min(100%, 1220px);
  margin: 0 auto;
  text-align: center;
  position: relative;
  overflow: hidden;
  padding: 4.45rem 1rem 4.7rem;
  border-top: 1px solid rgba(255, 255, 255, 0.06);
  border-bottom: 1px solid rgba(255, 255, 255, 0.06);
  background:
    linear-gradient(180deg, rgba(7, 11, 18, 0.55), rgba(7, 11, 18, 0.2));
}

.about-final-cta::before {
  content: "";
  position: absolute;
  inset: 0;
  background-image:
    linear-gradient(rgba(255, 255, 255, 0.04) 1px, transparent 1px),
    linear-gradient(90deg, rgba(255, 255, 255, 0.04) 1px, transparent 1px);
  background-size: 42px 42px;
  opacity: 0.12;
  pointer-events: none;
}

.about-final-cta > * {
  position: relative;
  z-index: 1;
}

.about-final-cta h2 {
  margin: 0 0 0.75rem;
  font-family: "Sora", sans-serif;
  font-size: clamp(1.9rem, 3vw, 3rem);
  line-height: 1.04;
  letter-spacing: -0.05em;
}

.about-final-cta p {
  margin: 0 auto;
  max-width: 640px;
  color: rgba(232, 241, 255, 0.78);
  font-size: 1.05rem;
  line-height: 1.7;
}

.about-sales-cta {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  gap: 0.75rem;
  min-width: 184px;
  min-height: 50px;
  margin-top: 1.8rem;
  padding: 0 1.55rem;
  border-radius: 18px;
  border: 0;
  background: linear-gradient(90deg, #7a4fff 0%, #4c8af4 52%, #7a4fff 100%);
  box-shadow: 0 0 28px rgba(122, 67, 255, 0.28);
  color: #eef5ff;
  font-family: "Sora", sans-serif;
  font-size: 0.96rem;
  font-weight: 700;
  text-decoration: none;
}

.st-key-about_contact_sales_button div.stButton > button,
.st-key-about_contact_sales_button button {
  border: 1px solid rgba(122, 67, 255, 0.3) !important;
  background: linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--purple-strong) 100%) !important;
  box-shadow: 0 14px 38px rgba(122, 67, 255, 0.24) !important;
}

.st-key-about_contact_sales_button div.stButton > button:hover,
.st-key-about_contact_sales_button button:hover {
  box-shadow: 0 16px 42px rgba(122, 67, 255, 0.3) !important;
}

.st-key-about_contact_sales_button div.stButton > button[kind="primary"],
.st-key-about_contact_sales_button button[kind="primary"] {
  border-color: rgba(122, 67, 255, 0.34) !important;
  background: linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--purple-strong) 100%) !important;
  box-shadow: 0 14px 38px rgba(122, 67, 255, 0.24) !important;
}

.about-sales-cta span {
  font-size: 1.05rem;
}

.about-inline-ai {
  font-weight: 800;
}

.detail-grid {
  display: grid;
  gap: 16px;
  grid-template-columns: repeat(3, minmax(0, 1fr));
}

.product-card-cta {
  display: flex;
  justify-content: center;
  margin-top: 1rem;
}

.product-card-cta-link {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  gap: 0.65rem;
  min-width: 210px;
  min-height: 48px;
  padding: 0 1.45rem;
  border-radius: 18px;
  color: #eef5ff;
  font-family: "Sora", sans-serif;
  font-size: 0.94rem;
  font-weight: 700;
  text-decoration: none;
  background: linear-gradient(90deg, #5f6bff 0%, #4f8ef5 52%, #28c2ea 100%);
  box-shadow: 0 0 26px rgba(76, 138, 244, 0.28);
  transition: transform 160ms ease, filter 160ms ease, box-shadow 160ms ease;
}

.product-card-cta-link:hover {
  transform: translateY(-1px);
  filter: brightness(1.03);
  box-shadow: 0 0 30px rgba(76, 138, 244, 0.34);
}

.guide-how-section {
  display: grid;
  gap: 1rem;
  margin: 1.25rem 0 1.5rem;
}

.guide-how-head {
  display: inline-flex;
  align-items: center;
  gap: 0.6rem;
  font-family: "Sora", sans-serif;
  font-size: 1.55rem;
  font-weight: 700;
  letter-spacing: -0.04em;
}

.guide-how-head svg {
  width: 18px;
  height: 18px;
  stroke: #46a7ff;
  stroke-width: 1.9;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.guide-how-shell {
  padding: 1.25rem 1.3rem;
}

.guide-how-grid {
  display: grid;
  grid-template-columns: repeat(5, minmax(0, 1fr));
  gap: 12px;
}

.guide-how-card {
  min-height: 114px;
  display: grid;
  justify-items: center;
  align-content: center;
  gap: 0.6rem;
  padding: 1.2rem 0.95rem 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: linear-gradient(180deg, rgba(8, 13, 20, 0.88), rgba(8, 13, 20, 0.96));
  text-align: center;
}

.guide-how-icon {
  width: 22px;
  height: 22px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  color: #46a7ff;
}

.guide-how-icon svg {
  width: 100%;
  height: 100%;
  stroke: currentColor;
  stroke-width: 1.85;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.guide-how-card strong {
  display: block;
  font-family: "Sora", sans-serif;
  font-size: 1rem;
}

.guide-how-card span {
  color: rgba(220, 233, 255, 0.62);
  font-size: 0.88rem;
  line-height: 1.55;
}

.guide-modules-section {
  display: block;
  margin: 0.35rem 0 1.6rem;
}

.guide-modules-layout {
  display: grid;
  grid-template-columns: minmax(0, 7fr) minmax(300px, 3fr);
  gap: 16px;
  align-items: start;
}

.guide-modules-main {
  display: grid;
  gap: 1rem;
}

.guide-modules-head {
  display: inline-flex;
  align-items: center;
  gap: 0.6rem;
  font-family: "Sora", sans-serif;
  font-size: 1.5rem;
  font-weight: 700;
  letter-spacing: -0.04em;
}

.guide-modules-head svg {
  width: 18px;
  height: 18px;
  stroke: #46a7ff;
  stroke-width: 1.9;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.guide-modules-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 16px;
}

.guide-module-card {
  display: grid;
  gap: 0.95rem;
  min-height: 238px;
  padding: 1.2rem 1.2rem 1.15rem;
  border-radius: 22px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: linear-gradient(180deg, rgba(8, 13, 20, 0.9), rgba(8, 13, 20, 0.98));
  text-decoration: none;
  color: inherit;
  transition: transform 180ms ease, box-shadow 180ms ease, border-color 180ms ease;
}

.guide-module-link {
  color: inherit;
  text-decoration: none;
}

.guide-module-link:focus-visible {
  outline: 2px solid rgba(70, 167, 255, 0.7);
  outline-offset: 3px;
  border-radius: 20px;
}

.guide-module-card:hover {
  transform: translateY(-3px);
  box-shadow: 0 24px 48px rgba(0, 0, 0, 0.24);
}

.guide-module-head {
  display: flex;
  align-items: center;
  gap: 0.85rem;
}

.guide-module-icon {
  width: 30px;
  height: 30px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 10px;
  color: #eef5ff;
}

.guide-module-icon svg {
  width: 17px;
  height: 17px;
  stroke: currentColor;
  stroke-width: 1.9;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.guide-module-icon.decision { background: linear-gradient(135deg, #2e77ff, #55a4ff); }
.guide-module-icon.task { background: linear-gradient(135deg, #8a43ff, #d66cff); }
.guide-module-icon.risk { background: linear-gradient(135deg, #d79312, #f0b738); }
.guide-module-icon.monte { background: linear-gradient(135deg, #16a56c, #43d07f); }
.guide-module-icon.summary { background: linear-gradient(135deg, #4c8af4, #74b5ff); }
.guide-module-icon.upload { background: linear-gradient(135deg, #7a8597, #adb6c4); }

.guide-module-title {
  font-family: "Sora", sans-serif;
  font-size: 1.2rem;
  font-weight: 700;
  letter-spacing: -0.04em;
}

.guide-module-arrow {
  margin-left: auto;
  color: rgba(220, 233, 255, 0.28);
  font-size: 1rem;
}

.guide-module-copy {
  margin: 0;
  color: rgba(220, 233, 255, 0.68);
  line-height: 1.7;
}

.guide-module-preview {
  margin-top: auto;
  padding: 0.95rem 1rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.04);
  background: rgba(255, 255, 255, 0.025);
}

.guide-module-note-kicker {
  color: #46a7ff;
  font-size: 0.82rem;
  font-weight: 700;
}

.guide-module-note-copy {
  margin: 0.45rem 0 0;
  color: rgba(220, 233, 255, 0.68);
  line-height: 1.72;
  font-size: 0.9rem;
}

.guide-sparkline {
  height: 84px;
  border-radius: 14px;
  background:
    linear-gradient(180deg, rgba(255, 255, 255, 0.02), rgba(255, 255, 255, 0.01)),
    rgba(7, 11, 17, 0.7);
  overflow: hidden;
}

.guide-sparkline svg,
.guide-monte-preview svg {
  width: 100%;
  height: 100%;
  display: block;
}

.guide-risk-list {
  display: grid;
  gap: 0.55rem;
}

.guide-risk-row {
  display: grid;
  grid-template-columns: auto 1fr auto;
  gap: 0.55rem;
  align-items: center;
  color: rgba(232, 241, 255, 0.82);
  font-size: 0.9rem;
}

.guide-risk-dot {
  width: 7px;
  height: 7px;
  border-radius: 999px;
}

.guide-risk-dot.high { background: #ff564d; }
.guide-risk-dot.medium { background: #ffb020; }
.guide-risk-dot.low { background: #99a2b2; }

.guide-risk-value.high { color: #ff564d; }
.guide-risk-value.medium { color: #ffb020; }
.guide-risk-value.low { color: rgba(232, 241, 255, 0.66); }

.guide-monte-preview {
  height: 84px;
  border-radius: 14px;
  background: rgba(7, 11, 17, 0.7);
  overflow: hidden;
}

.guide-activity-panel {
  display: grid;
  gap: 1.1rem;
  padding: 1.2rem 1.2rem 1.15rem;
  border-radius: 22px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: linear-gradient(180deg, rgba(8, 13, 20, 0.9), rgba(8, 13, 20, 0.98));
  box-shadow: 0 20px 46px rgba(0, 0, 0, 0.2);
}

.guide-activity-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: 1.35rem;
  font-weight: 700;
  letter-spacing: -0.04em;
}

.guide-activity-block {
  display: grid;
  gap: 0.75rem;
}

.guide-activity-subhead {
  margin: 0;
  color: rgba(232, 241, 255, 0.72);
  font-size: 0.82rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.guide-activity-list {
  display: grid;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.05);
  background: rgba(255, 255, 255, 0.02);
  overflow: hidden;
}

.guide-activity-item {
  display: grid;
  grid-template-columns: auto minmax(0, 1fr) auto;
  gap: 0.8rem;
  align-items: center;
  padding: 0.9rem 0.95rem;
}

.guide-activity-item + .guide-activity-item {
  border-top: 1px solid rgba(255, 255, 255, 0.05);
}

.guide-activity-icon {
  width: 28px;
  height: 28px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 9px;
  background: rgba(70, 167, 255, 0.12);
  color: #7fd3ff;
  box-shadow: inset 0 0 0 1px rgba(70, 167, 255, 0.12);
}

.guide-activity-icon svg {
  width: 15px;
  height: 15px;
  stroke: currentColor;
  stroke-width: 1.8;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.guide-activity-copy {
  color: rgba(232, 241, 255, 0.84);
  font-size: 0.92rem;
  line-height: 1.55;
}

.guide-activity-time {
  color: rgba(220, 233, 255, 0.5);
  font-size: 0.8rem;
  white-space: nowrap;
}

.guide-activity-actions {
  display: grid;
  gap: 0.75rem;
}

.guide-activity-action {
  display: flex;
  align-items: center;
  gap: 0.8rem;
  width: 100%;
  padding: 0.9rem 1rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.025);
  color: inherit;
  text-decoration: none;
  transition: transform 180ms ease, background 180ms ease, border-color 180ms ease, box-shadow 180ms ease;
}

.guide-activity-action:hover {
  transform: translateY(-2px);
  border-color: rgba(70, 167, 255, 0.22);
  background: rgba(70, 167, 255, 0.05);
  box-shadow: 0 14px 30px rgba(0, 0, 0, 0.18);
}

.guide-activity-action-label {
  flex: 1;
  font-family: "Sora", sans-serif;
  font-size: 0.95rem;
  font-weight: 600;
}

.guide-activity-action-arrow {
  color: rgba(232, 241, 255, 0.42);
  font-size: 1rem;
}

.ai-forecast-shell {
  display: grid;
  gap: 1.8rem;
}

.ai-forecast-kicker-card {
  padding: 1.05rem 1.3rem;
}

.ai-forecast-stage {
  display: grid;
  gap: 1.9rem;
  text-align: center;
}

.ai-forecast-copy {
  width: min(100%, 760px);
  margin-inline: auto;
}

.ai-forecast-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(2.2rem, 4.5vw, 3.45rem);
  line-height: 1.04;
  letter-spacing: -0.055em;
}

.ai-forecast-subtitle {
  margin: 1rem auto 0;
  max-width: 720px;
  color: rgba(232, 241, 255, 0.76);
  font-size: 1.04rem;
  line-height: 1.72;
}

.ai-forecast-grid {
  display: grid;
  grid-template-columns: repeat(2, minmax(0, 1fr));
  gap: 22px;
}

.ai-forecast-card {
  min-height: 182px;
  text-align: left;
  padding: 1.5rem 1.55rem;
}

.ai-forecast-icon {
  width: 28px;
  height: 28px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  margin-bottom: 1rem;
  color: #eef5ff;
}

.ai-forecast-icon svg {
  width: 100%;
  height: 100%;
  stroke: currentColor;
  stroke-width: 1.8;
  fill: none;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.ai-forecast-process-card {
  padding: 1.55rem 1.7rem 1.7rem;
  text-align: left;
}

.ai-forecast-process-card h3 {
  margin: 0 0 1.5rem;
  font-family: "Sora", sans-serif;
  font-size: 1.08rem;
}

.ai-forecast-steps {
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 1.25rem;
}

.ai-forecast-step {
  display: grid;
  justify-items: center;
  gap: 1rem;
  text-align: center;
}

.ai-forecast-step-dot {
  width: 38px;
  height: 38px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 999px;
  color: #46a7ff;
  font-family: "Sora", sans-serif;
  font-size: 0.95rem;
  font-weight: 700;
  background: rgba(43, 110, 214, 0.18);
  box-shadow: 0 0 0 10px rgba(43, 110, 214, 0.08);
}

.ai-forecast-step strong {
  font-size: 0.97rem;
  font-weight: 500;
}

.ai-forecast-cta-wrap {
  display: flex;
  justify-content: center;
  margin-top: 0.15rem;
}

.ai-forecast-cta {
  margin-top: 0;
  min-width: 248px;
  min-height: 52px;
  border-radius: 18px;
}

.dashboard-banner {
  padding: 1.6rem 1.8rem;
}

.task-architect-shell {
  display: grid;
  gap: 1.85rem;
}

.task-architect-hero {
  border: 1px solid transparent;
  border-radius: 30px;
  padding: 1.45rem 1.65rem 1.7rem;
  background:
    linear-gradient(180deg, rgba(16, 24, 40, 0.94), rgba(7, 18, 32, 0.94)) padding-box,
    linear-gradient(90deg, rgba(123, 78, 255, 0.44) 0%, rgba(79, 123, 255, 0.26) 58%, rgba(17, 225, 245, 0.42) 100%) border-box;
  box-shadow: var(--shadow);
}

.task-architect-hero-kicker {
  display: inline-flex;
  align-items: center;
  min-height: 34px;
  padding: 0 0.9rem;
  border-radius: 999px;
  border: 1px solid rgba(74, 193, 255, 0.34);
  background: rgba(8, 18, 34, 0.52);
  color: #7fe5ff;
  font-size: 0.74rem;
  font-weight: 700;
  letter-spacing: 0.13em;
  text-transform: uppercase;
  white-space: nowrap;
}

.task-architect-hero-title {
  margin: 1.05rem 0 0;
  font-family: "Sora", sans-serif;
  font-size: clamp(1.6rem, 2.45vw, 2.6rem);
  line-height: 1.02;
  letter-spacing: -0.06em;
  max-width: none;
  white-space: nowrap;
}

.task-architect-panel-head svg,
.task-architect-primary-btn svg,
.task-architect-empty-icon svg {
  width: 16px;
  height: 16px;
  stroke: currentColor;
  fill: none;
  stroke-width: 1.8;
  stroke-linecap: round;
  stroke-linejoin: round;
}

.task-architect-panel {
  padding: 1.35rem 1.5rem 1.5rem;
}

.task-architect-panel-head {
  display: inline-flex;
  align-items: center;
  gap: 0.6rem;
  color: rgba(232, 241, 255, 0.62);
  font-size: 0.83rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.task-architect-parameter-grid {
  display: grid;
  grid-template-columns: repeat(3, minmax(0, 1fr));
  gap: 14px;
  margin-top: 1.35rem;
}

.task-architect-field {
  display: grid;
  gap: 0.6rem;
}

.task-architect-field-label {
  color: rgba(232, 241, 255, 0.8);
  font-size: 0.98rem;
}

.task-architect-input {
  min-height: 56px;
  display: flex;
  align-items: center;
  padding: 0 1rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.035);
  color: #f2f8ff;
  font-size: 1rem;
  line-height: 1.4;
}

.task-architect-input.placeholder {
  color: rgba(232, 241, 255, 0.32);
}

.task-architect-panel-actions {
  margin-top: 1.35rem;
}

.task-architect-primary-btn {
  display: inline-flex;
  align-items: center;
  gap: 0.6rem;
  min-height: 52px;
  padding: 0 1.25rem;
  border-radius: 16px;
  border: none;
  background: linear-gradient(90deg, #5c69ff 0%, #29c4f2 100%);
  box-shadow: 0 14px 34px rgba(67, 140, 255, 0.3);
  color: #ffffff;
  font-family: "Sora", sans-serif;
  font-size: 0.98rem;
  font-weight: 700;
}

.task-architect-empty {
  min-height: 260px;
  display: grid;
  place-items: center;
  padding: 2.25rem 1.5rem;
}

.task-architect-empty-inner {
  display: grid;
  justify-items: center;
  gap: 1rem;
  text-align: center;
}

.task-architect-empty-icon {
  width: 66px;
  height: 66px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 19px;
  background: linear-gradient(180deg, rgba(29, 48, 98, 0.96), rgba(18, 30, 67, 0.98));
  color: #eef5ff;
  box-shadow:
    0 0 0 1px rgba(118, 231, 255, 0.12) inset,
    0 0 24px rgba(92, 169, 255, 0.18);
}

.task-architect-empty-icon svg {
  width: 26px;
  height: 26px;
}

.task-architect-empty-title {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: 1.8rem;
  letter-spacing: -0.05em;
}

.task-architect-empty-copy {
  max-width: 540px;
  margin: 0;
  color: var(--muted);
  font-size: 1.02rem;
  line-height: 1.7;
}

.kpi-grid {
  display: grid;
  grid-template-columns: repeat(4, minmax(0, 1fr));
  gap: 16px;
  margin: 1.5rem 0;
}

.metric-card {
  border: 1px solid rgba(255, 255, 255, 0.06);
  border-radius: 22px;
  background: rgba(255, 255, 255, 0.03);
  padding: 1rem 1.1rem;
}

.metric-card span {
  color: var(--muted);
  font-size: 0.88rem;
}

.metric-card strong {
  margin-top: 0.38rem;
  font-size: 1.75rem;
  letter-spacing: -0.05em;
}

.metric-card small {
  display: block;
  margin-top: 0.5rem;
  color: var(--muted);
}

.status-chip {
  width: fit-content;
  padding: 0.4rem 0.8rem;
  border-radius: 999px;
  font-weight: 700;
}

.sidebar-brand-card {
  border-bottom: 1px solid rgba(255, 255, 255, 0.08);
  padding: 0.55rem 0 1rem;
  margin-bottom: 0.35rem;
}

.sidebar-mini-label {
  color: var(--muted);
  font-size: 0.8rem;
  line-height: 1.5;
}

.shortcut-stack {
  display: grid;
  gap: 0.55rem;
  justify-items: center;
}

.shortcut-caption {
  color: var(--muted);
  font-size: 0.8rem;
  line-height: 1.5;
  text-align: center;
}

.command-bar-grid {
  display: grid;
  grid-template-columns: minmax(0, 1fr) auto;
  gap: 1rem;
  align-items: center;
}

.command-bar-copy {
  min-width: 0;
}

.command-bar-status {
  display: flex;
  justify-content: flex-end;
}

.sidebar-shortcut-wrap {
  display: grid;
  justify-items: center;
  gap: 0.55rem;
  margin-top: 0.15rem;
}

.sidebar-shortcut-pill {
  min-width: 112px;
  min-height: 44px;
  justify-content: center;
  color: var(--text);
  font-family: "Sora", sans-serif;
  font-weight: 700;
  letter-spacing: -0.01em;
  border: 1.5px solid transparent;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box;
  box-shadow: none;
}

.sidebar-shortcut-trigger {
  appearance: none;
  cursor: pointer;
  text-decoration: none;
  transition: transform 160ms ease, border-color 160ms ease, background 160ms ease, box-shadow 160ms ease;
}

.sidebar-shortcut-trigger:hover {
  transform: translateY(-1px);
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.22),
    0 14px 34px rgba(0, 0, 0, 0.18);
}

.sidebar-shortcut-trigger:focus-visible {
  outline: none;
  box-shadow:
    0 0 0 3px rgba(47, 162, 255, 0.12),
    0 0 0 1px rgba(117, 231, 255, 0.22);
}

.sidebar-shortcut-caption {
  text-align: center;
}

.sidebar-shortcut-key {
  color: var(--text);
  opacity: 1;
}

.floating-shortcut-wrap {
  position: fixed;
  left: 1.1rem;
  bottom: 1.2rem;
  z-index: 26;
  display: grid;
  justify-items: center;
  gap: 0.55rem;
  padding: 0.7rem 0.8rem 0.35rem;
  border-top: 1px solid rgba(255, 255, 255, 0.08);
  transition: opacity 180ms ease, transform 180ms ease;
}

.floating-shortcut-wrap .sidebar-shortcut-pill {
  min-width: 124px;
}

.floating-shortcut-wrap .sidebar-shortcut-caption {
  max-width: 150px;
}

.hero-shortcut-wrap {
  margin: 0.9rem 0 0.2rem;
}

.command-palette-overlay {
  position: fixed;
  inset: 0;
  z-index: 70;
  display: flex;
  align-items: flex-start;
  justify-content: center;
  padding: 5rem 1rem 1rem;
  opacity: 0;
  pointer-events: none;
  transition: opacity 180ms ease;
}

.command-palette-overlay.is-open {
  opacity: 1;
  pointer-events: auto;
}

.command-palette-backdrop {
  position: absolute;
  inset: 0;
  background: rgba(2, 7, 14, 0.72);
  backdrop-filter: blur(10px);
}

.command-palette-card {
  position: relative;
  width: min(100%, 760px);
  max-height: min(72vh, 680px);
  overflow: hidden;
  border-radius: 28px;
  border: 1px solid rgba(109, 165, 255, 0.24);
  background:
    linear-gradient(180deg, rgba(7, 13, 24, 0.98), rgba(6, 11, 21, 0.97));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 28px 96px rgba(0, 0, 0, 0.42);
}

.command-palette-head {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 1rem;
  padding: 1.05rem 1.15rem 0.8rem;
  border-bottom: 1px solid rgba(255, 255, 255, 0.06);
}

.command-palette-title {
  display: grid;
  gap: 0.2rem;
}

.command-palette-title strong {
  font-family: "Sora", sans-serif;
  font-size: 1rem;
}

.command-palette-title span {
  color: var(--muted);
  font-size: 0.9rem;
}

.command-palette-close {
  width: 2.3rem;
  height: 2.3rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  cursor: pointer;
}

.command-palette-search {
  padding: 1rem 1.15rem 0.75rem;
}

.command-palette-input {
  width: 100%;
  min-height: 54px;
  padding: 0 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.035);
  color: var(--text);
  outline: none;
  box-sizing: border-box;
}

.command-palette-input:focus {
  border-color: rgba(109, 165, 255, 0.28);
  box-shadow: 0 0 0 3px rgba(47, 162, 255, 0.08);
}

.command-palette-list {
  display: grid;
  gap: 0.7rem;
  max-height: min(52vh, 520px);
  padding: 0 1.15rem 1.15rem;
  overflow: auto;
}

.command-palette-item {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 1rem;
  padding: 0.95rem 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  text-decoration: none;
  transition: transform 160ms ease, border-color 160ms ease, background 160ms ease;
}

.command-palette-item:hover {
  transform: translateY(-1px);
  border-color: rgba(117, 231, 255, 0.18);
  background: rgba(255, 255, 255, 0.05);
}

.command-palette-item.active {
  border-color: rgba(109, 165, 255, 0.24);
}

.command-palette-item-copy {
  min-width: 0;
  display: grid;
  gap: 0.2rem;
}

.command-palette-item-title {
  font-family: "Sora", sans-serif;
  font-size: 0.98rem;
}

.command-palette-item-desc {
  color: var(--muted);
  font-size: 0.88rem;
  line-height: 1.45;
}

.command-palette-item-group {
  flex: 0 0 auto;
  color: var(--cyan);
  font-size: 0.76rem;
  font-weight: 700;
  letter-spacing: 0.08em;
  text-transform: uppercase;
}

.command-palette-empty {
  padding: 1.4rem 1rem;
  border-radius: 18px;
  border: 1px dashed rgba(255, 255, 255, 0.08);
  color: var(--muted);
  text-align: center;
}

.ai-assistant-shell {
  position: fixed;
  right: 1.25rem;
  bottom: 1.25rem;
  z-index: 52;
  display: grid;
  justify-items: end;
  gap: 0.85rem;
  transition: opacity 180ms ease, transform 180ms ease;
}

.ai-assistant-launcher {
  min-width: 178px;
  min-height: 56px;
  padding: 0.7rem 0.95rem;
  display: inline-flex;
  align-items: center;
  gap: 0.8rem;
  border: 1.5px solid transparent;
  border-radius: 20px;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box;
  box-shadow:
    0 16px 38px rgba(0, 0, 0, 0.28),
    0 0 0 1px rgba(255, 255, 255, 0.02) inset;
  color: var(--text);
  cursor: pointer;
  transition: transform 160ms ease, border-color 160ms ease, box-shadow 160ms ease;
}

.ai-assistant-launcher:hover {
  transform: translateY(-1px);
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.22),
    0 18px 42px rgba(0, 0, 0, 0.34),
    0 0 0 1px rgba(255, 255, 255, 0.02) inset;
}

.ai-assistant-launcher:focus-visible {
  outline: none;
  box-shadow:
    0 0 0 3px rgba(47, 162, 255, 0.12),
    0 18px 42px rgba(0, 0, 0, 0.34);
}

.ai-assistant-launcher-icon {
  width: 2.45rem;
  height: 2.45rem;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 999px;
  background: linear-gradient(180deg, rgba(115, 87, 255, 0.95), rgba(47, 162, 255, 0.98));
  box-shadow: 0 0 26px rgba(78, 143, 255, 0.22);
  font-size: 1.02rem;
  font-weight: 800;
}

.ai-assistant-launcher-copy {
  display: grid;
  justify-items: start;
  gap: 0.14rem;
  text-align: left;
}

.ai-assistant-launcher-copy strong {
  font-family: "Sora", sans-serif;
  font-size: 0.96rem;
}

.ai-assistant-launcher-copy span {
  color: var(--muted);
  font-size: 0.82rem;
}

.ai-assistant-panel {
  width: min(380px, calc(100vw - 1.5rem));
  max-height: min(68vh, 620px);
  overflow: hidden;
  border-radius: 26px;
  border: 1px solid rgba(107, 169, 255, 0.24);
  background: linear-gradient(180deg, rgba(7, 13, 24, 0.985), rgba(5, 10, 19, 0.975));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 28px 88px rgba(0, 0, 0, 0.42);
  opacity: 0;
  pointer-events: none;
  transform: translateY(16px) scale(0.98);
  transform-origin: bottom right;
  transition: opacity 180ms ease, transform 180ms ease;
}

.ai-assistant-shell.is-open .ai-assistant-panel {
  opacity: 1;
  pointer-events: auto;
  transform: translateY(0) scale(1);
}

.ai-assistant-head {
  display: flex;
  align-items: flex-start;
  justify-content: space-between;
  gap: 0.85rem;
  padding: 1rem 1rem 0.8rem;
  border-bottom: 1px solid rgba(255, 255, 255, 0.06);
}

.ai-assistant-head-copy {
  display: grid;
  gap: 0.22rem;
}

.ai-assistant-kicker {
  color: var(--cyan);
  font-size: 0.76rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.ai-assistant-head-copy strong {
  font-family: "Sora", sans-serif;
  font-size: 1rem;
}

.ai-assistant-head-copy span {
  color: var(--muted);
  font-size: 0.88rem;
  line-height: 1.45;
}

.ai-assistant-close {
  width: 2.15rem;
  height: 2.15rem;
  flex: 0 0 auto;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  cursor: pointer;
}

.ai-assistant-suggestions {
  display: flex;
  flex-wrap: wrap;
  gap: 0.55rem;
  padding: 0.85rem 1rem 0;
}

.ai-assistant-chip {
  padding: 0.5rem 0.72rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  cursor: pointer;
  font-size: 0.82rem;
}

.ai-assistant-messages {
  min-height: 210px;
  max-height: min(42vh, 360px);
  display: grid;
  gap: 0.8rem;
  padding: 1rem;
  overflow: auto;
}

.ai-assistant-message {
  max-width: 88%;
  padding: 0.82rem 0.9rem;
  border-radius: 18px;
  line-height: 1.55;
  font-size: 0.92rem;
  white-space: pre-wrap;
}

.ai-assistant-message.assistant {
  justify-self: start;
  border: 1px solid rgba(117, 231, 255, 0.14);
  background: rgba(255, 255, 255, 0.04);
}

.ai-assistant-message.user {
  justify-self: end;
  background: linear-gradient(135deg, rgba(115, 87, 255, 0.96), rgba(47, 162, 255, 0.96));
  box-shadow: 0 14px 30px rgba(76, 138, 244, 0.22);
}

.ai-assistant-form {
  display: grid;
  grid-template-columns: minmax(0, 1fr) auto;
  gap: 0.75rem;
  padding: 0 1rem 1rem;
}

.ai-assistant-input {
  min-width: 0;
  min-height: 50px;
  padding: 0 0.95rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.035);
  color: var(--text);
}

.ai-assistant-input:focus {
  outline: none;
  border-color: rgba(109, 165, 255, 0.28);
  box-shadow: 0 0 0 3px rgba(47, 162, 255, 0.08);
}

.ai-assistant-send {
  min-width: 78px;
  min-height: 50px;
  padding: 0 1rem;
  border: none;
  border-radius: 16px;
  background: linear-gradient(90deg, #7357ff 0%, #2fa2ff 100%);
  box-shadow: 0 12px 28px rgba(76, 138, 244, 0.24);
  color: #ffffff;
  font-weight: 700;
  cursor: pointer;
}

@media (max-width: 720px) {
  .ai-assistant-shell {
    right: 0.8rem;
    bottom: 0.9rem;
  }

  .ai-assistant-panel {
    width: min(100vw - 1rem, 380px);
  }

  .ai-assistant-launcher {
    min-width: 164px;
  }
}

.sidebar-section-label {
  margin: 1rem 0 0.7rem;
  color: var(--cyan);
  font-size: 0.76rem;
  font-weight: 700;
  letter-spacing: 0.1em;
  text-transform: uppercase;
}

.sidebar-divider {
  width: 100%;
  height: 1px;
  margin: 1rem 0;
  background: rgba(255, 255, 255, 0.08);
}

.sidebar-row-link {
  align-items: center;
  min-height: 42px;
  padding: 0 0.55rem;
  border-radius: 12px;
  color: var(--muted);
  text-decoration: none;
  font-weight: 600;
  transition: background 150ms ease, color 150ms ease;
}

.sidebar-row-link:hover {
  background: rgba(255, 255, 255, 0.04);
  color: var(--text);
}

.chip-high {
  color: #ffd4d8;
  border-color: rgba(255, 127, 140, 0.28);
  background: rgba(255, 127, 140, 0.12);
}

.chip-medium {
  color: #ffe7bc;
  border-color: rgba(245, 197, 107, 0.28);
  background: rgba(245, 197, 107, 0.11);
}

.chip-low {
  color: #cbfff7;
  border-color: rgba(24, 209, 199, 0.22);
  background: rgba(24, 209, 199, 0.11);
}

.workspace-card p,
.timeline-card p,
.history-card p {
  color: var(--muted);
  line-height: 1.65;
}

.data-pair {
  display: flex;
  justify-content: space-between;
  gap: 14px;
  padding: 0.55rem 0;
  border-bottom: 1px solid rgba(255, 255, 255, 0.05);
}

.data-pair:last-child {
  border-bottom: none;
}

.data-pair strong a {
  color: inherit;
  text-decoration: none;
}

.data-pair strong a:hover,
.data-pair strong a:focus {
  text-decoration: underline;
}

.faq-summary-card .data-pair {
  align-items: center;
  gap: 10px;
  white-space: nowrap;
}

.faq-summary-card .data-pair span {
  white-space: nowrap;
}

.faq-summary-card .data-pair strong {
  white-space: nowrap;
  font-size: 0.94rem;
  letter-spacing: -0.01em;
  text-align: right;
}

@media (min-width: 1101px) {
  .faq-summary-card {
    width: calc(100% + 2.3cm);
    margin-right: -2.3cm;
    justify-self: start;
  }
}

.contact-social-inline {
  margin-top: 1rem;
}

.contact-social-inline strong {
  display: block;
  font-family: "Sora", sans-serif;
  margin-bottom: 0.35rem;
}

.contact-social-inline p {
  margin: 0.5rem 0 0;
  color: var(--muted);
}

.contact-social-inline .footer-social-row {
  margin-top: 0;
}

.contact-top-stack {
  display: grid;
  gap: 22px;
}

.contact-right-balance {
  min-height: clamp(30rem, 46vh, 40rem);
  display: flex;
  flex-direction: column;
  justify-content: flex-end;
  gap: 1rem;
}

.timeline-grid,
.workspace-grid {
  display: grid;
  grid-template-columns: 1.15fr 0.85fr;
  gap: 16px;
}

.task-list {
  display: grid;
  gap: 12px;
}

.task-item {
  display: flex;
  justify-content: space-between;
  gap: 12px;
  padding: 0.85rem 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.05);
  background: rgba(255, 255, 255, 0.025);
}

.task-item span {
  color: var(--muted);
}

.scenario-card-grid {
  margin-top: 0.8rem;
}

.scenario-mini-card h4 {
  margin: 0.2rem 0 0.45rem;
  font-family: "Sora", sans-serif;
}

.scenario-mini-card p {
  margin: 0;
}

.history-list {
  display: grid;
  gap: 12px;
}

.history-item {
  padding: 0.95rem 1rem;
  border-radius: 18px;
  border: 1px solid rgba(255, 255, 255, 0.06);
  background: rgba(255, 255, 255, 0.03);
}

.history-item strong {
  display: block;
  font-family: "Sora", sans-serif;
}

.history-item span {
  color: var(--muted);
  font-size: 0.94rem;
}

.footer-shell {
  width: min(100%, var(--nav-width));
  margin-inline: auto;
  margin-top: 4rem;
}

.footer-panel {
  position: relative;
  overflow: visible;
  border-top: 1px solid rgba(255, 255, 255, 0.06);
  background:
    radial-gradient(circle at 12% 34%, rgba(24, 209, 199, 0.08), transparent 18%),
    radial-gradient(circle at 36% 78%, rgba(76, 138, 244, 0.12), transparent 16%),
    radial-gradient(circle at 58% 46%, rgba(122, 67, 255, 0.09), transparent 18%),
    radial-gradient(circle at 86% 22%, rgba(24, 209, 199, 0.08), transparent 14%),
    linear-gradient(180deg, rgba(3, 7, 14, 0.985), rgba(2, 6, 13, 0.995));
}

.footer-panel::before,
.footer-panel::after {
  content: "";
  position: absolute;
  inset: 0;
  pointer-events: none;
}

.footer-panel::before {
  background:
    linear-gradient(112deg, transparent 0 46%, rgba(100, 162, 255, 0.1) 46.3%, transparent 46.8%) 12% 18% / 340px 220px no-repeat,
    linear-gradient(80deg, transparent 0 48%, rgba(24, 209, 199, 0.1) 48.4%, transparent 48.9%) 72% 14% / 420px 260px no-repeat,
    linear-gradient(100deg, transparent 0 49%, rgba(122, 67, 255, 0.1) 49.4%, transparent 49.9%) 46% 60% / 360px 240px no-repeat;
  opacity: 0.28;
}

.footer-panel::after {
  background:
    radial-gradient(circle, rgba(76, 138, 244, 0.65) 0 1.5px, transparent 2px) 9% 34% / 240px 180px,
    radial-gradient(circle, rgba(24, 209, 199, 0.55) 0 1.5px, transparent 2px) 58% 62% / 300px 220px,
    radial-gradient(circle, rgba(76, 138, 244, 0.5) 0 1.5px, transparent 2px) 88% 20% / 260px 180px;
  opacity: 0.52;
}

.footer-content {
  position: relative;
  z-index: 1;
  width: min(calc(100% - 48px), var(--nav-width));
  margin-inline: auto;
  padding: 4rem 0 1.2rem;
}

.footer-top-grid,
.footer-mid-row,
.footer-bottom-row,
.footer-link-grid,
.footer-social-row,
.footer-subscribe-row,
.footer-compliance-row {
  display: flex;
}

.footer-top-grid,
.footer-mid-row {
  display: grid;
  gap: 2.75rem;
}

.footer-top-grid {
  grid-template-columns: minmax(280px, 1.05fr) minmax(0, 1.55fr);
}

.footer-brand-column {
  max-width: 31rem;
}

.footer-brand-row {
  display: inline-flex;
  align-items: center;
  gap: 0.95rem;
  text-decoration: none !important;
  border-bottom: none !important;
  box-shadow: none !important;
}

.footer-brand-row:hover,
.footer-brand-row:focus,
.footer-brand-row:visited {
  text-decoration: none !important;
  border-bottom: none !important;
  box-shadow: none !important;
}

.footer-brand-box {
  width: 64px;
  height: 64px;
  padding: 0.72rem;
  border-radius: 18px;
  background: linear-gradient(180deg, rgba(6, 10, 20, 0.98), rgba(4, 8, 16, 0.96));
  border: 1px solid rgba(255, 255, 255, 0.08);
  box-shadow: 0 0 0 1px rgba(122, 67, 255, 0.18), 0 18px 42px rgba(76, 138, 244, 0.18);
  display: grid;
  place-items: center;
}

.footer-brand-icon {
  display: block;
  width: 100%;
  height: auto;
}

.footer-brand-wordmark {
  font-family: "Sora", sans-serif;
  font-size: 2rem;
  font-weight: 700;
  letter-spacing: -0.045em;
  color: #eef5ff;
  text-decoration: none !important;
  border-bottom: none !important;
}

.footer-brand-accent {
  color: var(--brand-blue-accent);
  text-decoration: none !important;
  border-bottom: none !important;
}

.footer-brand-copy {
  margin: 1.45rem 0 0;
  color: var(--muted);
  font-size: 1.08rem;
  line-height: 1.72;
}

.footer-contact-stack {
  display: grid;
  gap: 0.9rem;
  margin-top: 1.5rem;
}

.footer-contact-row {
  display: inline-flex;
  align-items: center;
  gap: 0.7rem;
  color: var(--muted);
  text-decoration: none;
  font-size: 1.02rem;
}

.footer-contact-icon {
  width: 22px;
  height: 22px;
  border-radius: 999px;
  display: grid;
  place-items: center;
  color: var(--blue-strong);
  background: rgba(76, 138, 244, 0.08);
  border: 1px solid rgba(76, 138, 244, 0.18);
  font-size: 0.8rem;
  font-weight: 700;
}

.footer-social-row {
  gap: 0.85rem;
  align-items: center;
  margin-top: 1.9rem;
}

.footer-social-pill {
  width: 38px;
  height: 38px;
  border-radius: 12px;
  display: grid;
  place-items: center;
  background: rgba(255, 255, 255, 0.03);
  border: 1px solid rgba(255, 255, 255, 0.08);
  color: #d8e6fb;
  font-size: 0.76rem;
  font-weight: 700;
  letter-spacing: 0.04em;
  text-decoration: none;
  transition: transform 160ms ease, border-color 160ms ease, background 160ms ease;
}

.footer-social-pill:hover {
  transform: translateY(-1px);
  border-color: rgba(109, 165, 255, 0.2);
  background: rgba(255, 255, 255, 0.05);
}

.footer-social-pill:focus-visible {
  outline: none;
  box-shadow: 0 0 0 3px rgba(47, 162, 255, 0.12);
}

.footer-social-icon {
  width: 18px;
  height: 18px;
  display: block;
}

.footer-social-icon svg {
  width: 100%;
  height: 100%;
  display: block;
}

.footer-link-grid {
  justify-content: space-between;
  gap: 1.8rem;
  flex-wrap: nowrap;
}

.footer-link-column {
  min-width: 132px;
  display: grid;
  gap: 0.92rem;
  align-content: start;
}

.footer-link-column h4 {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: 1rem;
}

.footer-link-column a,
.footer-link-item,
.footer-static-link {
  color: var(--muted);
  text-decoration: none;
  font-size: 1.01rem;
  line-height: 1.45;
}

.footer-link-column a:hover,
.footer-link-item:hover {
  color: var(--text);
}

.footer-legal-link {
  display: inline-flex;
  align-items: center;
  width: fit-content;
  position: relative;
  z-index: 3;
  cursor: pointer;
  color: var(--blue-strong);
  text-decoration: underline;
  text-decoration-color: rgba(92, 169, 255, 0.78);
  text-decoration-thickness: 1.45px;
  text-underline-offset: 0.18em;
}

.footer-legal-link:hover,
.footer-legal-link:focus-visible {
  color: #eef5ff;
  text-decoration-color: rgba(92, 169, 255, 0.98);
}

.footer-link-note {
  color: rgba(158, 176, 203, 0.65);
  font-size: 0.82rem;
  font-style: italic;
}

.footer-mid-row {
  grid-template-columns: minmax(320px, 1fr) minmax(0, 1.2fr);
  align-items: center;
  margin-top: 3.25rem;
  padding-top: 2rem;
  border-top: 1px solid rgba(255, 255, 255, 0.07);
}

.footer-newsletter h4 {
  margin: 0;
  font-family: "Sora", sans-serif;
  font-size: 1.5rem;
}

.footer-newsletter p {
  margin: 0.7rem 0 0;
  color: var(--muted);
  font-size: 1rem;
}

.footer-subscribe-row {
  gap: 0.9rem;
  align-items: center;
  flex-wrap: wrap;
  margin-top: 1.35rem;
}

.footer-input-shell {
  min-height: 52px;
  min-width: 240px;
  flex: 1 1 260px;
  display: flex;
  align-items: center;
  padding: 0 1rem;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.03);
  color: #eef5ff;
  font: inherit;
  appearance: none;
}

.footer-input-shell::placeholder {
  color: rgba(158, 176, 203, 0.78);
}

.footer-input-shell:focus {
  outline: none;
  border-color: rgba(109, 165, 255, 0.22);
  box-shadow: 0 0 0 3px rgba(47, 162, 255, 0.08);
}

.footer-subscribe-btn {
  min-width: 168px;
  cursor: pointer;
  appearance: none;
  color: var(--blue-strong);
  font-family: "Sora", sans-serif;
  letter-spacing: -0.01em;
  text-decoration: underline;
  text-decoration-thickness: 1.5px;
  text-underline-offset: 0.12em;
}

.footer-compliance-row {
  justify-content: flex-end;
  gap: 1rem;
  flex-wrap: wrap;
}

.footer-compliance-item {
  min-width: 62px;
  display: grid;
  justify-items: center;
  gap: 0.42rem;
  text-align: center;
}

.footer-compliance-mark {
  width: 44px;
  height: 44px;
  display: grid;
  place-items: center;
  color: #eef5ff;
}

.footer-compliance-mark svg {
  width: 100%;
  height: 100%;
  display: block;
}

.footer-compliance-mark .compliance-stroke {
  fill: rgba(255, 255, 255, 0.02);
  stroke: rgba(255, 255, 255, 0.9);
  stroke-width: 1.5;
}

.footer-compliance-mark .compliance-glyph {
  fill: #eef5ff;
  font-family: "Sora", sans-serif;
  font-size: 8px;
  font-weight: 700;
  letter-spacing: 0.08em;
}

.footer-compliance-mark .compliance-detail {
  fill: rgba(238, 245, 255, 0.9);
  stroke: rgba(255, 255, 255, 0.9);
  stroke-width: 1.25;
}

.footer-compliance-code {
  color: #eef5ff;
  font-size: 0.72rem;
  font-weight: 700;
  letter-spacing: 0.01em;
}

.footer-compliance-region {
  color: rgba(158, 176, 203, 0.78);
  font-size: 0.72rem;
  line-height: 1.35;
}

.footer-bottom-row {
  justify-content: space-between;
  align-items: center;
  gap: 1rem;
  flex-wrap: wrap;
  margin-top: 2rem;
  padding-top: 1.5rem;
  border-top: 1px solid rgba(255, 255, 255, 0.07);
  color: var(--muted);
}

.footer-language-shell {
  display: inline-flex;
  align-items: center;
  gap: 0.75rem;
  position: relative;
  z-index: 4;
  margin-right: clamp(7.5rem, 10vw, 9.5rem);
}

.footer-language-globe {
  width: 18px;
  height: 18px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  color: rgba(219, 231, 250, 0.84);
}

.footer-language-globe svg {
  width: 18px;
  height: 18px;
  stroke: currentColor;
}

.footer-language-globe img {
  display: block;
  width: 18px;
  height: 18px;
}

.footer-language-dropdown {
  position: relative;
}

.footer-language-toggle {
  position: absolute;
  opacity: 0;
  pointer-events: none;
}

.footer-language-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
  min-height: 42px;
  padding: 0 0.95rem;
  border-radius: 999px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(255, 255, 255, 0.02);
  color: #dbe7fa;
  cursor: pointer;
}

.footer-language-toggle:checked + .footer-language-pill {
  border-color: rgba(109, 165, 255, 0.18);
  background: rgba(255, 255, 255, 0.03);
}

.footer-language-flag {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 18px;
  height: 13px;
  line-height: 1;
  border-radius: 2px;
  overflow: hidden;
}

.footer-language-flag svg {
  display: block;
  width: 18px;
  height: 13px;
}

.footer-language-flag img {
  display: block;
  width: 18px;
  height: 13px;
}

.footer-language-caret {
  width: 24px;
  height: 24px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  border-radius: 999px;
  border: 1px solid rgba(109, 165, 255, 0.18);
  background: rgba(255, 255, 255, 0.04);
  color: rgba(219, 231, 250, 0.82);
  line-height: 1;
  transition: transform 160ms ease;
}

.footer-language-caret svg {
  width: 12px;
  height: 12px;
  display: block;
  stroke: currentColor;
}

.footer-language-toggle:checked + .footer-language-pill .footer-language-caret {
  transform: rotate(180deg);
}

.footer-language-menu {
  position: absolute;
  right: 0;
  bottom: calc(100% + 0.7rem);
  min-width: 188px;
  display: grid;
  gap: 0;
  padding: 0.5rem 0;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  background: rgba(10, 14, 24, 0.98);
  box-shadow: 0 22px 50px rgba(0, 0, 0, 0.34);
  overflow: hidden;
  opacity: 0;
  visibility: hidden;
  pointer-events: none;
  transform: translateY(6px);
  transition: opacity 160ms ease, transform 160ms ease, visibility 160ms ease;
}

.footer-language-toggle:checked ~ .footer-language-menu {
  opacity: 1;
  visibility: visible;
  pointer-events: auto;
  transform: translateY(0);
}

.footer-language-option {
  display: flex;
  align-items: center;
  gap: 0.75rem;
  padding: 0.78rem 1rem;
  color: #eef5ff;
  font-size: 0.98rem;
  font-weight: 500;
  line-height: 1.2;
  white-space: nowrap;
  text-decoration: none;
}

.footer-language-option.selected {
  color: var(--blue-strong);
  font-weight: 600;
}

.footer-language-option:hover {
  background: rgba(255, 255, 255, 0.03);
}

.footer-language-option-flag {
  min-width: 18px;
  display: inline-flex;
  align-items: center;
  justify-content: center;
  font-size: 0.95rem;
  line-height: 1;
}

.footer-language-option-flag img {
  display: block;
  width: 18px;
  height: 13px;
  border-radius: 2px;
}

.dashboard-note {
  margin-top: 0.6rem;
  color: var(--muted);
  line-height: 1.75;
}

.login-panel {
  padding: 1.7rem;
}

.auth-page-shell {
  min-height: calc(100vh - 16rem);
  display: flex;
  flex-direction: column;
  align-items: center;
  justify-content: center;
  gap: 1.2rem;
  padding-top: 2.6rem;
}

.auth-card {
  width: min(100%, 420px);
  padding: 1.85rem 1.9rem 1.7rem;
  border-radius: 24px;
  border: 1px solid rgba(87, 160, 255, 0.58);
  background:
    linear-gradient(180deg, rgba(10, 14, 22, 0.96), rgba(9, 12, 18, 0.94));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 0 36px rgba(67, 129, 255, 0.14),
    0 28px 72px rgba(0, 0, 0, 0.46);
}

.auth-logo-shell {
  width: 96px;
  height: 48px;
  margin: 0 auto 1.25rem;
  display: flex;
  align-items: center;
  justify-content: center;
  border-radius: 0.35rem;
  background: linear-gradient(180deg, rgba(9, 13, 22, 0.98), rgba(6, 9, 16, 0.98));
  box-shadow:
    0 0 0 1px rgba(255, 255, 255, 0.02) inset,
    0 0 32px rgba(116, 81, 255, 0.2);
}

.auth-logo-image {
  width: 68px;
  height: auto;
  display: block;
}

.auth-logo-fallback {
  font-family: "Sora", sans-serif;
  font-size: 1.5rem;
  font-weight: 700;
  color: var(--text);
}

.auth-title {
  margin: 0;
  text-align: center;
  font-family: "Sora", sans-serif;
  font-size: clamp(2rem, 3vw, 2.25rem);
  letter-spacing: -0.06em;
}

.auth-title-helper {
  margin: 0.45rem 0 0;
  text-align: center;
  color: var(--muted);
}

.auth-kicker {
  margin-bottom: 0.65rem;
  text-align: center;
  color: var(--cyan);
  font-size: 0.8rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
}

.auth-subtitle {
  margin: 0.7rem 0 0;
  text-align: center;
  color: var(--cyan);
  font-size: 0.8rem;
  font-weight: 700;
  letter-spacing: 0.12em;
}

.auth-register-button {
  display: flex;
  align-items: center;
  justify-content: center;
  min-height: 48px;
  margin-top: 1rem;
  border-radius: 16px;
  background: linear-gradient(90deg, #7357ff 0%, #2fa2ff 100%);
  box-shadow: 0 14px 38px rgba(76, 138, 244, 0.3);
  color: #ffffff;
  font-weight: 700;
  text-decoration: none;
}

.auth-form-shell {
  display: grid;
}

.auth-form-shell input,
.auth-form-shell button {
  font: inherit;
}

.auth-social-stack {
  display: grid;
  gap: 0.7rem;
  margin-top: 0.65rem;
}

.auth-social-button {
  display: flex;
  align-items: center;
  justify-content: center;
  gap: 0.7rem;
  min-height: 48px;
  padding: 0 1rem;
  border-radius: 15px;
  border: 1px solid rgba(255, 255, 255, 0.07);
  background: rgba(255, 255, 255, 0.03);
  color: var(--text);
  text-decoration: none;
  transition: transform 160ms ease, border-color 160ms ease, background 160ms ease;
}

.auth-social-button:hover {
  transform: translateY(-1px);
  border-color: rgba(105, 168, 255, 0.28);
  background: rgba(255, 255, 255, 0.045);
}

.auth-social-icon {
  display: inline-flex;
  align-items: center;
  justify-content: center;
  width: 18px;
  font-size: 1.05rem;
  font-weight: 800;
  line-height: 1;
}

.auth-social-icon.google {
  color: #7ce6ff;
}

.auth-social-icon.apple {
  color: #f4f8ff;
}

.auth-separator {
  position: relative;
  margin: 1.35rem 0 1.05rem;
  text-align: center;
  color: var(--muted);
  font-size: 0.9rem;
}

.auth-separator::before,
.auth-separator::after {
  content: "";
  position: absolute;
  top: 50%;
  width: calc(50% - 92px);
  height: 1px;
  background: rgba(255, 255, 255, 0.08);
}

.auth-separator::before {
  left: 0;
}

.auth-separator::after {
  right: 0;
}

.auth-field-group {
  display: grid;
  gap: 0.5rem;
  margin-top: 0.95rem;
}

.auth-preview-stack {
  margin-top: 0.9rem;
}

.auth-preview-stack .auth-field-group {
  margin-top: 0.8rem;
}

.auth-preview-stack .auth-field-group:first-child {
  margin-top: 0;
}

.auth-field-group label {
  color: var(--text);
  font-weight: 600;
}

.auth-input-shell {
  min-height: 50px;
  display: flex;
  align-items: center;
  padding: 0 1rem;
  border-radius: 15px;
  border: 1px solid rgba(255, 255, 255, 0.07);
  background: rgba(255, 255, 255, 0.05);
  color: rgba(244, 248, 255, 0.56);
}

.auth-input-shell.password {
  letter-spacing: 0.25em;
  font-size: 1.05rem;
}

.auth-meta-row {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 0.75rem;
  margin: 1rem 0 1.3rem;
}

.auth-check-row {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
  color: var(--muted);
  font-size: 0.92rem;
  cursor: pointer;
  user-select: none;
}

.auth-checkbox-input {
  appearance: none;
  width: 14px;
  height: 14px;
  margin: 0;
  display: inline-grid;
  place-items: center;
  border-radius: 4px;
  border: 1px solid rgba(255, 255, 255, 0.22);
  background: rgba(255, 255, 255, 0.02);
  cursor: pointer;
}

.auth-checkbox-input::after {
  content: "";
  width: 7px;
  height: 4px;
  border: 2px solid #ffffff;
  border-top: none;
  border-right: none;
  transform: rotate(-45deg) scale(0);
  transition: transform 140ms ease;
}

.auth-checkbox-input:checked {
  background: linear-gradient(180deg, #5d97ff 0%, #2fa2ff 100%);
  border-color: rgba(93, 151, 255, 0.78);
  box-shadow: 0 0 0 1px rgba(93, 151, 255, 0.16);
}

.auth-checkbox-input:checked::after {
  transform: rotate(-45deg) scale(1);
}

.auth-checkbox-input:focus-visible {
  outline: none;
  box-shadow: 0 0 0 3px rgba(47, 162, 255, 0.12);
}

.auth-forgot-link,
.auth-inline-link {
  color: #5d97ff;
  text-decoration: none;
}

.auth-forgot-link {
  font-size: 0.92rem;
}

.auth-submit-button {
  display: flex;
  align-items: center;
  justify-content: center;
  width: 100%;
  min-height: 52px;
  border-radius: 16px;
  border: none;
  background: linear-gradient(90deg, #7357ff 0%, #2fa2ff 100%);
  box-shadow: 0 14px 38px rgba(76, 138, 244, 0.3);
  color: #ffffff;
  font-weight: 700;
  text-decoration: none;
  cursor: pointer;
  appearance: none;
}

.auth-bottom-copy {
  margin-top: 1.15rem;
  text-align: center;
  color: var(--muted);
}

.auth-inline-link {
  font-weight: 700;
}

.auth-trust-row {
  display: flex;
  flex-wrap: wrap;
  align-items: center;
  justify-content: center;
  gap: 1.1rem;
  color: var(--muted);
  font-size: 0.9rem;
}

.auth-trust-pill {
  display: inline-flex;
  align-items: center;
  gap: 0.55rem;
}

.auth-trust-dot {
  width: 8px;
  height: 8px;
  border-radius: 999px;
  background: linear-gradient(180deg, var(--teal), var(--blue));
  box-shadow: 0 0 0 6px rgba(24, 209, 199, 0.08);
}

div[data-testid="stTextArea"] textarea,
div[data-testid="stTextInput"] input,
div[data-testid="stNumberInput"] input,
div[data-baseweb="select"] > div,
div[data-baseweb="textarea"] textarea {
  background: rgba(255, 255, 255, 0.03) !important;
  border: 1px solid rgba(255, 255, 255, 0.08) !important;
  color: var(--text) !important;
  border-radius: 16px !important;
}

div[data-testid="stTextArea"] label,
div[data-testid="stTextInput"] label,
div[data-testid="stNumberInput"] label,
div[data-testid="stSelectbox"] label,
div[data-testid="stRadio"] label {
  color: var(--text) !important;
  font-weight: 600 !important;
}

div[data-testid="stForm"] div[data-baseweb="base-input"],
div[data-testid="stForm"] div[data-baseweb="base-input"] > div,
div[data-testid="stForm"] div[data-baseweb="textarea"],
div[data-testid="stForm"] div[data-baseweb="select"] > div {
  border: 1.5px solid transparent !important;
  border-radius: 22px !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow: none !important;
}

div[data-testid="stForm"] div[data-baseweb="base-input"]:focus-within,
div[data-testid="stForm"] div[data-baseweb="textarea"]:focus-within,
div[data-testid="stForm"] div[data-baseweb="select"] > div:focus-within {
  box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.22) !important;
}

div[data-testid="stForm"] div[data-testid="stTextInput"] input,
div[data-testid="stForm"] div[data-testid="stTextArea"] textarea,
div[data-testid="stForm"] div[data-testid="stSelectbox"] input {
  background: transparent !important;
  color: #f5fbff !important;
}

div[data-testid="stForm"] div[data-testid="stTextInput"] input::placeholder,
div[data-testid="stForm"] div[data-testid="stTextArea"] textarea::placeholder {
  color: rgba(219, 231, 250, 0.54) !important;
}

div[data-testid="stForm"] div[data-testid="stTextArea"] textarea {
  min-height: 176px !important;
}

div[data-testid="stForm"] div[data-testid="stTextArea"] label,
div[data-testid="stForm"] div[data-testid="stTextInput"] label,
div[data-testid="stForm"] div[data-testid="stSelectbox"] label {
  font-family: "Sora", sans-serif !important;
  font-weight: 700 !important;
}

div[data-testid="stForm"] div[data-baseweb="input"],
div[data-testid="stForm"] div[data-baseweb="input"] > div,
div[data-testid="stForm"] div[data-baseweb="textarea"],
div[data-testid="stForm"] div[data-baseweb="select"] > div,
div[data-testid="stForm"] [data-testid="stTextInputRootElement"],
div[data-testid="stForm"] [data-testid="stTextAreaRootElement"] {
  border: 1.5px solid transparent !important;
  border-radius: 22px !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow: none !important;
}

div[data-testid="stForm"] div[data-baseweb="input"]:focus-within,
div[data-testid="stForm"] div[data-baseweb="textarea"]:focus-within,
div[data-testid="stForm"] div[data-baseweb="select"] > div:focus-within {
  box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.22) !important;
}

div[data-testid="stForm"] div[data-baseweb="input"] input,
div[data-testid="stForm"] div[data-baseweb="textarea"] textarea,
div[data-testid="stForm"] div[data-baseweb="select"] input {
  background: transparent !important;
  border: none !important;
  color: #f5fbff !important;
}

div[data-testid="stForm"] div[data-baseweb="input"] input {
  min-height: 50px !important;
}

div[data-testid="stForm"] div[data-testid="stFormSubmitButton"] button {
  min-height: 52px;
  border: none !important;
  border-radius: 18px !important;
  color: #ffffff !important;
  background: linear-gradient(180deg, rgba(115, 87, 255, 0.95), rgba(47, 162, 255, 0.98)) !important;
  box-shadow: 0 0 26px rgba(78, 143, 255, 0.22) !important;
  font-weight: 700 !important;
}

div[data-testid="stForm"] div[data-testid="stFormSubmitButton"] button:hover {
  filter: brightness(1.03);
  box-shadow: 0 0 32px rgba(78, 143, 255, 0.28) !important;
}

div[data-testid="stNumberInput"] button {
  color: var(--muted) !important;
}

div[data-testid="stFileUploader"] {
  margin: 1.15rem 0 1rem;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] {
  border: 1.5px solid transparent !important;
  border-radius: 22px !important;
  padding: 0.9rem 1rem !important;
  background:
    linear-gradient(180deg, rgba(7, 14, 28, 0.99), rgba(6, 12, 24, 0.97)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow:
    0 0 0 1px rgba(68, 157, 255, 0.08),
    0 0 26px rgba(17, 225, 245, 0.08),
    0 14px 34px rgba(0, 0, 0, 0.18) !important;
  transition: box-shadow 180ms ease, transform 180ms ease, filter 180ms ease !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"]:hover {
  box-shadow:
    0 0 0 1px rgba(17, 225, 245, 0.18),
    0 0 30px rgba(79, 123, 255, 0.12),
    0 18px 38px rgba(0, 0, 0, 0.22) !important;
  filter: brightness(1.02);
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] {
  color: #f4f8ff !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] svg {
  color: #73d7ff !important;
  stroke: currentColor !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] p,
div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] [data-testid="stFileUploaderDropzoneInstructions"] span,
div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] small {
  color: inherit !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] small {
  color: rgba(220, 233, 255, 0.66) !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] button {
  min-height: 44px !important;
  display: inline-flex !important;
  align-items: center !important;
  justify-content: center !important;
  border: 1.5px solid transparent !important;
  border-radius: 18px !important;
  color: #eef5ff !important;
  background:
    linear-gradient(180deg, rgba(10, 16, 28, 0.98), rgba(8, 13, 24, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.08),
    0 10px 24px rgba(0, 0, 0, 0.18) !important;
  font-weight: 700 !important;
  cursor: pointer !important;
  pointer-events: auto !important;
  position: relative !important;
  z-index: 2 !important;
}

div[data-testid="stFileUploader"] section[data-testid="stFileUploaderDropzone"] button:hover {
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.16),
    0 12px 28px rgba(0, 0, 0, 0.22) !important;
  filter: brightness(1.03);
}

.st-key-task_architect_upload_generate {
  margin-top: 1.2rem;
}

.st-key-task_architect_upload_generate div.stButton > button,
.st-key-task_architect_upload_generate button {
  min-height: 44px !important;
  border: 1.5px solid transparent !important;
  border-radius: 18px !important;
  color: #eef5ff !important;
  background:
    linear-gradient(180deg, rgba(109, 83, 255, 0.98), rgba(51, 169, 255, 0.98)) padding-box,
    linear-gradient(90deg, rgba(129, 89, 255, 0.98), rgba(42, 191, 234, 0.98)) border-box !important;
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.1),
    0 14px 30px rgba(59, 111, 255, 0.2) !important;
  font-family: "Sora", sans-serif !important;
  font-weight: 700 !important;
  white-space: nowrap !important;
}

.st-key-task_architect_upload_generate div.stButton > button:hover,
.st-key-task_architect_upload_generate button:hover {
  box-shadow:
    0 0 0 1px rgba(117, 231, 255, 0.16),
    0 18px 36px rgba(59, 111, 255, 0.26) !important;
  filter: brightness(1.04);
}

div.stButton > button,
div.stDownloadButton > button {
  min-height: 48px;
  border-radius: 16px;
  border: 1px solid rgba(255, 255, 255, 0.08);
  color: white;
  background: linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%);
  box-shadow: 0 14px 38px rgba(122, 67, 255, 0.24);
  font-weight: 700;
}

[data-testid="stSidebar"] div.stButton > button {
  width: 100%;
  justify-content: center;
  min-height: 52px;
  padding: 0 1rem;
  border: 1.5px solid transparent !important;
  border-radius: 16px;
  color: #f5fbff !important;
  background-color: transparent !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow: none !important;
  font-weight: 700;
  transition: transform 160ms ease, box-shadow 160ms ease, filter 160ms ease;
}

[data-testid="stSidebar"] div.stButton > button:hover {
  transform: translateY(-1px);
  filter: brightness(1.03);
  box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.16) !important;
}

[data-testid="stSidebar"] div.stButton > button[kind="primary"] {
  color: #ffffff;
  border-color: transparent !important;
  background-color: transparent !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.28) !important;
}

div.stButton > button[kind="secondary"] {
  background: rgba(255, 255, 255, 0.04);
  box-shadow: none;
}

[data-testid="stSidebar"] div.stButton > button[kind="secondary"] {
  color: #f5fbff !important;
  background-color: transparent !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
  box-shadow: none !important;
}

div[data-testid="stMarkdownContainer"] p code {
  background: rgba(255, 255, 255, 0.08);
  color: var(--text);
}

div[data-testid="stPlotlyChart"] {
  border: 1px solid transparent;
  border-radius: 24px;
  background:
    linear-gradient(180deg, rgba(10, 20, 34, 0.96), rgba(6, 12, 24, 0.88)) padding-box,
    linear-gradient(135deg, rgba(155, 92, 255, 0.32), rgba(100, 162, 255, 0.18), rgba(24, 209, 199, 0.18)) border-box;
  padding: 0.6rem;
  box-shadow: 0 26px 54px rgba(0, 0, 0, 0.22);
}

div[data-testid="stPlotlyChart"] > div {
  border-radius: 18px;
  overflow: hidden;
}

div[data-testid="stDataFrame"] {
  border: 1px solid rgba(255, 255, 255, 0.06);
  border-radius: 24px;
  overflow: hidden;
}

div[data-testid="stAlert"] {
  border-radius: 18px;
}

div[data-testid="stExpander"] {
  border: 1.5px solid transparent !important;
  border-radius: 22px !important;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, rgba(155, 92, 255, 0.5) 0%, rgba(100, 162, 255, 0.34) 52%, rgba(24, 209, 199, 0.36) 100%) border-box !important;
  box-shadow: none !important;
  overflow: hidden;
  margin-top: 1rem;
}

div[data-testid="stExpander"] details,
div[data-testid="stExpander"] > div {
  background: transparent !important;
}

div[data-testid="stExpander"] summary {
  min-height: 56px;
  align-items: center;
}

div[data-testid="stExpander"] summary p {
  color: #f5fbff !important;
  font-weight: 700 !important;
}

.faq-empty-box {
  min-height: 56px;
  margin-top: 1rem;
  border: 1.5px solid transparent;
  border-radius: 22px;
  background:
    linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
    linear-gradient(90deg, rgba(155, 92, 255, 0.5) 0%, rgba(100, 162, 255, 0.34) 52%, rgba(24, 209, 199, 0.36) 100%) border-box;
}

@media (max-width: 1100px) {
  .nav-shell,
  .hero-grid,
  .page-grid,
  .dashboard-hero-grid,
  .timeline-grid,
  .workspace-grid,
  .footer-top-grid,
  .footer-mid-row {
    grid-template-columns: 1fr;
  }

  .metric-showcase-grid,
  .process-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .nav-shell {
    position: relative;
    flex-direction: column;
    align-items: flex-start;
  }

  .top-nav-links,
  .nav-actions,
  .top-nav-meta {
    width: 100%;
  }

  .home-preview-grid {
    flex-direction: column;
  }

  .contact-right-balance {
    min-height: auto;
  }

  .home-proof-strip {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .home-intelligence-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .home-testimonial-grid {
    grid-template-columns: 1fr;
  }

  .footer-link-grid {
    flex-wrap: wrap;
  }

  .footer-compliance-row {
    justify-content: flex-start;
  }
}

@media (max-width: 900px) {
  .kpi-grid,
  .detail-grid,
  .ai-forecast-grid,
  .guide-how-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .task-architect-hero {
    padding: 1.25rem 1.2rem 1.4rem;
    border-radius: 24px;
  }

  .task-architect-parameter-grid {
    grid-template-columns: 1fr;
  }

  .guide-modules-layout {
    grid-template-columns: 1fr;
  }

  .pricing-showcase {
    flex-direction: column;
  }

  .pricing-showcase .pricing-plan-card {
    max-width: none;
  }

  .about-mv-grid,
  .about-values-grid,
  .about-team-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .about-timeline::before {
    left: 18px;
    transform: none;
  }

  .about-timeline-row {
    grid-template-columns: 36px minmax(0, 1fr);
    gap: 0.9rem;
  }

  .about-timeline-axis {
    grid-column: 1;
  }

  .about-timeline-side {
    grid-column: 2;
  }

  .about-timeline-row .about-timeline-side:empty {
    display: none;
  }

  .metric-showcase-grid,
  .process-grid {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .command-bar-grid {
    grid-template-columns: 1fr;
    justify-items: center;
  }

  .ai-forecast-steps {
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }

  .command-bar-copy,
  .command-bar-status {
    text-align: center;
    justify-content: center;
  }

  .footer-content {
    width: min(calc(100% - 32px), var(--nav-width));
    padding-top: 3.3rem;
  }

  .footer-link-grid {
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
  }
}

@media (max-width: 720px) {
  .block-container {
    padding-left: 0.9rem !important;
    padding-right: 0.9rem !important;
  }

  .top-nav-meta,
  .nav-actions {
    flex-direction: column;
    align-items: stretch;
  }

  .kpi-grid,
  .detail-grid,
  .ai-forecast-grid,
  .guide-modules-grid,
  .guide-how-grid,
  .metric-showcase-grid,
  .process-grid {
    grid-template-columns: 1fr;
  }

  .about-mv-grid,
  .about-values-grid,
  .about-team-grid {
    grid-template-columns: 1fr;
  }

  .about-team-head {
    justify-content: center;
  }

  .about-team-section {
    text-align: center;
  }

  .hero-card,
  .insight-card,
  .hero-visual,
  .page-hero-card,
  .dashboard-banner {
    padding: 1.45rem;
  }

  .task-architect-panel,
  .task-architect-empty {
    padding-left: 1.2rem;
    padding-right: 1.2rem;
  }

  .footer-link-grid {
    grid-template-columns: 1fr;
  }

  .home-proof-strip {
    grid-template-columns: 1fr;
  }

  .home-intelligence-grid {
    grid-template-columns: 1fr;
  }

  .footer-subscribe-row {
    flex-direction: column;
    align-items: stretch;
  }

  .footer-input-shell,
  .footer-subscribe-btn {
    width: 100%;
  }

  .footer-bottom-row {
    align-items: flex-start;
  }

  .footer-language-shell {
    margin-right: 0;
  }

  .floating-shortcut-wrap {
    left: 0.7rem;
    bottom: 0.8rem;
  }

  .about-proto-logo-shell {
    width: min(100%, 228px);
    height: 146px;
  }

  .about-mv-card,
  .about-timeline-card,
  .about-member-card {
    padding-left: 1.15rem;
    padding-right: 1.15rem;
  }

  .ai-forecast-steps {
    grid-template-columns: 1fr;
  }

  .ai-forecast-process-card,
  .ai-forecast-card {
    padding-left: 1.2rem;
    padding-right: 1.2rem;
  }

  .guide-activity-item {
    grid-template-columns: auto 1fr;
  }

  .guide-activity-time {
    grid-column: 2;
    justify-self: start;
  }
}
</style>
"""


def _get_query_param(name: str, default: str) -> str:
    try:
        value = st.query_params.get(name, default)
    except Exception:
        legacy = st.experimental_get_query_params()
        value = legacy.get(name, [default])

    if isinstance(value, list):
        return str(value[0]) if value else default
    return str(value)


def _get_query_flag(name: str) -> bool:
    return _get_query_param(name, "").strip().lower() in {"1", "true", "yes", "y", "on"}


def _set_query_page(page: str) -> None:
    language_code = st.session_state.get("saas_language_code", "en")
    try:
        st.query_params["page"] = page
        st.query_params["lang"] = language_code
    except Exception:
        st.experimental_set_query_params(page=page, lang=language_code)


def _sync_auth_email_from_query() -> None:
    auth_email = _get_query_param("auth_email", "").strip()
    if auth_email:
        st.session_state.saas_user_email = auth_email


def _sync_contact_email_from_query() -> None:
    contact_email = _get_query_param("newsletter_email", "").strip()
    if contact_email:
        st.session_state.contact_email = contact_email


def _sync_registration_email_from_query() -> None:
    registration_email = _get_query_param("register_email", "").strip()
    if registration_email:
        st.session_state.registration_email = registration_email
        st.session_state.saas_user_email = registration_email


def _sync_language_from_query() -> None:
    language_code = _get_query_param("lang", "en").strip() or "en"
    if language_code not in SUPPORTED_LANGUAGE_CODES:
        language_code = "en"
    st.session_state.saas_language_code = language_code
    st.session_state.saas_language = LANGUAGE_LABELS.get(language_code, "English")


def _sync_newsletter_subscription_flow(page: str) -> None:
    if page != "subscription-confirmation":
        return

    recipient_email = _get_query_param("newsletter_email", "").strip()
    if recipient_email:
        st.session_state.contact_email = recipient_email

    confirmed = _get_query_flag("newsletter_confirmed")
    st.session_state.newsletter_confirmed = confirmed
    if confirmed:
        st.session_state.newsletter_delivery_status = "confirmed"
        st.session_state.newsletter_delivery_message = f"{recipient_email} has been confirmed for certAIn updates."
        return

    if not recipient_email:
        return

    request_key = f"{recipient_email}|{st.session_state.get('saas_language_code', 'en')}"
    if st.session_state.get("newsletter_last_request_key") == request_key:
        return

    confirmation_link = build_subscription_confirmation_link(
        recipient_email,
        settings.app_base_url,
        lang=st.session_state.get("saas_language_code", "en"),
    )
    result = send_subscription_confirmation_email(recipient_email, confirmation_link, settings)
    if result.sent:
        st.session_state.newsletter_last_request_key = request_key
    st.session_state.newsletter_delivery_status = result.status
    st.session_state.newsletter_delivery_message = result.message


def _sync_registration_flow(page: str) -> None:
    if page != "registration-confirmation":
        return

    recipient_email = _get_query_param("register_email", "").strip()
    if recipient_email:
        st.session_state.registration_email = recipient_email
        st.session_state.saas_user_email = recipient_email

    if not recipient_email:
        return

    request_key = f"{recipient_email}|{st.session_state.get('saas_language_code', 'en')}"
    if st.session_state.get("registration_last_request_key") == request_key:
        return

    result = send_registration_confirmation_email(recipient_email, settings)
    if result.sent:
        st.session_state.registration_last_request_key = request_key
    st.session_state.registration_delivery_status = result.status
    st.session_state.registration_delivery_message = result.message


def _resolve_page() -> str:
    page = _get_query_param("page", "home").strip().lower() or "home"
    return page if page in PAGE_TITLES else "home"


def _href(page: str, lang: str | None = None) -> str:
    language_code = lang or st.session_state.get("saas_language_code", "en")
    return f"?page={page}&lang={language_code}"


def _trial_href() -> str:
    return _href("login")


def _is_dashboard_page(page: str) -> bool:
    return page in DASHBOARD_PAGE_KEYS


def _top_nav_active_key(page: str) -> str:
    return "dashboard" if _is_dashboard_page(page) else page


def _github_profile_href() -> str:
    email = st.session_state.get("saas_user_email", "").strip().lower()
    if not email or "@" not in email:
        return FOOTER_SOCIAL_LINKS["github"]

    local_part = email.split("@", 1)[0]
    username = re.sub(r"[^a-z0-9-]+", "-", local_part)
    username = re.sub(r"-{2,}", "-", username).strip("-")[:39]
    if not username:
        return FOOTER_SOCIAL_LINKS["github"]
    return f"https://github.com/{username}"


def _assistant_profile(page: str) -> dict[str, Any]:
    if page == "home":
        return {
            "title": "Platform Guide",
            "summary": "Use this space to understand the product story, core workflow, and where to begin.",
            "suggestions": [
                "What can certAIn do?",
                "Where should I start?",
                "Explain the platform flow",
            ],
        }
    if page == "product":
        return {
            "title": "Product Walkthrough",
            "summary": "This page explains how certAIn connects planning, forecasting, simulation, and reporting.",
            "suggestions": [
                "Summarize the product",
                "What makes this different?",
                "How do the modules connect?",
            ],
        }
    if page == "forecasting":
        return {
            "title": "Forecasting Help",
            "summary": "This page focuses on predictive risk signals, model evidence, and earlier intervention.",
            "suggestions": [
                "Explain forecasting outputs",
                "What is the bundled model?",
                "How should I present this page?",
            ],
        }
    if page == "monte-carlo":
        return {
            "title": "Monte Carlo Help",
            "summary": "This page shows schedule uncertainty, completion ranges, and confidence-based delivery planning.",
            "suggestions": [
                "What does Monte Carlo show?",
                "How do I explain confidence levels?",
                "What should I look at first?",
            ],
        }
    if page == "pricing":
        return {
            "title": "Pricing Help",
            "summary": "Use this page to compare packages and guide trial or enterprise conversations.",
            "suggestions": [
                "Which plan fits a small team?",
                "Explain the trial path",
                "What should I say about enterprise?",
            ],
        }
    if page == "login":
        return {
            "title": "Access Help",
            "summary": "This page supports the demo entry flow before users continue into the platform experience.",
            "suggestions": [
                "How does the sign-in flow work?",
                "Where do these buttons go?",
                "What should a user do next?",
            ],
        }
    if page == "about":
        return {
            "title": "Team Story",
            "summary": "Use this page to explain the capstone story, product vision, and why the platform matters.",
            "suggestions": [
                "Summarize the project story",
                "How should I introduce certAIn?",
                "What problem does this solve?",
            ],
        }
    if page == "contact":
        return {
            "title": "Contact Help",
            "summary": "This page is the handoff point for questions, demos, and follow-up conversations.",
            "suggestions": [
                "What should happen after contact?",
                "How do I present this page?",
                "What is the next step after a demo?",
            ],
        }
    if _is_dashboard_page(page):
        return {
            "title": PAGE_TITLES.get(page, "Workspace"),
            "summary": "You are inside the working platform where teams generate plans, inspect risks, compare scenarios, and prepare executive outputs.",
            "suggestions": [
                "Explain this workspace page",
                "What should I do next here?",
                "How do I present this module?",
            ],
        }
    return {
        "title": PAGE_TITLES.get(page, "Platform"),
        "summary": "I can help explain the current page, point to the next step, or answer questions about the product workflow.",
        "suggestions": [
            "Explain this page",
            "What should I do next?",
            "Summarize certAIn",
        ],
    }


def _render_command_palette(active_page: str) -> None:
    items_markup = "".join(
        dedent(
            f"""
            <a
              href="{_href(page)}"
              class="command-palette-item {'active' if page == active_page else ''}"
              data-command-item
              data-command-text="{(group + ' ' + label + ' ' + description).lower()}"
            >
              <div class="command-palette-item-copy">
                <span class="command-palette-item-title">{label}</span>
                <span class="command-palette-item-desc">{description}</span>
              </div>
              <span class="command-palette-item-group">{group}</span>
            </a>
            """
        ).strip()
        for group, page, label, description in COMMAND_PALETTE_ITEMS
    )
    st.markdown(
        dedent(
            f"""
            <div class="command-palette-overlay" data-command-palette aria-hidden="true">
              <div class="command-palette-backdrop" data-command-close="true"></div>
              <div class="command-palette-card" role="dialog" aria-modal="true" aria-label="Quick command search">
                <div class="command-palette-head">
                  <div class="command-palette-title">
                    <strong>Quick command search</strong>
                    <span>Press ⌘K on Mac or Ctrl+K on Windows/Linux from any platform page.</span>
                  </div>
                  <button type="button" class="command-palette-close" data-command-close="true" aria-label="Close command search">×</button>
                </div>
                <div class="command-palette-search">
                  <input
                    class="command-palette-input"
                    type="text"
                    placeholder="Search pages, modules, and actions"
                    data-command-input
                  />
                </div>
                <div class="command-palette-list" data-command-list>
                  {items_markup}
                  <div class="command-palette-empty" data-command-empty hidden>No matching command found.</div>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _mount_command_palette_shortcuts() -> None:
    components.html(
        """
        <script>
        (() => {
          const parentWindow = window.parent;
          const parentDocument = parentWindow.document;

          const getPalette = () => parentDocument.querySelector("[data-command-palette]");
          const getInput = () => parentDocument.querySelector("[data-command-input]");
          const getItems = () => Array.from(parentDocument.querySelectorAll("[data-command-item]"));
          const getEmpty = () => parentDocument.querySelector("[data-command-empty]");

          const filterItems = () => {
            const input = getInput();
            if (!input) return;
            const term = (input.value || "").trim().toLowerCase();
            let visibleCount = 0;
            getItems().forEach((item) => {
              const haystack = item.dataset.commandText || item.textContent.toLowerCase();
              const show = !term || haystack.includes(term);
              item.style.display = show ? "" : "none";
              if (show) visibleCount += 1;
            });
            const empty = getEmpty();
            if (empty) empty.hidden = visibleCount > 0;
          };

          const openPalette = () => {
            const palette = getPalette();
            if (!palette) return;
            palette.classList.add("is-open");
            palette.setAttribute("aria-hidden", "false");
            parentDocument.body.classList.add("palette-open");
            parentWindow.requestAnimationFrame(() => {
              const input = getInput();
              if (input) {
                input.focus();
                input.select();
              }
              filterItems();
            });
          };

          const closePalette = () => {
            const palette = getPalette();
            if (!palette) return;
            palette.classList.remove("is-open");
            palette.setAttribute("aria-hidden", "true");
            parentDocument.body.classList.remove("palette-open");
          };

          const attachHandlers = () => {
            if (parentWindow.__certainPaletteInit) return;

            parentDocument.addEventListener("click", (event) => {
              const trigger = event.target.closest("[data-command-trigger]");
              if (trigger) {
                event.preventDefault();
                openPalette();
                return;
              }

              const closer = event.target.closest("[data-command-close]");
              if (closer) {
                event.preventDefault();
                closePalette();
              }
            });

            parentDocument.addEventListener("input", (event) => {
              if (event.target.matches("[data-command-input]")) {
                filterItems();
              }
            });

            parentDocument.addEventListener("keydown", (event) => {
              const isShortcut = (event.metaKey || event.ctrlKey) && event.key.toLowerCase() === "k";
              if (isShortcut) {
                event.preventDefault();
                const palette = getPalette();
                if (palette && palette.classList.contains("is-open")) {
                  closePalette();
                } else {
                  openPalette();
                }
                return;
              }

              if (event.key === "Escape") {
                closePalette();
                return;
              }

              if (event.key === "Enter") {
                const palette = getPalette();
                if (!palette || !palette.classList.contains("is-open")) return;
                const firstVisible = getItems().find((item) => item.style.display !== "none");
                if (firstVisible) {
                  event.preventDefault();
                  firstVisible.click();
                }
              }
            });

            parentWindow.__certainPaletteInit = true;
          };

          const syncPalette = () => {
            if (!getPalette()) {
              parentWindow.setTimeout(syncPalette, 120);
              return;
            }
            attachHandlers();
            filterItems();
          };

          syncPalette();
        })();
        </script>
        """,
        height=0,
        width=0,
    )


def _mount_page_translation() -> None:
    widget_code = LANGUAGE_WIDGET_CODES.get(
        st.session_state.get("saas_language_code", "en"),
        "en",
    )
    included_languages = ",".join(widget_code for _, _, widget_code in LANGUAGE_OPTIONS)
    components.html(
        dedent(
            f"""
            <script>
            (() => {{
              const parentWindow = window.parent;
              const parentDocument = parentWindow.document;
              const targetLanguage = "{widget_code}";
              const includedLanguages = "{included_languages}";

              const ensureContainer = () => {{
                let container = parentDocument.getElementById("certain_google_translate_element");
                if (!container) {{
                  container = parentDocument.createElement("div");
                  container.id = "certain_google_translate_element";
                  container.style.display = "none";
                  parentDocument.body.appendChild(container);
                }}
                return container;
              }};

              const ensureGoogleInit = () => {{
                if (parentWindow.google?.translate?.TranslateElement && !parentWindow.__certainTranslateWidgetReady) {{
                  ensureContainer();
                  new parentWindow.google.translate.TranslateElement(
                    {{
                      pageLanguage: "en",
                      includedLanguages,
                      autoDisplay: false,
                      layout: parentWindow.google.translate.TranslateElement.InlineLayout.SIMPLE,
                    }},
                    "certain_google_translate_element",
                  );
                  parentWindow.__certainTranslateWidgetReady = true;
                }}
              }};

              const persistLanguage = () => {{
                const languagePath = targetLanguage === "en" ? "/auto/en" : `/auto/${{targetLanguage}}`;
                parentDocument.cookie = `googtrans=${{languagePath}};path=/;SameSite=Lax`;
                try {{
                  parentWindow.localStorage.setItem("certain.language", targetLanguage);
                }} catch (error) {{
                  // Ignore storage issues in restricted browser contexts.
                }}
              }};

              const applyLanguage = () => {{
                persistLanguage();
                parentDocument.documentElement.lang = targetLanguage === "en" ? "en" : targetLanguage;
                const combo = parentDocument.querySelector(".goog-te-combo");
                if (!combo) return false;
                combo.value = targetLanguage === "en" ? "" : targetLanguage;
                combo.dispatchEvent(new parentWindow.Event("change", {{ bubbles: true }}));
                if (targetLanguage === "en") {{
                  parentWindow.setTimeout(() => {{
                    combo.value = "";
                    combo.dispatchEvent(new parentWindow.Event("change", {{ bubbles: true }}));
                  }}, 60);
                }}
                return true;
              }};

              const bootTranslation = () => {{
                ensureGoogleInit();
                if (applyLanguage()) return;
                parentWindow.setTimeout(bootTranslation, 180);
              }};

              if (!parentWindow.__certainTranslateScriptLoaded) {{
                parentWindow.googleTranslateElementInit = () => {{
                  parentWindow.__certainTranslateScriptLoaded = true;
                  bootTranslation();
                }};
                const script = parentDocument.createElement("script");
                script.src = "https://translate.google.com/translate_a/element.js?cb=googleTranslateElementInit";
                script.async = true;
                parentDocument.body.appendChild(script);
                parentWindow.__certainTranslateScriptLoaded = true;
              }}

              bootTranslation();
            }})();
            </script>
            """
        ),
        height=0,
        width=0,
    )


def _render_ai_assistant(active_page: str) -> None:
    profile = _assistant_profile(active_page)
    suggestions_markup = "".join(
        f'<button type="button" class="ai-assistant-chip" data-ai-suggestion="{suggestion}">{suggestion}</button>'
        for suggestion in profile["suggestions"]
    )
    st.markdown(
        dedent(
            f"""
            <div class="ai-assistant-shell" data-ai-assistant data-ai-page="{active_page}">
              <div class="ai-assistant-panel" data-ai-assistant-panel aria-hidden="true">
                <div class="ai-assistant-head">
                  <div class="ai-assistant-head-copy">
                    <span class="ai-assistant-kicker">AI Assistant</span>
                    <strong>{profile["title"]}</strong>
                    <span>{profile["summary"]}</span>
                  </div>
                  <button type="button" class="ai-assistant-close" data-ai-close="true" aria-label="Close AI assistant">×</button>
                </div>
                <div class="ai-assistant-suggestions">
                  {suggestions_markup}
                </div>
                <div class="ai-assistant-messages" data-ai-messages></div>
                <form class="ai-assistant-form" data-ai-form>
                  <input
                    class="ai-assistant-input"
                    type="text"
                    placeholder="Ask about risks, pricing, forecasting, or next steps"
                    data-ai-input
                  />
                  <button type="submit" class="ai-assistant-send">Send</button>
                </form>
              </div>
              <button type="button" class="ai-assistant-launcher" data-ai-toggle="true" aria-label="Open AI assistant">
                <span class="ai-assistant-launcher-icon">✦</span>
                <span class="ai-assistant-launcher-copy">
                  <strong>AI Assistant</strong>
                  <span>Ask {_brand_name_html()}</span>
                </span>
              </button>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _mount_ai_assistant(active_page: str) -> None:
    profile_json = json.dumps(_assistant_profile(active_page))
    active_page_json = json.dumps(active_page)
    components.html(
        dedent(
            f"""
            <script>
            (() => {{
              const parentWindow = window.parent;
              const parentDocument = parentWindow.document;
              const profile = {profile_json};
              const activePage = {active_page_json};
              const storeKey = "certain.ai.assistant.messages";
              const openKey = "certain.ai.assistant.open";

              const escapeHtml = (value) =>
                String(value)
                  .replace(/&/g, "&amp;")
                  .replace(/</g, "&lt;")
                  .replace(/>/g, "&gt;")
                  .replace(/"/g, "&quot;")
                  .replace(/'/g, "&#39;");

              const getShell = () => parentDocument.querySelector("[data-ai-assistant]");
              const getPanel = () => parentDocument.querySelector("[data-ai-assistant-panel]");
              const getMessages = () => parentDocument.querySelector("[data-ai-messages]");
              const getInput = () => parentDocument.querySelector("[data-ai-input]");

              const readMessages = () => {{
                try {{
                  return JSON.parse(parentWindow.sessionStorage.getItem(storeKey) || "[]");
                }} catch (error) {{
                  return [];
                }}
              }};

              const saveMessages = (messages) => {{
                const trimmed = messages.slice(-14);
                try {{
                  parentWindow.sessionStorage.setItem(storeKey, JSON.stringify(trimmed));
                }} catch (error) {{
                  // Ignore storage issues in restricted browsers.
                }}
              }};

              const introMessage = () =>
                `Hi, I'm your certAIn AI assistant. You're on ${{profile.title}}. ${{profile.summary}}`;

              const ensureMessages = () => {{
                const current = readMessages();
                if (current.length) return current;
                const seeded = [{{ role: "assistant", text: introMessage() }}];
                saveMessages(seeded);
                return seeded;
              }};

              const renderMessages = () => {{
                const messages = ensureMessages();
                const container = getMessages();
                if (!container) return;
                container.innerHTML = messages
                  .map(
                    (message) =>
                      `<div class="ai-assistant-message ${{message.role}}"><span>${{escapeHtml(message.text)}}</span></div>`,
                  )
                  .join("");
                container.scrollTop = container.scrollHeight;
              }};

              const nextStepReply = () => {{
                if (activePage === "home") {{
                  return "Start with Product for the platform story, then move to AI Forecasting or Monte Carlo to show the evidence-driven workflow.";
                }}
                if (activePage === "pricing") {{
                  return "Use Pricing to position the trial flow first, then move into the dashboard or simulation pages to demonstrate value.";
                }}
                if (activePage === "login") {{
                  return "From here, the sign-in and register actions move users into the pricing flow before they continue deeper into the platform.";
                }}
                if (activePage === "forecasting" || activePage === "monte-carlo") {{
                  return "Highlight the key metric first, explain how the result supports better commitments, then connect it to dashboard decisions or scenario testing.";
                }}
                if (activePage === "about" || activePage === "contact") {{
                  return "Use this page as a closer: summarize the story, invite the next conversation, and route viewers back into the working product pages if needed.";
                }}
                return "Lead with the key signal on this page, explain the decision it supports, and then move to the next workspace or reporting screen.";
              }};

              const composeReply = (message) => {{
                const text = message.toLowerCase();

                if (/pricing|plan|trial|enterprise|team/.test(text)) {{
                  return "The pricing flow is best framed as a guided entry point: start with Team Pilot for smaller evaluations and position Enterprise for broader governance and rollout support.";
                }}

                if (/forecast|predict|prediction|delay|deadline/.test(text)) {{
                  return "Forecasting helps explain which signals suggest delay or delivery risk early enough for teams to intervene before commitments slip.";
                }}

                if (/monte|simulation|probability|confidence/.test(text)) {{
                  return "Monte Carlo is strongest when you describe it as a range-based planning tool: it shows likely finish windows and confidence levels instead of a single brittle date.";
                }}

                if (/risk|classifier|score|driver/.test(text)) {{
                  return "Use the risk views to explain what is driving exposure, which tasks or signals are contributing most, and where mitigation should start first.";
                }}

                if (/task|plan|dependency|dag|critical path/.test(text)) {{
                  return "certAIn turns a plain-language brief into structured tasks, dependencies, and critical-path logic so teams can move from idea to defensible plan faster.";
                }}

                if (/upload|csv|file|import/.test(text)) {{
                  return "Project Upload is presented as a guided intake experience in the product story, with CSV previewing and broader file support represented in the UI.";
                }}

                if (/summary|report|executive|stakeholder/.test(text)) {{
                  return "The reporting flow is meant to translate technical forecast signals into business-ready summaries that stakeholders can review quickly.";
                }}

                if (/language|translate/.test(text)) {{
                  return "The platform language can be changed from the language selector, and the selected language is carried across the shared platform pages.";
                }}

                if (/command|search|cmd\\+k|ctrl\\+k/.test(text)) {{
                  return "Quick command search is available from the shortcut pill or with Command+K on Mac and Ctrl+K on Windows or Linux.";
                }}

                if (/what can you do|help|support|assist/.test(text)) {{
                  return "I can explain the current page, suggest the next best step, summarize product capabilities, and help you present forecasts, risks, Monte Carlo results, pricing, or reporting.";
                }}

                if (/next|start|where should i begin|what should i do/.test(text)) {{
                  return nextStepReply();
                }}

                if (/this page|what is this|where am i|here/.test(text)) {{
                  return `This is the ${{profile.title}} page. ${{profile.summary}}`;
                }}

                return `You're on ${{profile.title}}. ${{profile.summary}} Ask me about forecasting, Monte Carlo, pricing, reporting, or the next best step.`;
              }};

              const updateOpenState = (open, focusInput = false) => {{
                const shell = getShell();
                const panel = getPanel();
                if (!shell || !panel) return;
                shell.classList.toggle("is-open", open);
                panel.setAttribute("aria-hidden", open ? "false" : "true");
                try {{
                  parentWindow.sessionStorage.setItem(openKey, open ? "1" : "0");
                }} catch (error) {{
                  // Ignore storage issues in restricted browsers.
                }}
                if (open && focusInput) {{
                  parentWindow.requestAnimationFrame(() => {{
                    const input = getInput();
                    if (input) input.focus();
                  }});
                }}
              }};

              const handlePrompt = (rawMessage) => {{
                const message = (rawMessage || "").trim();
                if (!message) return;
                const messages = ensureMessages();
                messages.push({{ role: "user", text: message }});
                messages.push({{ role: "assistant", text: composeReply(message) }});
                saveMessages(messages);
                renderMessages();
                const input = getInput();
                if (input) {{
                  input.value = "";
                  input.focus();
                }}
                updateOpenState(true);
              }};

              const attachHandlers = () => {{
                if (parentWindow.__certainAssistantInit) return;

                parentDocument.addEventListener("click", (event) => {{
                  const toggle = event.target.closest("[data-ai-toggle]");
                  if (toggle) {{
                    event.preventDefault();
                    const shell = getShell();
                    const isOpen = shell?.classList.contains("is-open");
                    updateOpenState(!isOpen, true);
                    renderMessages();
                    return;
                  }}

                  const closeButton = event.target.closest("[data-ai-close]");
                  if (closeButton) {{
                    event.preventDefault();
                    updateOpenState(false);
                    return;
                  }}

                  const suggestion = event.target.closest("[data-ai-suggestion]");
                  if (suggestion) {{
                    event.preventDefault();
                    handlePrompt(suggestion.dataset.aiSuggestion || suggestion.textContent || "");
                  }}
                }});

                parentDocument.addEventListener("submit", (event) => {{
                  if (!event.target.matches("[data-ai-form]")) return;
                  event.preventDefault();
                  const input = getInput();
                  if (!input) return;
                  handlePrompt(input.value);
                }});

                parentWindow.__certainAssistantInit = true;
              }};

              const syncAssistant = () => {{
                if (!getShell()) {{
                  parentWindow.setTimeout(syncAssistant, 120);
                  return;
                }}
                attachHandlers();
                renderMessages();
                const shouldOpen = parentWindow.sessionStorage.getItem(openKey) === "1";
                updateOpenState(shouldOpen);
              }};

              syncAssistant();
            }})();
            </script>
            """
        ),
        height=0,
        width=0,
    )


@st.cache_data
def _logo_data_uri() -> str:
    for path, mime_type in (
        (LOGO_PNG_PATH, "image/png"),
        (LEGACY_LOGO_PNG_PATH, "image/png"),
        (LOGO_SVG_PATH, "image/svg+xml"),
    ):
        if path.exists():
            encoded = base64.b64encode(path.read_bytes()).decode("ascii")
            return f"data:{mime_type};base64,{encoded}"
    return ""


@st.cache_data
def _nav_logo_data_uri() -> str:
    for path in (NAV_LOGO_PNG_PATH, LOGO_PNG_PATH, LEGACY_LOGO_PNG_PATH):
        if path.exists():
            encoded = base64.b64encode(path.read_bytes()).decode("ascii")
            return f"data:image/png;base64,{encoded}"
    return ""


@st.cache_data
def _image_data_uri(path: str) -> str:
    image_path = Path(path)
    if not image_path.exists():
        return ""

    mime_type = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
    }.get(image_path.suffix.lower(), "application/octet-stream")
    encoded = base64.b64encode(image_path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"


def _svg_data_uri(svg: str) -> str:
    encoded = base64.b64encode(svg.encode("utf-8")).decode("ascii")
    return f"data:image/svg+xml;base64,{encoded}"


def _brand_name_html() -> str:
    return 'cert<span class="ai-gradient">AI</span>n'


def _brand_name_in_text_html(text: str) -> str:
    return html.escape(text).replace("certAIn", 'cert<span class="ai-gradient about-inline-ai">AI</span>n')


def _brand_lockup(location: str, caption: str | None = None) -> str:
    wordmark = f'<span class="brand-word brand-word-inline notranslate" translate="no">{_brand_name_html()}</span>'
    logo_uri = _logo_data_uri()
    if logo_uri:
        media = (
            '<span class="brand-media with-wordmark">'
            '<span class="brand-logo-shell">'
            f'<img src="{logo_uri}" class="brand-logo" alt="certAIn logo" />'
            "</span>"
            f"{wordmark}"
            "</span>"
        )
    else:
        media = (
            '<span class="brand-media with-wordmark">'
            '<span class="brand-mark" aria-hidden="true">'
            '<span class="brand-triangle"></span>'
            '<span class="brand-core">A</span>'
            "</span>"
            f"{wordmark}"
            "</span>"
        )

    caption_html = ""
    if caption:
        caption_html = f'<span class="brand-copy"><span class="brand-caption">{caption}</span></span>'

    return f'<a class="brand brand-lockup {location}" href="{_href("home")}">{media}{caption_html}</a>'


def _nav_brand_markup() -> str:
    logo_uri = _nav_logo_data_uri()
    if logo_uri:
        media = (
            '<span class="nav-brand-box">'
            f'<img src="{logo_uri}" class="nav-brand-icon" alt="certAIn logo" />'
            "</span>"
        )
    else:
        media = (
            '<span class="nav-brand-box">'
            '<span class="brand-mark" aria-hidden="true">'
            '<span class="brand-triangle"></span>'
            '<span class="brand-core">A</span>'
            "</span>"
            "</span>"
        )

    return (
        f'<a class="nav-brand-zone" href="{_href("home")}">'
        f"{media}"
        f'<span class="nav-brand-word notranslate" translate="no">{_brand_name_html()}</span>'
        "</a>"
    )


def _footer_brand_markup() -> str:
    logo_uri = _nav_logo_data_uri()
    if logo_uri:
        media = (
            '<span class="footer-brand-box">'
            f'<img src="{logo_uri}" class="footer-brand-icon" alt="certAIn logo" />'
            "</span>"
        )
    else:
        media = (
            '<span class="footer-brand-box">'
            '<span class="brand-mark" aria-hidden="true">'
            '<span class="brand-triangle"></span>'
            '<span class="brand-core">A</span>'
            "</span>"
            "</span>"
        )

    return (
        f'<a class="footer-brand-row" href="{_href("home")}">'
        f"{media}"
        f'<span class="footer-brand-wordmark notranslate" translate="no">{_brand_name_html()}</span>'
        "</a>"
    )


def _footer_social_icon_markup(platform: str, label: str) -> str:
    href = _github_profile_href() if platform == "github" else FOOTER_SOCIAL_LINKS.get(platform, "#")
    icons = {
        "linkedin": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.8" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<path d="M16 8a6 6 0 0 1 6 6v7h-4v-7a2 2 0 0 0-2-2 2 2 0 0 0-2 2v7h-4v-7a6 6 0 0 1 6-6z"/>'
            '<rect x="2" y="9" width="4" height="12"/>'
            '<circle cx="4" cy="4" r="2"/>'
            "</svg>"
        ),
        "x": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<path d="M4 4L20 20"/>'
            '<path d="M20 4L4 20"/>'
            "</svg>"
        ),
        "github": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.7" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<path d="M15 22v-4a4.8 4.8 0 0 0-1-3.75c3.14-.35 6.44-1.54 6.44-7a5.44 5.44 0 0 0-1.42-3.78A5.07 5.07 0 0 0 18.91 1S17.73.65 15 2.48a13.38 13.38 0 0 0-6 0C6.27.65 5.09 1 5.09 1A5.07 5.07 0 0 0 5 4.77a5.44 5.44 0 0 0-1.42 3.78c0 5.42 3.3 6.61 6.44 7A4.8 4.8 0 0 0 9 18v4"/>'
            '<path d="M9 18c-4.51 2-5-2-7-2"/>'
            "</svg>"
        ),
        "youtube": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.7" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<path d="M2.5 17a24.12 24.12 0 0 1 0-10 2 2 0 0 1 1.4-1.4 49.56 49.56 0 0 1 16.2 0 2 2 0 0 1 1.4 1.4 24.12 24.12 0 0 1 0 10 2 2 0 0 1-1.4 1.4 49.56 49.56 0 0 1-16.2 0 2 2 0 0 1-1.4-1.4"/>'
            '<path d="M10 15l5-3-5-3z"/>'
            "</svg>"
        ),
        "instagram": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.7" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<rect x="3.5" y="3.5" width="17" height="17" rx="5"/>'
            '<circle cx="12" cy="12" r="4.1"/>'
            '<circle cx="17.3" cy="6.7" r="0.8" fill="currentColor" stroke="none"/>'
            "</svg>"
        ),
        "facebook": (
            '<svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.7" '
            'stroke-linecap="round" stroke-linejoin="round" aria-hidden="true">'
            '<path d="M14 8h3V4h-3c-3 0-5 2-5 5v3H6v4h3v4h4v-4h3.2l.8-4H13V9c0-.7.3-1 1-1z"/>'
            "</svg>"
        ),
    }
    icon = icons.get(platform, "")
    return (
        f'<a class="footer-social-pill" href="{href}" target="_blank" rel="noreferrer noopener" aria-label="{label}" title="{label}">'
        f'<span class="footer-social-icon">{icon}</span>'
        "</a>"
    )


def _footer_compliance_logo_markup(mark: str) -> str:
    logos = {
        "ISO": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<circle class="compliance-stroke" cx="24" cy="24" r="12"/>'
            '<text class="compliance-glyph" x="24" y="27" text-anchor="middle">ISO</text>'
            "</svg>"
        ),
        "S2": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<path class="compliance-stroke" d="M24 7l11 4v9c0 8.2-5.4 14.7-11 17.8C18.4 34.7 13 28.2 13 20v-9l11-4z"/>'
            '<text class="compliance-glyph" x="24" y="27" text-anchor="middle">S2</text>'
            "</svg>"
        ),
        "CA": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<path class="compliance-stroke" d="M9 19h30l-4 8H13l-4-8z"/>'
            '<path class="compliance-detail" d="M13 30h22"/>'
            '<text class="compliance-glyph" x="24" y="18" text-anchor="middle">CA</text>'
            "</svg>"
        ),
        "EU": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<circle class="compliance-stroke" cx="24" cy="24" r="10"/>'
            '<circle class="compliance-detail" cx="24" cy="13" r="1.2"/>'
            '<circle class="compliance-detail" cx="32.5" cy="16.5" r="1.2"/>'
            '<circle class="compliance-detail" cx="35" cy="24" r="1.2"/>'
            '<circle class="compliance-detail" cx="32.5" cy="31.5" r="1.2"/>'
            '<circle class="compliance-detail" cx="24" cy="35" r="1.2"/>'
            '<circle class="compliance-detail" cx="15.5" cy="31.5" r="1.2"/>'
            '<circle class="compliance-detail" cx="13" cy="24" r="1.2"/>'
            '<circle class="compliance-detail" cx="15.5" cy="16.5" r="1.2"/>'
            '<text class="compliance-glyph" x="24" y="27" text-anchor="middle">EU</text>'
            "</svg>"
        ),
        "PI": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<path class="compliance-stroke" d="M24 8l11 6v10l-11 16-11-16V14l11-6z"/>'
            '<text class="compliance-glyph" x="24" y="27" text-anchor="middle">PI</text>'
            "</svg>"
        ),
        "LG": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<path class="compliance-stroke" d="M24 8l12 5v9c0 8.5-5.6 14.7-12 17.7C17.6 36.7 12 30.5 12 22v-9l12-5z"/>'
            '<path class="compliance-detail" d="M18 28h12"/>'
            '<text class="compliance-glyph" x="24" y="23" text-anchor="middle">LG</text>'
            "</svg>"
        ),
        "AP": (
            '<svg viewBox="0 0 48 48" aria-hidden="true">'
            '<path class="compliance-stroke" d="M24 8l13 5v11c0 8.3-5.7 14.2-13 16.8C16.7 38.2 11 32.3 11 24V13l13-5z"/>'
            '<path class="compliance-detail" d="M17 31h14"/>'
            '<text class="compliance-glyph" x="24" y="23" text-anchor="middle">AP</text>'
            "</svg>"
        ),
    }
    return logos.get(mark, mark)


def _inject_styles() -> None:
    st.markdown(GLOBAL_CSS, unsafe_allow_html=True)
    st.markdown(
        """
        <div class="saas-shell">
          <div class="particle-network" aria-hidden="true">
            <div class="particle-group group-a">
              <span class="particle-link" style="left:42px;top:40px;width:54px;transform:rotate(22deg);"></span>
              <span class="particle-link" style="left:95px;top:62px;width:62px;transform:rotate(10deg);"></span>
              <span class="particle-link" style="left:154px;top:73px;width:66px;transform:rotate(8deg);"></span>
              <span class="particle-node teal" style="left:42px;top:40px;--size:10px;--pulse:4.8s;"></span>
              <span class="particle-node blue" style="left:94px;top:61px;--size:8px;--pulse:5.7s;--delay:-1.2s;"></span>
              <span class="particle-node soft" style="left:154px;top:72px;--size:6px;--pulse:6.4s;--delay:-0.8s;"></span>
              <span class="particle-node soft" style="left:220px;top:82px;--size:5px;--pulse:5.4s;--delay:-2.4s;"></span>
              <span class="particle-node violet" style="left:244px;top:18px;--size:6px;--pulse:7.1s;--delay:-3s;"></span>
            </div>
            <div class="particle-group group-b">
              <span class="particle-link" style="left:52px;top:108px;width:58px;transform:rotate(-33deg);"></span>
              <span class="particle-link" style="left:101px;top:75px;width:70px;transform:rotate(29deg);"></span>
              <span class="particle-node blue" style="left:52px;top:108px;--size:9px;--pulse:5.1s;"></span>
              <span class="particle-node teal" style="left:100px;top:74px;--size:8px;--pulse:6.2s;--delay:-1.6s;"></span>
              <span class="particle-node soft" style="left:162px;top:109px;--size:5px;--pulse:5.8s;--delay:-0.9s;"></span>
              <span class="particle-node soft" style="left:34px;top:152px;--size:4px;--pulse:6.6s;--delay:-2.2s;"></span>
            </div>
            <div class="particle-group group-c">
              <span class="particle-link" style="left:70px;top:72px;width:54px;transform:rotate(-28deg);"></span>
              <span class="particle-node blue" style="left:70px;top:72px;--size:7px;--pulse:5.3s;"></span>
              <span class="particle-node violet" style="left:118px;top:46px;--size:5px;--pulse:6.9s;--delay:-1.4s;"></span>
              <span class="particle-node soft" style="left:152px;top:64px;--size:4px;--pulse:7.4s;--delay:-3.2s;"></span>
            </div>
          </div>
          <div class="ambient-orb orb-a"></div>
          <div class="ambient-orb orb-b"></div>
          <div class="ambient-orb orb-c"></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _inject_layout_state(active_page: str) -> None:
    if not _is_dashboard_page(active_page):
        st.markdown(
            """
            <style>
            [data-testid="stSidebar"] {
              display: none !important;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )


def _render_nav(active_page: str) -> None:
    active_key = _top_nav_active_key(active_page)
    nav_links = "".join(
        f'<a class="nav-link {"active" if page == active_key else ""}" href="{_href("simulator" if page == "monte-carlo" else page)}">{label}</a>'
        for page, label in NAV_ITEMS
    )
    st.markdown(
        dedent(
            f"""
            <div class="nav-shell">
              {_nav_brand_markup()}
              <div class="top-nav-pill">
                <span class="top-nav-pill-dot"></span>
                <span>AI-Powered</span>
              </div>
              <div class="top-nav-links">
                {nav_links}
              </div>
              <div class="top-nav-meta">
                <a href="{_trial_href()}" class="btn btn-primary btn-outline-cta">Start Free Trial</a>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_sidebar(active_page: str) -> None:
    with st.sidebar:
        st.markdown(
            f"""
            <div class="sidebar-brand-card">
              {_brand_lockup("sidebar", "AI-Powered Project Management")}
              <div style="margin-top:0.85rem;" class="status-chip chip-low">AI Engine: Active</div>
              <div class="sidebar-mini-label" style="margin-top:0.7rem;">Forecasts, risk intelligence, and executive reporting in one workspace.</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            """
            <div class="sidebar-divider"></div>
            <div class="sidebar-section-label">Dashboard</div>
            """,
            unsafe_allow_html=True,
        )

        for page, label in DASHBOARD_ITEMS:
            button_type = "primary" if page == active_page else "secondary"
            if st.button(label, key=f"dashboard_{page}", use_container_width=True, type=button_type):
                _set_query_page(page)
                st.rerun()

        st.markdown(
            """
            <div class="sidebar-divider"></div>
            <div class="sidebar-shortcut-wrap">
              <button type="button" class="mini-pill sidebar-shortcut-pill sidebar-shortcut-trigger" data-command-trigger="true" aria-label="Open quick command search"><span class="sidebar-shortcut-key">⌘K / Ctrl+K</span></button>
              <div class="sidebar-mini-label sidebar-shortcut-caption">Quick command search</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def _render_command_bar() -> None:
    st.markdown(
        """
        <div class="glass" style="padding:1rem 1.2rem;margin-top:1rem;">
          <div class="command-bar-grid">
            <div class="command-bar-copy">
              <div class="kicker">AI command center</div>
              <div style="margin-top:0.55rem;color:var(--muted);">Project Spotlight available from anywhere in the workspace.</div>
            </div>
            <div class="command-bar-status">
              <div class="status-chip chip-low">AI Engine: Active</div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _render_monte_carlo_explainer() -> None:
    st.markdown(
        """
        <div class="page-shell">
          <div class="page-grid">
            <div class="glass page-hero-card">
              <span class="kicker">Monte Carlo</span>
              <h1 class="page-title">Replace one optimistic date with a defensible range.</h1>
              <p class="page-copy">
                Run repeated simulations to estimate completion distributions, surface safer commitments,
                and identify how deadline risk changes under different scenarios.
              </p>
            </div>
            <div class="glass page-hero-card">
              <span class="eyebrow">Forecast anchors</span>
              <div class="data-pair"><span>P50 commitment</span><strong>Balanced target</strong></div>
              <div class="data-pair"><span>P80 commitment</span><strong>Safer promise</strong></div>
              <div class="data-pair"><span>Delay probability</span><strong>Leadership signal</strong></div>
              <div class="data-pair"><span>Main driver</span><strong>Mitigation focus</strong></div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _render_monte_carlo_lens() -> None:
    _section_intro(
        "Simulation lens",
        "Monte Carlo is the credibility engine behind the dashboard",
        "It transforms schedule uncertainty into metrics the audience can understand: mean completion, P50, P80, and exposure to the chosen deadline.",
    )

    st.markdown(
        """
        <div class="feature-grid">
          <div class="feature-card"><h3>P50</h3><p>The most likely commitment point when normal delivery conditions hold.</p></div>
          <div class="feature-card"><h3>P80</h3><p>A safer promise for executive communication when teams want higher certainty.</p></div>
          <div class="feature-card"><h3>Delay drivers</h3><p>Critical-path frequency reveals which tasks deserve mitigation before launch.</p></div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _render_floating_shortcut(active_page: str) -> None:
    if _is_dashboard_page(active_page):
        return

    st.markdown(
        """
        <div class="floating-shortcut-wrap">
          <button type="button" class="mini-pill sidebar-shortcut-pill sidebar-shortcut-trigger" data-command-trigger="true" aria-label="Open quick command search"><span class="sidebar-shortcut-key">⌘K / Ctrl+K</span></button>
          <div class="sidebar-mini-label sidebar-shortcut-caption">Quick command search</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _render_footer(active_page: str) -> None:
    globe_icon_uri = _svg_data_uri(
        """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24" fill="none" stroke="#e8f1ff" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round">
          <circle cx="12" cy="12" r="9"/>
          <path d="M3 12h18"/>
          <path d="M12 3a14.5 14.5 0 0 1 0 18"/>
          <path d="M12 3a14.5 14.5 0 0 0 0 18"/>
        </svg>
        """
    )
    flag_icon_uri = _svg_data_uri(
        """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 19 13">
          <rect width="19" height="13" fill="#B22234"/>
          <rect y="1" width="19" height="1" fill="#ffffff"/>
          <rect y="3" width="19" height="1" fill="#ffffff"/>
          <rect y="5" width="19" height="1" fill="#ffffff"/>
          <rect y="7" width="19" height="1" fill="#ffffff"/>
          <rect y="9" width="19" height="1" fill="#ffffff"/>
          <rect y="11" width="19" height="1" fill="#ffffff"/>
          <rect width="8.2" height="7" fill="#3C3B6E"/>
          <circle cx="1.3" cy="1.3" r="0.35" fill="#ffffff"/>
          <circle cx="3" cy="1.3" r="0.35" fill="#ffffff"/>
          <circle cx="4.7" cy="1.3" r="0.35" fill="#ffffff"/>
          <circle cx="6.4" cy="1.3" r="0.35" fill="#ffffff"/>
          <circle cx="2.15" cy="2.45" r="0.35" fill="#ffffff"/>
          <circle cx="3.85" cy="2.45" r="0.35" fill="#ffffff"/>
          <circle cx="5.55" cy="2.45" r="0.35" fill="#ffffff"/>
          <circle cx="1.3" cy="3.6" r="0.35" fill="#ffffff"/>
          <circle cx="3" cy="3.6" r="0.35" fill="#ffffff"/>
          <circle cx="4.7" cy="3.6" r="0.35" fill="#ffffff"/>
          <circle cx="6.4" cy="3.6" r="0.35" fill="#ffffff"/>
          <circle cx="2.15" cy="4.75" r="0.35" fill="#ffffff"/>
          <circle cx="3.85" cy="4.75" r="0.35" fill="#ffffff"/>
          <circle cx="5.55" cy="4.75" r="0.35" fill="#ffffff"/>
        </svg>
        """
    )
    farsi_flag_icon_uri = _svg_data_uri(
        """
        <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 19 13">
          <rect width="19" height="13" fill="#da1f26"/>
          <rect width="19" height="8.6667" fill="#ffffff"/>
          <rect width="19" height="4.3333" fill="#239f40"/>
          <g fill="none" stroke="#d8a828" stroke-width="0.22" stroke-linecap="round" stroke-linejoin="round">
            <circle cx="9.5" cy="6.55" r="1.15" fill="#e3b63a" stroke="#c99512"/>
            <path d="M9.5 4.6v-1.1M8.85 4.85l-.35-.85M10.15 4.85l.35-.85M8.25 5.3l-.8-.35M10.75 5.3l.8-.35M8.2 6.15H7.2M10.8 6.15h1M8.25 7l-.8.35M10.75 7l.8.35"/>
            <path d="M7.95 8.35c.1-.6.55-.95 1.2-.95h1.25c.8 0 1.15.55 1.18 1.05"/>
            <path d="M8.2 8.25l-.22.7M10.9 8.25l.28.68M8.9 7.4c-.22-.52-.18-1.05.12-1.45M10.05 7.32c.48-.18.78-.55.95-.98"/>
            <path d="M6.65 8.95h5.7"/>
            <path d="M8.15 5.15l.45-.75h1.8l.45.75"/>
            <path d="M8.95 4.15l.55-.45.55.45"/>
            <path d="M9.5 3.1l.28-.32.28.32"/>
            <path d="M7.05 5.2c-.6.65-.9 1.45-.88 2.3M11.95 5.2c.58.65.9 1.45.88 2.3"/>
          </g>
        </svg>
        """
    )
    flag_markup_by_code = {
        "en": f'<img src="{flag_icon_uri}" alt="" />',
        "fa": f'<img src="{farsi_flag_icon_uri}" alt="" />',
        "iw": "🇮🇱",
        "de": "🇩🇪",
        "it": "🇮🇹",
        "es": "🇪🇸",
        "fr": "🇫🇷",
        "pt": "🇧🇷",
        "zh-CN": "🇨🇳",
        "ja": "🇯🇵",
        "ar": "🇸🇦",
    }
    current_language_code = st.session_state.get("saas_language_code", "en")
    current_language_label = LANGUAGE_LABELS.get(current_language_code, "English")
    current_flag_markup = flag_markup_by_code.get(current_language_code, f'<img src="{flag_icon_uri}" alt="" />')
    language_toggle_id = f"footer-language-toggle-{re.sub(r'[^a-z0-9_-]+', '-', active_page.lower())}"
    language_options = [
        (
            code,
            LANGUAGE_LABELS.get(code, label),
            flag_markup_by_code.get(code, "🌐"),
            code == current_language_code,
        )
        for code, label, _ in LANGUAGE_OPTIONS
    ]
    language_menu_markup = "".join(
        (
            f'<a href="{_href(active_page, lang=code)}" class="footer-language-option'
            + (" selected" if selected else "")
            + '">'
            + f'<span class="footer-language-option-flag">{flag_markup}</span>'
            + f"<span>{label}</span>"
            + "</a>"
        )
        for code, label, flag_markup, selected in language_options
    )
    link_columns_markup = []
    for title, items in FOOTER_LINK_COLUMNS:
        item_markup = []
        for label, page, note in items:
            note_markup = f' <span class="footer-link-note">({note})</span>' if note else ""
            if page:
                link_class = "footer-link-item footer-legal-link" if title == "Legal" else "footer-link-item"
                item_markup.append(f'<a class="{link_class}" href="{_href(page)}">{label}{note_markup}</a>')
            else:
                item_markup.append(f'<span class="footer-static-link">{label}{note_markup}</span>')
        link_columns_markup.append(
            "<div class=\"footer-link-column\">"
            f"<h4>{title}</h4>"
            + "".join(item_markup)
            + "</div>"
        )

    social_markup = "".join(
        _footer_social_icon_markup(platform, label)
        for platform, label in FOOTER_SOCIALS
    )
    compliance_markup = "".join(
        (
            '<div class="footer-compliance-item">'
            f'<div class="footer-compliance-mark">{_footer_compliance_logo_markup(mark)}</div>'
            f'<div class="footer-compliance-code">{code}</div>'
            f'<div class="footer-compliance-region">{region}</div>'
            "</div>"
        )
        for mark, code, region in FOOTER_COMPLIANCE
    )
    footer_html = dedent(
        f"""
        <div class="footer-shell">
          <div class="footer-panel">
            <div class="footer-content">
              <div class="footer-top-grid">
                <div class="footer-brand-column">
                  {_footer_brand_markup()}
                  <p class="footer-brand-copy">
                    AI-powered project forecasting using Monte Carlo simulation and dependency analysis.
                    Predict delays before they happen.
                  </p>
                  <div class="footer-contact-stack">
                    <a class="footer-contact-row" href="mailto:info@certain-pm.com">
                      <span class="footer-contact-icon">✉</span>
                      <span>info@certain-pm.com</span>
                    </a>
                    <div class="footer-contact-row">
                      <span class="footer-contact-icon">⌖</span>
                      <span>Beverly Hills, CA</span>
                    </div>
                  </div>
                  <div class="footer-social-row">{social_markup}</div>
                </div>
                <div class="footer-link-grid">{"".join(link_columns_markup)}</div>
              </div>
              <div class="footer-mid-row">
                <div class="footer-newsletter">
                  <h4>Stay up to date</h4>
                  <p>Get the latest on AI-powered project management.</p>
                  <form class="footer-subscribe-row" method="get" action="">
                    <input type="hidden" name="page" value="subscription-confirmation" />
                    <input type="hidden" name="lang" value="{st.session_state.get("saas_language_code", "en")}" />
                    <input
                      class="footer-input-shell"
                      type="email"
                      name="newsletter_email"
                      placeholder="Enter your email"
                      autocomplete="email"
                    />
                    <button type="submit" class="btn btn-primary btn-outline-cta footer-subscribe-btn">Subscribe</button>
                  </form>
                </div>
                <div class="footer-compliance-row">{compliance_markup}</div>
              </div>
              <div class="footer-bottom-row">
                <span>&copy; 2026 {_brand_name_html()}. All rights reserved.</span>
                <div class="footer-language-shell">
                  <span class="footer-language-globe" aria-hidden="true">
                    <img src="{globe_icon_uri}" alt="" />
                  </span>
                  <div class="footer-language-dropdown">
                    <input id="{language_toggle_id}" class="footer-language-toggle" type="checkbox" />
                    <label class="footer-language-pill" for="{language_toggle_id}">
                      <span class="footer-language-flag" aria-hidden="true">
                        {current_flag_markup}
                      </span>
                      <span>{current_language_label}</span>
                      <span class="footer-language-caret" aria-hidden="true">
                        <svg viewBox="0 0 16 16" fill="none" xmlns="http://www.w3.org/2000/svg">
                          <path d="M4 6.5L8 10L12 6.5" stroke-width="1.8" stroke-linecap="round" stroke-linejoin="round"/>
                        </svg>
                      </span>
                    </label>
                    <div class="footer-language-menu">
                      {language_menu_markup}
                    </div>
                  </div>
                </div>
              </div>
            </div>
          </div>
        </div>
        """
    ).strip()

    st.markdown(footer_html, unsafe_allow_html=True)


def _section_intro(label: str, title: str, body: str) -> None:
    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="section-label">{label}</div>
              <h2 class="section-heading">{title}</h2>
              <p class="section-copy">{body}</p>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _risk_level(delay_probability: float) -> str:
    if delay_probability < 0.25:
        return "low"
    if delay_probability < 0.55:
        return "medium"
    return "high"


def _confidence_score(metrics: dict[str, float]) -> float:
    spread_penalty = max(0.0, metrics["p80"] - metrics["p50"]) * 0.42
    delay_penalty = metrics["delay_probability"] * 48.0
    confidence = 96.0 - spread_penalty - delay_penalty
    return round(max(58.0, min(96.0, confidence)), 1)


def _narrative(metrics: dict[str, float], top_driver: str | None) -> str:
    risk = _risk_level(metrics["delay_probability"]).upper()
    driver_text = f" Primary pressure point: {top_driver}." if top_driver else ""
    return (
        f"Forecast posture: {risk}. The model estimates a {metrics['delay_probability'] * 100:.1f}% chance "
        f"of missing the target date. A safer stakeholder commitment is {metrics['p80']:.1f} days.{driver_text}"
    )


def _theme_chart(figure: Any) -> Any:
    figure.update_layout(
        template="plotly",
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(5,12,24,0.58)",
        font=dict(color="#f4f8ff", family="Manrope, sans-serif", size=13),
        title=dict(x=0.02, xanchor="left", font=dict(color="#f4f8ff", family="Sora, sans-serif", size=22)),
        margin=dict(l=30, r=24, t=76, b=34),
        colorway=["#9b5cff", "#64a2ff", "#18d1c7", "#f5c56b", "#ff7f8c"],
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="left",
            x=0,
            bgcolor="rgba(0,0,0,0)",
            font=dict(color="#9eb0cb"),
        ),
        hoverlabel=dict(
            bgcolor="rgba(5,12,24,0.96)",
            bordercolor="rgba(100,162,255,0.32)",
            font=dict(color="#f4f8ff", family="Manrope, sans-serif"),
        ),
        coloraxis_colorbar=dict(
            outlinewidth=0,
            tickfont=dict(color="#9eb0cb"),
            title=dict(font=dict(color="#f4f8ff")),
        ),
    )
    figure.update_xaxes(
        gridcolor="rgba(155,92,255,0.08)",
        zerolinecolor="rgba(255,255,255,0.08)",
        linecolor="rgba(255,255,255,0.12)",
        showline=True,
        ticks="outside",
        tickfont=dict(color="#b9c9df"),
        title_font=dict(color="#dce9ff"),
    )
    figure.update_yaxes(
        gridcolor="rgba(100,162,255,0.08)",
        zerolinecolor="rgba(255,255,255,0.08)",
        linecolor="rgba(255,255,255,0.12)",
        showline=True,
        ticks="outside",
        tickfont=dict(color="#b9c9df"),
        title_font=dict(color="#dce9ff"),
    )
    for trace in figure.data:
        trace_type = getattr(trace, "type", "")
        if trace_type in {"scatter", "scattergl"}:
            line_width = getattr(getattr(trace, "line", None), "width", None) or 0
            marker_size = getattr(getattr(trace, "marker", None), "size", None) or 7
            trace.update(
                line=dict(width=max(float(line_width), 3.0)),
                marker=dict(size=marker_size, line=dict(width=0)),
            )
        elif trace_type in {"bar", "histogram"}:
            marker = getattr(trace, "marker", None)
            marker_line = getattr(marker, "line", None)
            existing_line_color = getattr(marker_line, "color", None)
            existing_line_width = getattr(marker_line, "width", None)
            trace.update(
                opacity=0.92,
                marker_line_color=existing_line_color if existing_line_color is not None else "rgba(255,255,255,0.14)",
                marker_line_width=existing_line_width if existing_line_width is not None else 1.2,
            )
        elif trace_type == "indicator":
            trace.update(
                number_font=dict(color="#f4f8ff", family="Sora, sans-serif"),
                title_font=dict(color="#9eb0cb"),
            )
    return figure


def _init_state() -> None:
    defaults = {
        "saas_selected_sample": DEFAULT_SAMPLE,
        "saas_project_text": SAMPLE_SCENARIOS[DEFAULT_SAMPLE]["description"],
        "saas_mode": "mock",
        "saas_deadline": SAMPLE_SCENARIOS[DEFAULT_SAMPLE]["deadline"],
        "saas_iterations": SAMPLE_SCENARIOS[DEFAULT_SAMPLE]["iterations"],
        "saas_max_tasks": SAMPLE_SCENARIOS[DEFAULT_SAMPLE]["max_tasks"],
        "saas_seed": 17,
        "saas_payload": None,
        "saas_history": [],
        "saas_user_email": "team@certain.ai",
        "saas_bootstrapped": False,
        "saas_ai_sensitivity": 0.65,
        "saas_confidence_threshold": 0.7,
        "saas_email_alerts": True,
        "saas_in_app_alerts": True,
        "saas_ai_alerts": True,
        "saas_language_code": "en",
        "saas_language": "English",
        "saas_region": "Global",
        "contact_email": "",
        "registration_email": "",
        "newsletter_last_request_key": "",
        "newsletter_delivery_status": "",
        "newsletter_delivery_message": "",
        "newsletter_confirmed": False,
        "registration_last_request_key": "",
        "registration_delivery_status": "",
        "registration_delivery_message": "",
        "task_architect_goal": "",
        "task_architect_timeline": 12,
        "task_architect_team_size": 5,
        "task_architect_payload": None,
        "decision_hub_last_analysis": "",
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


@st.cache_resource
def _load_risk_model_bundle() -> tuple[Any | None, RiskModelStatus]:
    model_path = settings.risk_model_path_file
    if not model_path.is_absolute():
        model_path = ROOT / model_path

    return load_risk_model(
        enabled=settings.risk_model_enabled,
        model_path=model_path,
        model_version=settings.risk_model_version,
    )


@st.cache_resource
def _initialize_database() -> bool:
    init_db()
    return True


@st.cache_data
def _load_model_metrics() -> dict[str, Any]:
    metrics_path = settings.risk_model_metrics_file
    if not metrics_path.is_absolute():
        metrics_path = ROOT / metrics_path
    if not metrics_path.exists():
        return {}
    return json.loads(metrics_path.read_text(encoding="utf-8"))


@st.cache_data
def _load_portfolio_history() -> pd.DataFrame:
    return pd.read_csv(ROOT / "data" / "project_portfolio_history.csv")


@st.cache_data
def _load_monitoring_snapshot() -> pd.DataFrame:
    return pd.read_csv(ROOT / "data" / "risk_monitoring_snapshot.csv")


def _portfolio_summary(portfolio_df: pd.DataFrame) -> dict[str, float]:
    delayed = (portfolio_df["Delay_Flag"] == "Delayed").mean()
    return {
        "projects": float(len(portfolio_df)),
        "delayed_pct": float(delayed * 100.0),
        "avg_budget_overrun": float(portfolio_df["Budget_Overrun_Pct"].mean()),
        "avg_forecast_error": float(portfolio_df["Forecast_Error_Days"].mean()),
        "avg_complexity": float(portfolio_df["Complexity_Score"].mean()),
    }


def _monitoring_summary(monitoring_df: pd.DataFrame) -> dict[str, float]:
    accuracy = (monitoring_df["Predicted_Risk_Level"] == monitoring_df["Actual_Risk_Level"]).mean()
    return {
        "accuracy_pct": float(accuracy * 100.0),
        "avg_confidence": float(monitoring_df["Prediction_Confidence"].mean() * 100.0),
        "avg_drift": float(monitoring_df["Drift_Score"].mean()),
        "alerts_open": float(monitoring_df["Alerts_Open"].sum()),
        "avg_abs_error": float(monitoring_df["Absolute_Forecast_Error_Days"].mean()),
    }


def _safe_metric(value: float, digits: int = 1) -> str:
    return f"{value:.{digits}f}"


def _persist_run(payload: dict[str, Any]) -> None:
    try:
        _initialize_database()
        save_session_run(
            project_text=payload["project_text"],
            mode=payload["mode"],
            params={
                "sample": payload["sample"],
                "max_tasks": payload["max_tasks"],
                "portfolio_projects": int(payload["portfolio_summary"]["projects"]),
            },
            tasks=payload["tasks"],
            deadline_days=payload["deadline_days"],
            iterations=payload["iterations"],
            seed=payload["seed"],
            metrics=payload["metrics"],
            completion_times=payload["completion"],
            scenarios_df=payload["scenarios_df"],
            ml_features_df=payload["ml_features_df"],
            ml_predictions_df=payload["ml_predictions_df"],
            model_version=str(payload["model_status"].get("model_version", settings.risk_model_version)),
        )
    except Exception:
        # Keep the presentation flow resilient even if SQLite is unavailable.
        return


def _load_selected_sample() -> None:
    sample = SAMPLE_SCENARIOS.get(st.session_state.saas_selected_sample)
    if not sample:
        return
    st.session_state.saas_project_text = sample["description"]
    st.session_state.saas_deadline = sample["deadline"]
    st.session_state.saas_iterations = sample["iterations"]
    st.session_state.saas_max_tasks = sample["max_tasks"]


def _build_tasks() -> list[dict[str, Any]]:
    selected_sample = st.session_state.saas_selected_sample
    if st.session_state.saas_mode == "mock" and selected_sample in SAMPLE_SCENARIOS:
        return copy.deepcopy(SAMPLE_SCENARIOS[selected_sample]["tasks"])

    plan = generate_task_plan(
        st.session_state.saas_project_text,
        max_tasks=int(st.session_state.saas_max_tasks),
        mode=st.session_state.saas_mode,
    )
    return [task.model_dump() for task in plan.tasks]


def _parse_timeline_input(timeline_text: str | int | float) -> tuple[int, float]:
    if isinstance(timeline_text, (int, float)):
        weeks = max(1, round(float(timeline_text)))
        return int(weeks), float(max(15, weeks * 5))

    clean = timeline_text.strip().lower()
    if not clean:
        raise TaskGenerationError("Timeline is required. Try something like `12 weeks`.")

    match = re.search(r"(\d+(?:\.\d+)?)", clean)
    if not match:
        raise TaskGenerationError("Timeline must include a number, for example `12 weeks`.")

    value = float(match.group(1))
    if "quarter" in clean:
        weeks = max(1, round(value * 13))
    elif "month" in clean:
        weeks = max(1, round(value * 4))
    elif "day" in clean:
        weeks = max(1, round(value / 5))
    else:
        weeks = max(1, round(value))

    return int(weeks), float(max(15, weeks * 5))


def _build_task_architect_tasks(project_goal: str, timeline_weeks: int, team_size: int) -> list[dict[str, Any]]:
    base_plan = generate_task_plan("Create a delivery plan", max_tasks=8, mode="mock")
    base_tasks = [task.model_dump() for task in base_plan.tasks]
    goal_label = project_goal.strip().rstrip(".")
    short_goal = goal_label if len(goal_label) <= 42 else f"{goal_label[:39].rstrip()}..."

    custom_names = [
        f"Scope and success criteria for {short_goal}",
        "Roadmap, milestones, and workstream sequencing",
        "Architecture and workflow design",
        f"Core build for {short_goal}",
        "Integrations and dependency hardening",
        "Validation, QA, and pilot readiness",
        "Documentation and team enablement",
        "Launch readiness and stakeholder handoff",
    ]

    target_days = max(20.0, timeline_weeks * 5.0)
    base_total = sum(float(task["mean_duration"]) for task in base_tasks)
    schedule_scale = max(0.6, min(1.45, target_days / base_total))
    team_factor = max(0.8, min(1.16, 1.08 - ((min(team_size, 12) - 5) * 0.035)))
    pressure_delta = max(0.0, (1.0 - schedule_scale) * 0.16)
    capacity_delta = max(0.0, (5 - min(team_size, 5)) * 0.025)

    tasks: list[dict[str, Any]] = []
    for index, task in enumerate(base_tasks):
        mean_duration = max(2, round(float(task["mean_duration"]) * schedule_scale * team_factor))
        std_dev = round(max(1.0, mean_duration * (0.17 + float(task["risk_factor"]) * 0.18)), 1)
        risk_factor = min(
            0.82,
            max(0.18, round(float(task["risk_factor"]) + pressure_delta + capacity_delta, 2)),
        )
        tasks.append(
            {
                **task,
                "name": custom_names[index] if index < len(custom_names) else task["name"],
                "mean_duration": mean_duration,
                "std_dev": std_dev,
                "risk_factor": risk_factor,
            }
        )
    return tasks


def _generate_task_architect_payload(project_goal: str, timeline_text: str | int | float, team_size: int) -> dict[str, Any]:
    goal = project_goal.strip()
    if not goal:
        raise TaskGenerationError("Project goal is required to generate an AI plan.")

    timeline_weeks, deadline_days = _parse_timeline_input(timeline_text)
    timeline_label = (
        f"{int(round(float(timeline_text)))} weeks"
        if isinstance(timeline_text, (int, float))
        else timeline_text.strip()
    )
    normalized_team_size = max(1, int(team_size))
    tasks = _build_task_architect_tasks(goal, timeline_weeks, normalized_team_size)
    iterations = 2000
    seed = int(st.session_state.get("saas_seed", 17))
    graph = build_project_graph(tasks)
    critical_path, critical_days = critical_path_by_mean(graph)
    completion = run_monte_carlo(graph, iterations=iterations, seed=seed)
    metrics = compute_metrics(completion, deadline_days)
    drivers = rank_delay_drivers(graph, iterations=min(500, iterations), seed=seed)
    scenarios_df = scenario_comparison(
        build_project_graph,
        tasks,
        deadline_days=deadline_days,
        iterations=iterations,
        seed=seed,
    )
    ml_features_df, ml_predictions_df, ml_summary_df, model_status = _run_ml_scoring(tasks)
    portfolio_df = _load_portfolio_history()
    monitoring_df = _load_monitoring_snapshot()
    model_metrics = _load_model_metrics()
    portfolio_summary = _portfolio_summary(portfolio_df)
    monitoring_summary = _monitoring_summary(monitoring_df)

    payload = {
        "project_text": goal,
        "sample": "AI Task Architect",
        "mode": "mock",
        "deadline_days": deadline_days,
        "iterations": iterations,
        "max_tasks": len(tasks),
        "seed": seed,
        "tasks": tasks,
        "graph": graph,
        "critical_path": critical_path,
        "critical_path_days": critical_days,
        "completion": completion,
        "metrics": metrics,
        "drivers": drivers,
        "scenarios_df": scenarios_df,
        "ml_features_df": ml_features_df,
        "ml_predictions_df": ml_predictions_df,
        "ml_summary_df": ml_summary_df,
        "model_status": asdict(model_status),
        "portfolio_df": portfolio_df,
        "monitoring_df": monitoring_df,
        "model_metrics": model_metrics,
        "portfolio_summary": portfolio_summary,
        "monitoring_summary": monitoring_summary,
        "architect_meta": {
            "project_goal": goal,
            "timeline_text": timeline_label,
            "timeline_weeks": timeline_weeks,
            "team_size": normalized_team_size,
        },
    }
    return payload


def _run_task_architect_generation(project_goal: str, timeline_text: str | int | float, team_size: int) -> str | None:
    try:
        generated_payload = _generate_task_architect_payload(project_goal, timeline_text, int(team_size))
        st.session_state.task_architect_payload = generated_payload
        st.session_state.saas_payload = generated_payload
        st.session_state.saas_bootstrapped = True
        st.session_state.saas_project_text = generated_payload["project_text"]
        st.session_state.saas_deadline = generated_payload["deadline_days"]
        st.session_state.saas_iterations = generated_payload["iterations"]
        st.session_state.saas_max_tasks = generated_payload["max_tasks"]
        _persist_run(generated_payload)

        top_driver = generated_payload["drivers"][0]["task_name"] if generated_payload["drivers"] else "No dominant driver"
        history_item = {
            "timestamp": datetime.now(tz=UTC).strftime("%Y-%m-%d %H:%M UTC"),
            "sample": generated_payload["sample"],
            "title": generated_payload["project_text"][:60] + ("..." if len(generated_payload["project_text"]) > 60 else ""),
            "delay_probability": generated_payload["metrics"]["delay_probability"],
            "p80": generated_payload["metrics"]["p80"],
            "top_driver": top_driver,
        }
        st.session_state.saas_history = [history_item, *st.session_state.saas_history[:5]]
        return None
    except (TaskGenerationError, GraphValidationError, ValueError) as exc:
        return str(exc)
    except Exception as exc:  # noqa: BLE001
        return f"Unable to generate the AI plan: {exc}"


def _run_task_architect_upload_generation(uploaded_file: Any) -> str | None:
    if uploaded_file is None:
        return "Upload a project file first."

    try:
        generated_payload = _generate_task_architect_payload_from_upload(uploaded_file.name, uploaded_file.getvalue())
        st.session_state.task_architect_payload = generated_payload
        st.session_state.saas_payload = generated_payload
        st.session_state.saas_bootstrapped = True
        st.session_state.saas_project_text = generated_payload["project_text"]
        st.session_state.saas_deadline = generated_payload["deadline_days"]
        st.session_state.saas_iterations = generated_payload["iterations"]
        st.session_state.saas_max_tasks = generated_payload["max_tasks"]
        _persist_run(generated_payload)

        top_driver = generated_payload["drivers"][0]["task_name"] if generated_payload["drivers"] else "No dominant driver"
        history_item = {
            "timestamp": datetime.now(tz=UTC).strftime("%Y-%m-%d %H:%M UTC"),
            "sample": generated_payload["sample"],
            "title": generated_payload["project_text"][:60] + ("..." if len(generated_payload["project_text"]) > 60 else ""),
            "delay_probability": generated_payload["metrics"]["delay_probability"],
            "p80": generated_payload["metrics"]["p80"],
            "top_driver": top_driver,
        }
        st.session_state.saas_history = [history_item, *st.session_state.saas_history[:5]]
        return None
    except (TaskGenerationError, GraphValidationError, ValueError) as exc:
        return str(exc)
    except Exception as exc:  # noqa: BLE001
        return f"Unable to generate the AI plan from the uploaded file: {exc}"


def _run_ml_scoring(
    tasks: list[dict[str, Any]],
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, RiskModelStatus]:
    model, status = _load_risk_model_bundle()
    features_df = build_default_feature_df(tasks)
    if model is None or not status.available:
        return (
            features_df,
            pd.DataFrame(),
            pd.DataFrame(columns=["Risk_Level", "Count"]),
            status,
        )

    validated_df, scored_df, summary_df = score_tasks(model, features_df)
    return validated_df, scored_df, summary_df, status


def _serialize_payload(payload: dict[str, Any]) -> str:
    safe_payload = {
        "project_text": payload["project_text"],
        "sample": payload["sample"],
        "mode": payload["mode"],
        "deadline_days": payload["deadline_days"],
        "iterations": payload["iterations"],
        "seed": payload["seed"],
        "metrics": payload["metrics"],
        "critical_path": payload["critical_path"],
        "drivers": payload["drivers"],
        "scenarios": payload["scenarios_df"].to_dict(orient="records"),
        "tasks": payload["tasks"],
        "ml_predictions": payload["ml_predictions_df"].to_dict(orient="records"),
        "model_status": payload["model_status"],
    }
    return json.dumps(safe_payload, indent=2)


def _run_forecast() -> None:
    project_text = st.session_state.saas_project_text.strip()
    if not project_text:
        st.error("Project description is required.")
        return

    try:
        with st.spinner("Designing the plan and running the forecast..."):
            tasks = _build_tasks()
            graph = build_project_graph(tasks)
            critical_path, critical_days = critical_path_by_mean(graph)
            completion = run_monte_carlo(
                graph,
                iterations=int(st.session_state.saas_iterations),
                seed=int(st.session_state.saas_seed),
            )
            metrics = compute_metrics(completion, float(st.session_state.saas_deadline))
            drivers = rank_delay_drivers(
                graph,
                iterations=min(500, int(st.session_state.saas_iterations)),
                seed=int(st.session_state.saas_seed),
            )
            scenarios_df = scenario_comparison(
                build_project_graph,
                tasks,
                deadline_days=float(st.session_state.saas_deadline),
                iterations=int(st.session_state.saas_iterations),
                seed=int(st.session_state.saas_seed),
            )
            ml_features_df, ml_predictions_df, ml_summary_df, model_status = _run_ml_scoring(tasks)
            portfolio_df = _load_portfolio_history()
            monitoring_df = _load_monitoring_snapshot()
            model_metrics = _load_model_metrics()
            portfolio_summary = _portfolio_summary(portfolio_df)
            monitoring_summary = _monitoring_summary(monitoring_df)

        payload = {
            "project_text": project_text,
            "sample": st.session_state.saas_selected_sample,
            "mode": st.session_state.saas_mode,
            "deadline_days": float(st.session_state.saas_deadline),
            "iterations": int(st.session_state.saas_iterations),
            "max_tasks": int(st.session_state.saas_max_tasks),
            "seed": int(st.session_state.saas_seed),
            "tasks": tasks,
            "graph": graph,
            "critical_path": critical_path,
            "critical_path_days": critical_days,
            "completion": completion,
            "metrics": metrics,
            "drivers": drivers,
            "scenarios_df": scenarios_df,
            "ml_features_df": ml_features_df,
            "ml_predictions_df": ml_predictions_df,
            "ml_summary_df": ml_summary_df,
            "model_status": asdict(model_status),
            "portfolio_df": portfolio_df,
            "monitoring_df": monitoring_df,
            "model_metrics": model_metrics,
            "portfolio_summary": portfolio_summary,
            "monitoring_summary": monitoring_summary,
        }
        st.session_state.saas_payload = payload
        _persist_run(payload)

        top_driver = drivers[0]["task_name"] if drivers else "No dominant driver"
        history_item = {
            "timestamp": datetime.now(tz=UTC).strftime("%Y-%m-%d %H:%M UTC"),
            "sample": payload["sample"],
            "title": project_text[:60] + ("..." if len(project_text) > 60 else ""),
            "delay_probability": metrics["delay_probability"],
            "p80": metrics["p80"],
            "top_driver": top_driver,
        }
        st.session_state.saas_history = [history_item, *st.session_state.saas_history[:5]]
    except TaskGenerationError as exc:
        st.error(str(exc))
    except GraphValidationError as exc:
        st.error(f"Generated task graph is invalid: {exc}")
    except Exception as exc:  # noqa: BLE001
        st.error(f"Unable to generate the forecast workspace: {exc}")


def _ensure_demo_payload() -> None:
    if st.session_state.saas_bootstrapped or st.session_state.saas_payload is not None:
        return
    _run_forecast()
    st.session_state.saas_bootstrapped = True


def _current_payload() -> dict[str, Any]:
    _ensure_demo_payload()
    return st.session_state.saas_payload


def _history_markup() -> str:
    if not st.session_state.saas_history:
        return "<div class='history-card'><p>No workspaces yet. Launch one from the console.</p></div>"

    items = "".join(
        f"""
        <div class="history-item">
          <strong>{item['sample']}</strong>
          <span>{item['timestamp']}</span>
          <span>{item['delay_probability'] * 100:.1f}% delay risk · P80 {item['p80']:.1f}d</span>
          <span>Primary driver: {item['top_driver']}</span>
        </div>
        """
        for item in st.session_state.saas_history
    )
    return f"<div class='history-list'>{items}</div>"


def _render_home() -> None:
    logo_uri = _logo_data_uri()

    ticker_markup = "".join(f"<span class='ticker-item'>{item}</span>" for item in HOME_TICKER_ITEMS)

    st.markdown(
        dedent(
            f"""
            <div class="hero-shell">
              <div class="home-hero-stage">
                <div class="glass ticker-shell home-hero-ticker">
                  <div class="ticker-mask">
                    <div class="ticker-track">{ticker_markup}{ticker_markup}</div>
                  </div>
                </div>
                <div class="home-logo-stack">
                  <div class="home-logo-frame">
                    {f'<img src="{logo_uri}" class="home-logo-image" alt="certAIn logo" />' if logo_uri else _footer_brand_markup()}
                  </div>
                  <div class="home-powered-pill">
                    <span class="home-powered-glyph">✦</span>
                    <span>Powered by Advanced AI</span>
                    <span class="home-powered-dot"></span>
                  </div>
                </div>
                <h1 class="home-hero-title">
                  <span>Predict Project Delays</span>
                  <span class="home-hero-emphasis">Before They Happen</span>
                </h1>
                <p class="home-hero-copy">
                  AI-powered forecasting using Monte Carlo simulation and dependency analysis.
                </p>
                <div class="home-live-pill">
                  <span class="home-live-dot"></span>
                  <span>AI analyzing 2,847 task dependencies...</span>
                </div>
                <div class="home-hero-actions">
                  <a href="{_href("product")}" class="btn btn-primary btn-outline-cta">View Product Demo</a>
                </div>
                <div class="home-preview-shell">
                  <div class="home-preview-grid">
                    <div class="home-preview-card">
                      <div class="home-preview-head">
                        <span class="home-preview-icon">◫</span>
                        <span>Completion Probability</span>
                      </div>
                      <div class="home-chart-wrap">
                        <div class="home-chart-bars">
                          <span class="home-chart-bar" style="height:18px;"></span>
                          <span class="home-chart-bar" style="height:34px;"></span>
                          <span class="home-chart-bar" style="height:54px;"></span>
                          <span class="home-chart-bar" style="height:78px;"></span>
                          <span class="home-chart-bar" style="height:96px;"></span>
                          <span class="home-chart-bar" style="height:102px;"></span>
                          <span class="home-chart-bar" style="height:90px;"></span>
                          <span class="home-chart-bar" style="height:60px;"></span>
                          <span class="home-chart-bar" style="height:34px;"></span>
                          <span class="home-chart-bar" style="height:14px;"></span>
                        </div>
                        <div class="home-chart-axis"></div>
                        <div class="home-chart-caption">P50: May 15 · P80: May 22</div>
                      </div>
                    </div>
                    <div class="home-preview-card">
                      <div class="home-preview-head">
                        <span class="home-preview-icon">⬡</span>
                        <span>Risk Alerts</span>
                      </div>
                      <div class="home-alert-stack">
                        <div class="home-alert-item high">
                          <span class="home-alert-dot"></span>
                          <span class="home-alert-copy">API Integration at risk — 78% delay probability</span>
                        </div>
                        <div class="home-alert-item medium">
                          <span class="home-alert-dot"></span>
                          <span class="home-alert-copy">Database migration — resource bottleneck</span>
                        </div>
                        <div class="home-alert-item low">
                          <span class="home-alert-dot"></span>
                          <span class="home-alert-copy">Frontend sprint on track — 95% confidence</span>
                        </div>
                      </div>
                    </div>
                    <div class="home-preview-card">
                      <div class="home-preview-head">
                        <span class="home-preview-icon">↗</span>
                        <span>Project Timeline</span>
                      </div>
                      <div class="home-timeline-stack">
                        <div class="home-timeline-row">
                          <div class="home-timeline-top"><span>Design Phase</span><span>100%</span></div>
                          <div class="home-timeline-bar"><span style="width:100%;"></span></div>
                        </div>
                        <div class="home-timeline-row">
                          <div class="home-timeline-top"><span>Development</span><span>72%</span></div>
                          <div class="home-timeline-bar"><span style="width:72%;"></span></div>
                        </div>
                        <div class="home-timeline-row">
                          <div class="home-timeline-top"><span>Testing</span><span>30%</span></div>
                          <div class="home-timeline-bar"><span style="width:30%; background:rgba(255,255,255,0.26);"></span></div>
                        </div>
                        <div class="home-timeline-row">
                          <div class="home-timeline-top"><span>Launch</span><span>0%</span></div>
                          <div class="home-timeline-bar"><span style="width:0%;"></span></div>
                        </div>
                      </div>
                    </div>
                  </div>
                </div>
                <div class="home-proof-strip">
                  <div class="home-proof-item">
                    <div class="home-proof-value">10000+</div>
                    <div class="home-proof-label">Monte Carlo Simulations</div>
                    <div class="home-proof-copy">Per project forecast</div>
                  </div>
                  <div class="home-proof-item">
                    <div class="home-proof-value">98.7%</div>
                    <div class="home-proof-label">Forecast Accuracy</div>
                    <div class="home-proof-copy">Validated across 500+ projects</div>
                  </div>
                  <div class="home-proof-item">
                    <div class="home-proof-value">47%</div>
                    <div class="home-proof-label">Risk Reduction</div>
                    <div class="home-proof-copy">Average across all clients</div>
                  </div>
                  <div class="home-proof-item">
                    <div class="home-proof-value">3x</div>
                    <div class="home-proof-label">Faster Decisions</div>
                    <div class="home-proof-copy">Compared to traditional PM</div>
                  </div>
                </div>
                <div class="home-intelligence-shell">
                  <div class="home-intelligence-head">
                    <div class="home-intelligence-pill">
                      <span class="home-intelligence-pill-glyph">⌘</span>
                      <span>AI-Powered Features</span>
                    </div>
                    <h2 class="home-intelligence-title">Intelligence at Every Level</h2>
                    <p class="home-intelligence-copy">
                      From task planning to executive reporting — {_brand_name_html()} covers the entire project lifecycle.
                    </p>
                  </div>
                  <div class="home-intelligence-grid">
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">⌘</div>
                      <h3>AI Project Planning</h3>
                      <p>Automatically structure projects with intelligent task decomposition and dependency mapping.</p>
                    </div>
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">∿</div>
                      <h3>Monte Carlo Forecasting</h3>
                      <p>Run thousands of simulations to predict completion dates with statistical confidence.</p>
                    </div>
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">◈</div>
                      <h3>Critical Path Risk Detection</h3>
                      <p>Identify bottlenecks and vulnerable dependencies before they cause delays.</p>
                    </div>
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">↗</div>
                      <h3>Explainable AI Insights</h3>
                      <p>Understand why delays happen with transparent AI reasoning and recommendations.</p>
                    </div>
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">▤</div>
                      <h3>Executive Reporting</h3>
                      <p>Generate stakeholder-ready reports with health scores and mitigation plans.</p>
                    </div>
                    <div class="home-intelligence-card">
                      <div class="home-intelligence-icon-shell">◫</div>
                      <h3>What-If Scenario Simulation</h3>
                      <p>Model budget, timeline, and resource changes to see their impact before committing decisions.</p>
                    </div>
                  </div>
                </div>
              </div>
            </div>
            """,
        ),
        unsafe_allow_html=True,
    )

    _section_intro(
        "Proof points",
        "Enterprise SaaS framing with metrics that feel presentation-ready",
        "These signals are designed to make the product story immediately legible before you move into the live dashboard.",
    )
    st.markdown(
        f"""
        <div class="home-testimonial-shell">
          <div class="home-testimonial-grid">
            <div class="home-testimonial-card">
              <p class="home-testimonial-quote">"{_brand_name_html()} reduced our project overruns by 40%. The Monte Carlo simulations gave us confidence we never had before."</p>
              <div class="home-testimonial-author">
                <strong>Sarah Chen</strong>
                <span>VP Engineering, TechCorp</span>
              </div>
            </div>
            <div class="home-testimonial-card">
              <p class="home-testimonial-quote">"The AI risk detection caught a critical dependency issue that would have delayed our launch by 3 months."</p>
              <div class="home-testimonial-author">
                <strong>Marcus Weber</strong>
                <span>Program Director, Siemens Digital</span>
              </div>
            </div>
            <div class="home-testimonial-card">
              <p class="home-testimonial-quote">"Finally, a tool that speaks the language of project managers. The executive reports are exactly what our stakeholders need."</p>
              <div class="home-testimonial-author">
                <strong>Elena Rodriguez</strong>
                <span>PMO Lead, Consulting Firm</span>
              </div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(
        dedent(
            f"""
            <div class="home-cta-shell">
              <div class="home-cta-pill">
                <span class="home-cta-pill-glyph">✦</span>
                <span>Start in under 2 minutes</span>
              </div>
              <h2 class="home-cta-title">Ready to predict your project's future?</h2>
              <p class="home-cta-copy">Start forecasting with AI in minutes. No credit card required.</p>
              <a href="{_trial_href()}" class="btn btn-primary btn-outline-cta">Start Free Trial</a>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_product() -> None:
    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="page-grid">
                <div class="glass page-hero-card">
                  <span class="kicker">Product</span>
                  <h1 class="page-title">Plan, forecast, and explain risk from one workflow.</h1>
                  <p class="page-copy">
                    {_brand_name_html()} connects project intake, dependency modeling, simulation, advisory ML, and scenario comparison
                    in one SaaS-style interface.
                  </p>
                  <div class="cta-row">
                    <a href="{_href("pricing")}" class="product-card-cta-link">View Pricing</a>
                  </div>
                </div>
                <div class="glass page-hero-card">
                  <span class="eyebrow">Core flow</span>
                  <div class="data-pair"><span>1. Project brief</span><strong>AI-assisted intake</strong></div>
                  <div class="data-pair"><span>2. Task graph</span><strong>Dependency reasoning</strong></div>
                  <div class="data-pair"><span>3. Simulation</span><strong>Monte Carlo</strong></div>
                  <div class="data-pair"><span>4. Scenarios</span><strong>Decision-ready</strong></div>
                  <div class="data-pair"><span>5. Executive brief</span><strong>Presentation-safe</strong></div>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )

    _section_intro(
        "Modules",
        "The platform is structured like a SaaS product, not just a notebook demo",
        "Each part of the interface tells a different part of the value story while staying anchored in the same underlying logic.",
    )

    st.markdown(
        dedent(
            f"""
            <div class="detail-grid">
              <div class="detail-card">
                <h3>Brief to plan</h3>
                <p>Generate structured project tasks with durations, dependencies, and risk factors.</p>
              </div>
              <div class="detail-card">
                <h3>Workflow graph</h3>
                <p>Highlight the critical path and expose where sequence risk becomes schedule risk.</p>
              </div>
              <div class="detail-card">
                <h3>Scenario lab</h3>
                <p>Compare baseline, mitigation, and acceleration paths before promising a deadline.</p>
              </div>
              <div class="detail-card">
                <h3>Advisory signals</h3>
                <p>Overlay task-level risk scoring without replacing the probabilistic schedule view.</p>
              </div>
              <div class="detail-card">
                <h3>Executive outputs</h3>
                <p>Present metrics, delay drivers, and safer commitments in business language.</p>
              </div>
              <div class="detail-card">
                <h3>Presentation handoff</h3>
                <p>Reuse the same product language in the prototype, Streamlit app, and final presentation.</p>
              </div>
            </div>
            <div class="product-card-cta">
              <a href="{_href("guide")}" class="product-card-cta-link">Try the Dashboard <span aria-hidden="true">→</span></a>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_forecasting() -> None:
    _ensure_demo_payload()
    st.markdown(
        dedent(
            f"""
            <div class="page-shell ai-forecast-shell">
              <div class="glass page-hero-card ai-forecast-kicker-card">
                <span class="kicker">AI Forecasting</span>
              </div>
              <div class="ai-forecast-stage">
                <div class="ai-forecast-copy">
                  <h1 class="ai-forecast-title">Predictive Intelligence for Projects</h1>
                  <p class="ai-forecast-subtitle">
                    Our AI engine analyzes project structures, dependencies, and historical data to predict risks
                    with unprecedented accuracy.
                  </p>
                </div>
                <div class="ai-forecast-grid">
                  <article class="detail-card ai-forecast-card">
                    <span class="ai-forecast-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M10 5a3 3 0 1 0-6 0c0 1.6 1.1 2.8 2.4 3.4"></path>
                        <path d="M14 5a3 3 0 1 1 6 0c0 1.6-1.1 2.8-2.4 3.4"></path>
                        <path d="M7 13a3 3 0 1 0 0 6"></path>
                        <path d="M17 13a3 3 0 1 1 0 6"></path>
                        <path d="M7 8h10"></path>
                        <path d="M12 8v8"></path>
                        <path d="M9.5 16H7"></path>
                        <path d="M17 16h-2.5"></path>
                      </svg>
                    </span>
                    <h3>Delay Prediction</h3>
                    <p>Machine learning models analyze historical patterns and current dependencies to forecast delays before they happen.</p>
                  </article>
                  <article class="detail-card ai-forecast-card">
                    <span class="ai-forecast-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M5 16l5-5 4 4 5-7"></path>
                        <path d="M15 8h4v4"></path>
                      </svg>
                    </span>
                    <h3>Velocity Forecasting</h3>
                    <p>Track team velocity trends and predict future sprint capacity with confidence intervals.</p>
                  </article>
                  <article class="detail-card ai-forecast-card">
                    <span class="ai-forecast-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M12 4l8 14H4L12 4z"></path>
                        <path d="M12 10v3.5"></path>
                        <path d="M12 17h.01"></path>
                      </svg>
                    </span>
                    <h3>Risk Scoring</h3>
                    <p>Every task receives an AI-generated risk score based on complexity, dependencies, and resource allocation.</p>
                  </article>
                  <article class="detail-card ai-forecast-card">
                    <span class="ai-forecast-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M5 19V9"></path>
                        <path d="M10 19V5"></path>
                        <path d="M15 19v-7"></path>
                        <path d="M20 19V8"></path>
                        <path d="M3 19h18"></path>
                      </svg>
                    </span>
                    <h3>Explainable Insights</h3>
                    <p>Understand why the AI predicts a delay with transparent reasoning and actionable recommendations.</p>
                  </article>
                </div>
                <div class="glass ai-forecast-process-card">
                  <h3>How AI Forecasting Works</h3>
                  <div class="ai-forecast-steps">
                    <div class="ai-forecast-step">
                      <span class="ai-forecast-step-dot">1</span>
                      <strong>Upload Project Plan</strong>
                    </div>
                    <div class="ai-forecast-step">
                      <span class="ai-forecast-step-dot">2</span>
                      <strong>AI Analyzes Dependencies</strong>
                    </div>
                    <div class="ai-forecast-step">
                      <span class="ai-forecast-step-dot">3</span>
                      <strong>Risk Scores Generated</strong>
                    </div>
                    <div class="ai-forecast-step">
                      <span class="ai-forecast-step-dot">4</span>
                      <strong>Actionable Insights</strong>
                    </div>
                  </div>
                </div>
                <div class="ai-forecast-cta-wrap">
                  <a href="?page=risk-intelligence&amp;lang=en" class="btn btn-primary ai-forecast-cta">
                    See Risk Intelligence <span aria-hidden="true">→</span>
                  </a>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )

def _render_monte_carlo() -> None:
    _render_monte_carlo_explainer()
    _render_monte_carlo_lens()


def _render_pricing() -> None:
    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="glass page-hero-card pricing-hero-card">
                <span class="kicker">Pricing</span>
                <h1 class="page-title">Simple pricing that scales from first forecast to enterprise rollout.</h1>
                <p class="page-copy">
                  Choose a plan that fits your team, from early project planning to full enterprise governance.
                </p>
              </div>
            </div>
            <div class="page-shell pricing-showcase-shell">
              <div class="pricing-grid pricing-showcase">
                <article class="pricing-card pricing-plan-card">
                  <div class="pricing-plan-head">
                    <h3 class="pricing-plan-name">Starter</h3>
                    <div class="pricing-price-row">
                      <span class="pricing-price-value">$29</span>
                      <span class="pricing-price-suffix">/month</span>
                    </div>
                    <p class="pricing-plan-copy">For small teams getting started with AI project planning.</p>
                  </div>
                  <ul class="pricing-feature-list">
                    <li>Up to 5 projects</li>
                    <li>1,000 Monte Carlo simulations/mo</li>
                    <li>Basic risk detection</li>
                    <li>CSV &amp; Excel upload</li>
                    <li>Email support</li>
                    <li>1 team member</li>
                  </ul>
                  <a href="{_trial_href()}" class="pricing-plan-cta pricing-plan-cta-muted">Start Free Trial</a>
                </article>
                <article class="pricing-card pricing-plan-card pricing-plan-card-featured">
                  <span class="pricing-plan-badge">Most Popular</span>
                  <div class="pricing-plan-head">
                    <h3 class="pricing-plan-name">Professional</h3>
                    <div class="pricing-price-row">
                      <span class="pricing-price-value">$99</span>
                      <span class="pricing-price-suffix">/month</span>
                    </div>
                    <p class="pricing-plan-copy">For growing teams that need advanced forecasting.</p>
                  </div>
                  <ul class="pricing-feature-list">
                    <li>Unlimited projects</li>
                    <li>50,000 simulations/mo</li>
                    <li>Advanced AI risk analysis</li>
                    <li>All file formats (incl. .mpp)</li>
                    <li>Executive reporting</li>
                    <li>Up to 25 team members</li>
                    <li>API access</li>
                    <li>Priority support</li>
                  </ul>
                  <a href="{_trial_href()}" class="pricing-plan-cta pricing-plan-cta-primary">Start Free Trial</a>
                </article>
                <article class="pricing-card pricing-plan-card">
                  <div class="pricing-plan-head">
                    <h3 class="pricing-plan-name">Enterprise</h3>
                    <div class="pricing-price-row pricing-price-row-custom">
                      <span class="pricing-price-value pricing-price-value-custom">Custom</span>
                    </div>
                    <p class="pricing-plan-copy">For organizations that need maximum control and scale.</p>
                  </div>
                  <ul class="pricing-feature-list">
                    <li>Everything in Professional</li>
                    <li>Unlimited simulations</li>
                    <li>SSO &amp; SAML</li>
                    <li>Custom integrations</li>
                    <li>Dedicated account manager</li>
                    <li>SLA guarantee</li>
                    <li>On-premise deployment option</li>
                    <li>Custom AI model training</li>
                  </ul>
                  <a href="{_href("about")}" class="pricing-plan-cta pricing-plan-cta-muted">Contact Sales</a>
                </article>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_about() -> None:
    values = [
        ("Innovation First", "We push the boundaries of what AI can do for project management.", "Innovation"),
        ("Trust & Security", "Enterprise-grade security and compliance are non-negotiable.", "Security"),
        ("Customer Obsession", "Every feature is driven by real feedback from project teams.", "Customer"),
        ("Global Perspective", "Built for teams of every size, across every continent.", "Global"),
    ]
    journey = [
        (
            "2026",
            "Idea & Foundation",
            "certAIn began from a clearly defined idea to bring probabilistic, AI-driven decision support into project planning. The initial phase focused on exploring how uncertainty, dependencies, and risk could be modeled more realistically.",
            "left",
        ),
        (
            "2026",
            "Method & Architecture",
            "The system approach was shaped through exploration of key methodologies, including machine learning, DAG-based dependency modeling, and Monte Carlo simulation, leading to a unified architecture connecting planning, forecasting, and monitoring.",
            "right",
        ),
        (
            "2026",
            "Data Strategy & Modeling",
            "Synthetic datasets were designed and generated to simulate realistic project conditions, followed by development and evaluation of machine learning models for risk prediction.",
            "left",
        ),
        (
            "2026",
            "Forecasting Engine",
            "A DAG-based Monte Carlo forecasting engine was implemented to simulate task-level uncertainty, identify the critical path, and generate P50/P80 delivery estimates.",
            "right",
        ),
        (
            "2026",
            "End-to-End System",
            "The full workflow was structured across notebooks, covering data analysis, modeling, forecasting, and monitoring, resulting in a reproducible and scalable pipeline.",
            "left",
        ),
        (
            "2026",
            "Product & Experience",
            "The system was translated into an interactive prototype, and consolidated into a clear product narrative aligned with user value and decision-making support.",
            "right",
        ),
    ]
    team_members = [
        (
            "Meysameh Shojaei",
            "Founder & CEO",
            "AI Product Leader building data-driven decision systems for next-generation project planning. Leads product strategy and execution from concept to deployment. Originator of the certAIn vision, creating a scalable AI platform that redefines how teams plan, forecast risk, and make critical decisions.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "CEO.jpeg")),
        ),
        (
            "Santiago Molina",
            "COO",
            "Industrial Engineer and AI Project Leader driving end-to-end digital initiatives in international environments. Specializes in product development, workflow optimization, and operational efficiency using Agile methodologies. Experienced in large-scale transformation, including process redesign and cloud migration.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "COO.jpg")),
        ),
        (
            "James Park",
            "CTO",
            "Ex-Microsoft engineer building scalable ML infrastructure and high-performance systems. Expert in data pipelines, model deployment, and simulation systems for complex forecasting and decision-making applications.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "CTO.png")),
        ),
        (
            "Maria Smith",
            "CFO",
            "Former Deloitte Partner with 18+ years in corporate finance and scaling technology ventures. Specializes in financial strategy, fundraising, and operational efficiency across AI and SaaS organizations.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "CFO.png")),
        ),
        (
            "Lisa Schmidt",
            "Head of AI",
            "PhD in Operations Research leading advanced AI and predictive analytics initiatives. Expert in forecasting models, optimization, and building scalable data-driven systems for complex planning challenges.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "Head AI.png")),
        ),
        (
            "David Chen",
            "Head of Product",
            "Product leader with experience building project management tools at Atlassian and Asana. Focused on scalable, user-centric products that improve workflow efficiency and team collaboration.",
            _image_data_uri(str(LEADERSHIP_IMAGES / "Head Product.png")),
        ),
    ]
    brand_name = "certAIn"
    logo_uri = _logo_data_uri()
    logo_markup = (
        f'<img src="{logo_uri}" alt="certAIn logo" />'
        if logo_uri
        else '<div class="brand-mark" aria-hidden="true">AI</div>'
    )
    st.markdown(
        f"""
        <div class="about-proto-shell">
          <div class="about-top-stage">
            <div class="about-proto-hero">
              <div class="about-proto-logo-shell">
                {logo_markup}
              </div>
              <h1 class="about-proto-title">About {_brand_name_html()}</h1>
              <p class="about-proto-subtitle">
                We're building the future of project intelligence — where AI helps teams deliver on time, every time.
              </p>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    st.markdown(
        """
        <div class="about-top-stage">
          <div class="about-mv-grid">
            <div class="about-mv-card">
              <div class="about-card-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <circle cx="12" cy="12" r="8"></circle>
                  <circle cx="12" cy="12" r="4"></circle>
                  <circle cx="12" cy="12" r="1.3"></circle>
                </svg>
              </div>
              <h3>Our Mission</h3>
              <p>
                To eliminate project failure by giving every team access to AI-powered forecasting and risk intelligence.
                We believe that project delays are predictable — and preventable — when teams have the right data at the
                right time.
              </p>
            </div>
            <div class="about-mv-card">
              <div class="about-card-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <path d="M2 12s3.5-6 10-6 10 6 10 6-3.5 6-10 6-10-6-10-6z"></path>
                  <circle cx="12" cy="12" r="2.8"></circle>
                </svg>
              </div>
              <h3>Our Vision</h3>
              <p>
                A world where every project plan comes with a probability of success. Where project managers can
                confidently tell stakeholders not just when a project will finish — but how certain they are about it.
              </p>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    st.markdown(
        """
        <div class="about-values-section">
          <div class="about-values-pill">
            <svg viewBox="0 0 24 24" aria-hidden="true">
              <path d="M12 3v4"></path>
              <path d="M12 17v4"></path>
              <path d="M4.93 4.93l2.83 2.83"></path>
              <path d="M16.24 16.24l2.83 2.83"></path>
              <path d="M3 12h4"></path>
              <path d="M17 12h4"></path>
              <path d="M4.93 19.07l2.83-2.83"></path>
              <path d="M16.24 7.76l2.83-2.83"></path>
            </svg>
            <span>What Drives Us</span>
          </div>
          <h2 class="about-section-heading">Our Core Values</h2>
          <div class="about-values-grid">
            <div class="about-value-card">
              <div class="about-value-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <path d="M13 2L4 14h7l-1 8 10-13h-7l0-7z"></path>
                </svg>
              </div>
              <h3>Innovation First</h3>
              <p>We push the boundaries of what AI can do for project management.</p>
            </div>
            <div class="about-value-card">
              <div class="about-value-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <path d="M12 3l7 4v5c0 5-3.5 8.5-7 9-3.5-.5-7-4-7-9V7l7-4z"></path>
                </svg>
              </div>
              <h3>Trust &amp; Security</h3>
              <p>Enterprise-grade security and compliance are non-negotiable.</p>
            </div>
            <div class="about-value-card">
              <div class="about-value-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <path d="M12 21s-7-4.35-7-10a4 4 0 0 1 7-2.53A4 4 0 0 1 19 11c0 5.65-7 10-7 10z"></path>
                </svg>
              </div>
              <h3>Customer Obsession</h3>
              <p>Every feature is driven by real feedback from project teams.</p>
            </div>
            <div class="about-value-card">
              <div class="about-value-icon" aria-hidden="true">
                <svg viewBox="0 0 24 24">
                  <circle cx="12" cy="12" r="9"></circle>
                  <path d="M3 12h18"></path>
                  <path d="M12 3a15 15 0 0 1 0 18"></path>
                  <path d="M12 3a15 15 0 0 0 0 18"></path>
                </svg>
              </div>
              <h3>Global Perspective</h3>
              <p>Built for teams of every size, across every continent.</p>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    timeline_rows = []
    for year, title, copy, side in journey:
        card_markup = (
            f'<article class="about-timeline-card">'
            f'<span class="about-timeline-year ai-gradient">{html.escape(year)}</span>'
            f'<h3 class="about-journey-step-title"><span class="ai-gradient">{html.escape(title)}</span></h3>'
            f'<p>{html.escape(copy)}</p>'
            f'</article>'
        )
        left_markup = card_markup if side == "left" else ""
        right_markup = card_markup if side == "right" else ""
        timeline_rows.append(
            f'<div class="about-timeline-row {html.escape(side)}">'
            f'<div class="about-timeline-side">{left_markup}</div>'
            f'<div class="about-timeline-axis"><span class="about-timeline-dot"></span></div>'
            f'<div class="about-timeline-side">{right_markup}</div>'
            f'</div>'
        )

    st.markdown(
        f"""
        <div class="about-journey-section" style="margin-bottom: 2.2rem;">
          <h2 class="about-section-heading">Transforming an initial idea into a risk-aware <span class="ai-gradient about-inline-ai">AI</span> planning platform</h2>
          <div class="about-timeline">
            {"".join(timeline_rows)}
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    team_cards_markup = "".join(
        f"""
        <article class="about-member-card">
          <div class="avatar-wrapper">
            <div class="avatar team-avatar">
              <img src="{html.escape(image_src, quote=True)}" alt="{html.escape(name)} portrait" loading="lazy" />
            </div>
          </div>
          <h3>{html.escape(name)}</h3>
          <div class="about-member-role">{html.escape(role)}</div>
          <p>{_brand_name_in_text_html(bio)}</p>
        </article>
        """
        for name, role, bio, image_src in team_members
    )

    st.write("")
    st.markdown(
        f"""
        <div class="about-team-section">
          <div class="about-team-head">
            <span class="about-team-icon" aria-hidden="true">
              <svg viewBox="0 0 24 24">
                <path d="M16 19v-1.2c0-1.8-1.8-3.3-4-3.3s-4 1.5-4 3.3V19"></path>
                <circle cx="12" cy="8.2" r="2.9"></circle>
                <path d="M20 18v-.9c0-1.3-1-2.5-2.5-3"></path>
                <path d="M17.5 5.8a2.5 2.5 0 0 1 0 4.8"></path>
                <path d="M4 18v-.9c0-1.3 1-2.5 2.5-3"></path>
                <path d="M6.5 5.8a2.5 2.5 0 0 0 0 4.8"></path>
              </svg>
            </span>
            <h2 class="about-section-heading">Leadership Team</h2>
          </div>
          <div class="about-team-grid">
            {team_cards_markup}
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.write("")
    cta_left, cta_center, cta_right = st.columns([1, 1.8, 1], gap="large")
    with cta_center:
        st.subheader("Ready to Transform Your Projects?")
        st.write(
            f"Join 500+ enterprise teams already using {brand_name} to deliver projects with AI-powered certainty."
        )
        if st.button("Contact Sales", key="about_contact_sales_button", type="primary", use_container_width=True):
            _set_query_page("contact")
            st.rerun()


def _render_contact() -> None:
    contact_social_markup = "".join(
        _footer_social_icon_markup(platform, label)
        for platform, label in FOOTER_SOCIALS
    )
    st.markdown(
        f"""
        <div class="page-shell">
          <div class="page-grid">
            <div class="glass page-hero-card">
              <span class="kicker">Contact Us</span>
              <h1 class="page-title">Talk to the team behind {_brand_name_html()}.</h1>
              <p class="page-copy">
                Book a demo, discuss enterprise rollout, or use this page in your capstone presentation to show a production-style contact journey.
              </p>
              <p class="page-copy">
                If you cannot find the answer to your question in the <a href="{_href("faq")}" class="auth-inline-link">FAQ</a>, please submit your inquiry using the contact form below.
              </p>
            </div>
            <div class="contact-top-stack">
              <div class="glass page-hero-card">
                <div class="data-pair"><span>Email</span><strong><a href="mailto:info@certain-pm.com">info@certain-pm.com</a></strong></div>
                <div class="data-pair"><span>Phone</span><strong><a href="tel:+13105552040">+1 (310) 555-2040</a></strong></div>
                <div class="data-pair"><span>Sales</span><strong><a href="mailto:enterprise@certain.ai.com">enterprise@certain.ai.com</a></strong></div>
                <div class="data-pair"><span>Support</span><strong><a href="mailto:support@certain-pm.com">support@certain-pm.com</a></strong></div>
              </div>
              <div class="glass page-hero-card">
                <span class="eyebrow">Global HQ</span>
                <h3>Beverly Hills</h3>
                <p>Rodeo Drive<br/><br/>300 Rodeo Dr<br/>Beverly Hills 90210<br/>CA, USA</p>
              </div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    left, right = st.columns([1.05, 0.95], gap="large")
    with left:
        with st.form("contact_form"):
            st.text_input("Name", key="contact_name")
            st.text_input("Email", key="contact_email")
            st.text_input("Company", key="contact_company")
            st.selectbox("Subject", ["Book demo", "Enterprise inquiry", "Partnership", "Support"], key="contact_subject")
            st.text_area("Message", height=180, key="contact_message")
            submitted = st.form_submit_button("Send message", type="primary", use_container_width=True)
            if submitted:
                st.success("Your request has been captured for the demo flow.")
    with right:
        st.markdown(
            f"""
            <div class="contact-right-balance">
                <div class="feature-grid">
                <div class="feature-card">
                  <span class="eyebrow">Tehran</span>
                  <h3>Niavaran Square</h3>
                  <p>20 Niavaran<br/>Tehran 19789<br/>Tehran, Iran</p>
                </div>
                <div class="feature-card">
                  <span class="eyebrow">Tel Aviv</span>
                  <h3>Rothschild Blvd</h3>
                  <p>50 Rothschild Blvd<br/>Tel Aviv 6688125<br/>Tel Aviv District, Israel</p>
                </div>
                <div class="feature-card">
                  <span class="eyebrow">New York</span>
                  <h3>Manhattan</h3>
                  <p>725 5th Ave<br/>New York 10022<br/>NY, USA</p>
                </div>
              </div>
              <div class="contact-social-inline">
                <strong>Connect with us</strong>
                <div class="footer-social-row">{contact_social_markup}</div>
                <p>LinkedIn · X · GitHub · YouTube · Instagram · Facebook</p>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

def _render_faq() -> None:
    left_questions = [
        (
            "Can certAIn connect to Jira, Asana, and MS Project?",
            "Yes. The Tools Integration page includes those connectors in the SaaS story.",
        ),
        (
            "Does the platform support executive reporting?",
            "Yes. Executive Summary and Decision Hub are designed for stakeholder-ready communication.",
        ),
        (
            "Can the AI model be monitored after deployment?",
            "Yes. The dashboard reads `risk_monitoring_snapshot.csv` to visualize monitoring drift and confidence.",
        ),
        (
            "Is certAIn limited to one industry?",
            "No. The platform flow is designed to be generalizable across project domains, while the bundled advisory model is just one current demonstration model.",
        ),
        (
            "What makes the forecasts credible?",
            "The app combines dependency-aware task graphs, Monte Carlo simulation, and risk signals so teams can discuss ranges and confidence levels instead of a single optimistic date.",
        ),
    ]
    middle_questions = [
        (
            "Can users upload project files in this demo?",
            "Yes. The Project Upload page shows the intake flow and previews CSV content in the current prototype experience.",
        ),
        (
            "Can certAIn compare multiple delivery scenarios?",
            "Yes. The What-If Scenarios page is designed to compare mitigation and acceleration options before teams commit to a plan.",
        ),
        (
            "Does the platform include a guided workspace experience?",
            "Yes. The dashboard modules walk users through planning, risk analysis, simulation, scenario review, and executive reporting.",
        ),
        (
            "Can results be shared with stakeholders?",
            "Yes. The platform includes executive-ready summaries, charts, and export-oriented flows to support stakeholder conversations.",
        ),
        (
            "Can certAIn help explain project risk to non-technical audiences?",
            "Yes. The platform is designed to turn simulation outputs, delay drivers, and confidence ranges into language that sponsors and stakeholders can understand quickly.",
        ),
    ]
    right_questions = [
        (
            "Why should project managers choose certAIn over traditional planning tools?",
            "certAIn brings AI task generation, dependency-aware planning, Monte Carlo forecasting, scenario comparison, and executive-ready reporting into one connected workflow. Instead of forcing PMs to switch between separate tools for planning, analysis, and stakeholder communication, certAIn helps teams move from project setup to risk-aware decision-making in one platform.",
        ),
        (
            "What do the Monte Carlo P50 and P80 completion bars mean?",
            "P50 is the date with a 50% chance of completion, while P80 is the more conservative date with an 80% chance of completion.",
        ),
    ]

    st.markdown(
        f"""
        <div class="page-shell">
          <div class="page-grid">
            <div class="glass page-hero-card">
              <span class="kicker">FAQ</span>
              <h1 class="page-title">Answers to the questions people ask during the {_brand_name_html()} walkthrough.</h1>
              <p class="page-copy">
                Use this page to explain the product story, platform scope, and the evidence behind the forecasting workflow without leaving the demo.
              </p>
            </div>
            <div class="glass page-hero-card faq-summary-card">
              <div class="data-pair"><span>Audience</span><strong>Project &amp; Product Managers, PMOs, Decision-Makers</strong></div>
              <div class="data-pair"><span>Focus</span><strong>AI-driven planning, dependency modeling, risk forecasting, decision support</strong></div>
              <div class="data-pair"><span>Coverage</span><strong>End-to-end workflow: planning → simulation → monitoring → iteration</strong></div>
              <div class="data-pair"><span>Format</span><strong>Clear, actionable, presentation-ready insights</strong></div>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    left_col, middle_col, right_col = st.columns(3, gap="large")
    with left_col:
        for question, answer in left_questions:
            with st.expander(question):
                st.write(answer)

    with middle_col:
        for question, answer in middle_questions:
            with st.expander(question):
                st.write(answer)

    with right_col:
        for question, answer in right_questions:
            with st.expander(question):
                st.write(answer)
        for _ in range(3):
            st.markdown('<div class="faq-empty-box"></div>', unsafe_allow_html=True)


def _render_legal_page(page: str) -> None:
    legal_pages: dict[str, dict[str, str]] = {
        "privacy-policy": {
            "kicker": "Privacy Policy",
            "title": "How certAIn handles personal and project-related information.",
            "intro": "This demo page represents the privacy layer of the platform and explains how contact details, project inputs, and workflow signals would be handled in a production-style environment.",
            "scope": "Data scope",
            "scope_value": "Contact details, workspace inputs, and planning metadata",
            "purpose": "Use of data",
            "purpose_value": "Platform access, forecasting workflows, and product communication",
            "control": "User controls",
            "control_value": "Review, update, or request removal through support channels",
            "note": "Presentation note",
            "note_value": "Shown here as a product-trust page for the capstone platform",
        },
        "terms-of-service": {
            "kicker": "Terms of Service",
            "title": "The product-use terms that frame access to certAIn.",
            "intro": "This page presents the platform terms at a product-story level, including acceptable use, trial access, and expectations around evaluation or enterprise rollout conversations.",
            "scope": "Applies to",
            "scope_value": "Trial users, team pilots, and enterprise evaluations",
            "purpose": "Use of platform",
            "purpose_value": "Project planning, forecasting, simulation, and reporting workflows",
            "control": "Account expectations",
            "control_value": "Responsible access, protected credentials, and lawful use",
            "note": "Presentation note",
            "note_value": "Structured as a SaaS trust page within the demo experience",
        },
        "cookie-settings": {
            "kicker": "Cookie Settings",
            "title": "How browser preferences and consent controls would be managed.",
            "intro": "This page represents the preferences layer for analytics, language persistence, and session experience in a production-style product environment.",
            "scope": "Preference areas",
            "scope_value": "Language, session continuity, analytics, and UI preferences",
            "purpose": "Primary use",
            "purpose_value": "Improve navigation, remember settings, and support product insights",
            "control": "User controls",
            "control_value": "Adjustable through browser and in-product preference settings",
            "note": "Presentation note",
            "note_value": "Included to make the platform footer feel complete and credible",
        },
        "security-data-protection": {
            "kicker": "Security & Data Protection",
            "title": "How certAIn frames security, privacy, and data protection.",
            "intro": "This page summarizes the product's trust posture, including access controls, monitoring, compliance framing, and enterprise-style data protection expectations.",
            "scope": "Security focus",
            "scope_value": "Access control, monitoring, privacy, and controlled sharing",
            "purpose": "Protection model",
            "purpose_value": "Safeguard user data, project context, and reporting outputs",
            "control": "Trust signals",
            "control_value": "Compliance badges, secure flows, and privacy-forward design",
            "note": "Presentation note",
            "note_value": "Supports the trust and enterprise-readiness story of the platform",
        },
    }

    content = legal_pages.get(page)
    if not content:
        _render_dashboard()
        return

    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="page-grid">
                <div class="glass page-hero-card">
                  <span class="kicker">{content["kicker"]}</span>
                  <h1 class="page-title">{content["title"]}</h1>
                  <p class="page-copy">
                    {content["intro"]}
                  </p>
                  <div style="margin-top:1.25rem; display:flex; gap:0.9rem; flex-wrap:wrap;">
                    <a href="{_href("home")}" class="btn btn-primary">Back to Home</a>
                    <a href="{_href("contact")}" class="btn btn-secondary">Contact Us</a>
                  </div>
                </div>
                <div class="glass page-hero-card">
                  <div class="data-pair"><span>{content["scope"]}</span><strong>{content["scope_value"]}</strong></div>
                  <div class="data-pair"><span>{content["purpose"]}</span><strong>{content["purpose_value"]}</strong></div>
                  <div class="data-pair"><span>{content["control"]}</span><strong>{content["control_value"]}</strong></div>
                  <div class="data-pair"><span>{content["note"]}</span><strong>{content["note_value"]}</strong></div>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_subscription_confirmation() -> None:
    subscribed_email = st.session_state.get("contact_email", "").strip()
    safe_email = html.escape(subscribed_email)
    confirmed = bool(st.session_state.get("newsletter_confirmed"))
    delivery_status = st.session_state.get("newsletter_delivery_status", "")
    delivery_message = html.escape(st.session_state.get("newsletter_delivery_message", "").strip())

    if confirmed:
        title = "Subscription confirmed."
        intro = (
            f"<strong>{safe_email}</strong> is now confirmed for certAIn updates."
            if subscribed_email
            else "Your subscription is now confirmed."
        )
        status_pairs = [
            ("Status", "Confirmed"),
            ("Updates", "Product and release news"),
            ("Flow", "Email confirmation completed"),
        ]
    else:
        title = "Check your email to confirm the subscription."
        if delivery_status == "sent":
            intro = delivery_message
        elif delivery_status == "disabled":
            intro = (
                "Email delivery is not configured yet, so the app could not send the confirmation message."
            )
        elif delivery_status == "error":
            intro = delivery_message or "We could not send the confirmation email right now."
        else:
            intro = (
                f"We are preparing the confirmation email for <strong>{safe_email}</strong>."
                if subscribed_email
                else "We are preparing your confirmation email."
            )
        status_pairs = [
            ("Confirmation", "Email verification"),
            ("Updates", "Product and release news"),
            ("Flow", "Footer subscribe experience"),
        ]

    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="page-grid">
                <div class="glass page-hero-card">
                  <span class="kicker">Newsletter</span>
                  <h1 class="page-title">{title}</h1>
                  <p class="page-copy">
                    {intro}
                  </p>
                  <div style="margin-top:1.25rem; display:flex; gap:0.9rem; flex-wrap:wrap;">
                    <a href="{_href("home")}" class="btn btn-primary">Back to Home</a>
                    <a href="{_href("product")}" class="btn btn-secondary">Explore Product</a>
                  </div>
                </div>
                <div class="glass page-hero-card">
                  <span class="eyebrow">What happens next</span>
                  {"".join(f'<div class="data-pair"><span>{label}</span><strong>{value}</strong></div>' for label, value in status_pairs)}
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_registration_confirmation() -> None:
    registered_email = st.session_state.get("registration_email", "").strip()
    safe_email = html.escape(registered_email)
    delivery_status = st.session_state.get("registration_delivery_status", "")
    delivery_message = html.escape(st.session_state.get("registration_delivery_message", "").strip())
    email_status = "Confirmation email sent" if delivery_status == "sent" else "Email delivery pending"

    if delivery_status == "sent":
        intro = delivery_message
    elif delivery_status == "disabled":
        intro = "Your registration has been confirmed, but the confirmation email could not be sent yet."
    elif delivery_status == "error":
        intro = delivery_message or "Your registration has been confirmed, but the confirmation email could not be sent right now."
    else:
        intro = (
            f"<strong>{safe_email}</strong> is now registered with certAIn."
            if registered_email
            else "Your certAIn registration is complete."
        )

    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="page-grid">
                <div class="glass page-hero-card">
                  <span class="kicker">Registration</span>
                  <h1 class="page-title">Your registration has been confirmed.</h1>
                  <p class="page-copy">
                    {intro}
                  </p>
                  <div style="margin-top:1.25rem; display:flex; gap:0.9rem; flex-wrap:wrap;">
                    <a href="{_href("pricing")}" class="btn btn-primary">Continue to Pricing</a>
                    <a href="{_href("dashboard")}" class="btn btn-secondary">Open Dashboard</a>
                  </div>
                </div>
                <div class="glass page-hero-card">
                  <span class="eyebrow">What happens next</span>
                  <div class="data-pair"><span>Status</span><strong>Registration complete</strong></div>
                  <div class="data-pair"><span>Email</span><strong>{safe_email or "Stored for this session"}</strong></div>
                  <div class="data-pair"><span>Flow</span><strong>{email_status}</strong></div>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_login() -> None:
    logo_uri = _nav_logo_data_uri()
    selected_language = st.session_state.get("saas_language_code", "en")
    welcome_title = "خوش آمدید" if selected_language == "fa" else "Welcome back"
    google_label = "ادامه با گوگل" if selected_language == "fa" else "Continue with Google"
    logo_markup = (
        f'<img src="{logo_uri}" class="auth-logo-image" alt="certAIn logo" />'
        if logo_uri
        else '<span class="auth-logo-fallback">A</span>'
    )
    st.markdown(
        dedent(
            f"""
            <div class="page-shell auth-page-shell">
              <div class="auth-card">
                <div class="auth-logo-shell">{logo_markup}</div>
                <h1 class="auth-title">{welcome_title}</h1>
                <p class="auth-title-helper">Don't have an account?</p>
                <form class="auth-form-shell" method="get" action="">
                  <input type="hidden" name="page" value="registration-confirmation" />
                  <input type="hidden" name="lang" value="{st.session_state.get("saas_language_code", "en")}" />
                  <div class="auth-preview-stack">
                    <div class="auth-field-group">
                      <label>Email</label>
                      <input class="auth-input-field" type="email" name="register_email" placeholder="team@certain.ai" autocomplete="email" />
                    </div>
                    <div class="auth-field-group">
                      <label>Password</label>
                      <input class="auth-input-field password" type="password" placeholder="........" autocomplete="new-password" />
                    </div>
                  </div>
                  <button type="submit" class="auth-register-button">Register Now</button>
                </form>
                <div class="auth-social-stack">
                  <a href="{_href("dashboard")}" class="auth-social-button">
                    <span>{google_label}</span>
                  </a>
                  <a href="{_href("dashboard")}" class="auth-social-button">
                    <span class="auth-social-icon apple"></span>
                    <span>Continue with Apple</span>
                  </a>
                </div>
                <div class="auth-separator"><span>or continue with email</span></div>
                <form class="auth-form-shell" method="get" action="">
                  <input type="hidden" name="page" value="pricing" />
                  <input type="hidden" name="lang" value="{st.session_state.get("saas_language_code", "en")}" />
                  <div class="auth-field-group">
                    <label>Email</label>
                    <input class="auth-input-field" type="email" name="auth_email" placeholder="you@company.com" autocomplete="email" />
                  </div>
                  <div class="auth-field-group">
                    <label>Password</label>
                    <input class="auth-input-field password" type="password" placeholder="........" autocomplete="current-password" />
                  </div>
                  <div class="auth-meta-row">
                    <label class="auth-check-row">
                      <input type="checkbox" class="auth-checkbox-input" name="remember_me" />
                      <span>Remember me</span>
                    </label>
                    <a href="{_href("contact")}" class="auth-forgot-link">Forgot password?</a>
                  </div>
                  <button type="submit" class="auth-submit-button">Sign In</button>
                </form>
              </div>
              <div class="auth-trust-row">
                <span class="auth-trust-pill"><span class="auth-trust-dot"></span>256-bit encryption</span>
                <span class="auth-trust-pill"><span class="auth-trust-dot"></span>SOC2 &amp; GDPR</span>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )


def _render_dashboard() -> None:
    _ensure_demo_payload()
    payload = st.session_state.saas_payload
    model_status = payload["model_status"] if payload else {}
    model_chip = (
        '<span class="status-chip chip-low">ML advisory ready</span>'
        if model_status.get("available", False)
        else '<span class="status-chip chip-medium">Monte Carlo mode</span>'
    )

    st.markdown(
        dedent(
            f"""
            <div class="page-shell">
              <div class="dashboard-hero-grid">
                <div class="glass dashboard-banner">
                  <span class="kicker">Forecast Console</span>
                  <h1 class="page-title">Operate the planning workspace like a real SaaS product.</h1>
                  <p class="page-copy">
                    Generate a plan, quantify deadline risk, compare scenarios, and present a stakeholder-safe recommendation
                    from a single interface.
                  </p>
                  <div class="cta-row">
                    {model_chip}
                    <span class="mini-pill">Default mode: {st.session_state.saas_mode.upper()}</span>
                    <span class="mini-pill">Workspace owner: {st.session_state.saas_user_email}</span>
                  </div>
                </div>
                <div class="glass dashboard-banner">
                  <span class="eyebrow">Workspace pulse</span>
                  <div class="data-pair"><span>Sample scenario</span><strong>{st.session_state.saas_selected_sample}</strong></div>
                  <div class="data-pair"><span>Iterations</span><strong>{int(st.session_state.saas_iterations):,}</strong></div>
                  <div class="data-pair"><span>Deadline</span><strong>{float(st.session_state.saas_deadline):.0f} days</strong></div>
                  <div class="data-pair"><span>Recent workspaces</span><strong>{len(st.session_state.saas_history)}</strong></div>
                </div>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )

    _section_intro(
        "Launch workspace",
        "Control the briefing flow and run a fresh forecast",
        "Use demo scenarios for presentation speed, or switch to real mode with your Groq key for live task generation.",
    )

    left, right = st.columns([1.18, 0.82], gap="large")
    with left:
        st.selectbox(
            "Demo scenario",
            options=list(SAMPLE_SCENARIOS),
            key="saas_selected_sample",
        )
        st.text_area(
            "Project brief",
            key="saas_project_text",
            height=180,
            placeholder="Describe the project, scope, constraints, and what success looks like.",
        )
        st.radio(
            "Operating mode",
            options=["mock", "real"],
            key="saas_mode",
            horizontal=True,
        )
        if st.session_state.saas_mode == "real" and not settings.groq_api_key:
            st.info("`real` mode needs `GROQ_API_KEY`. Keep `mock` mode for presentation-safe demos.")
    with right:
        st.number_input("Deadline (days)", min_value=10.0, step=1.0, key="saas_deadline")
        st.number_input("Simulation runs", min_value=200, max_value=settings.max_iterations, step=100, key="saas_iterations")
        st.number_input("Max tasks", min_value=8, max_value=18, step=1, key="saas_max_tasks")
        st.number_input("Random seed", min_value=0, max_value=100_000, step=1, key="saas_seed")
        action_a, action_b = st.columns(2)
        with action_a:
            if st.button("Load demo brief", use_container_width=True):
                _load_selected_sample()
                st.rerun()
        with action_b:
            if st.button("Generate forecast", type="primary", use_container_width=True):
                _run_forecast()
                st.rerun()
        st.markdown(
            """
            <div class="history-card">
              <strong style="font-family:Sora,sans-serif;">Presentation tip</strong>
              <p>Open with a saved sample, then edit the brief live so the judges see both product polish and analytics depth.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    if not payload:
        return

    metrics = payload["metrics"]
    drivers = payload["drivers"]
    top_driver = drivers[0]["task_name"] if drivers else None
    risk_level = _risk_level(metrics["delay_probability"])
    confidence = _confidence_score(metrics)
    critical_path = payload["critical_path"]
    ml_summary_df = payload["ml_summary_df"]
    scenarios_df = payload["scenarios_df"]
    recommendation = scenarios_df.sort_values("Delay Prob (%)", ascending=True).iloc[0]

    st.markdown(
        dedent(
            f"""
            <div class="kpi-grid">
              <div class="metric-card">
                <span>Delay probability</span>
                <strong>{metrics['delay_probability'] * 100:.1f}%</strong>
                <small>Risk posture: {risk_level.upper()}</small>
              </div>
              <div class="metric-card">
                <span>P80 commitment</span>
                <strong>{metrics['p80']:.1f}d</strong>
                <small>Recommended stakeholder date</small>
              </div>
              <div class="metric-card">
                <span>Forecast confidence</span>
                <strong>{confidence:.0f}%</strong>
                <small>Blends uncertainty spread and delay exposure</small>
              </div>
              <div class="metric-card">
                <span>Primary delay driver</span>
                <strong>{top_driver or "N/A"}</strong>
                <small>Most frequent critical-path pressure point</small>
              </div>
            </div>
            """
        ),
        unsafe_allow_html=True,
    )

    summary_col, workspace_col = st.columns([1.16, 0.84], gap="large")
    with summary_col:
        st.markdown(
            dedent(
                f"""
                <div class="workspace-grid">
                  <div class="workspace-card">
                    <span class="eyebrow">Executive narrative</span>
                    <p>{_narrative(metrics, top_driver)}</p>
                  </div>
                  <div class="workspace-card">
                    <span class="eyebrow">Recommended scenario</span>
                    <strong>{recommendation['Scenario']}</strong>
                    <p>
                      Best observed tradeoff in this run. Delay risk drops to {float(recommendation['Delay Prob (%)']):.1f}%
                      with a projected P80 of {float(recommendation['P80 (days)']):.1f} days.
                    </p>
                  </div>
                </div>
                """
            ),
            unsafe_allow_html=True,
        )
    with workspace_col:
        st.markdown(
            f"""
            <div class="history-card">
              <strong style="font-family:Sora,sans-serif;">Recent workspaces</strong>
              {_history_markup()}
            </div>
            """,
            unsafe_allow_html=True,
        )

    dataset_cols = st.columns(4)
    dataset_cols[0].metric("Portfolio Projects", f"{int(payload['portfolio_summary']['projects'])}")
    dataset_cols[1].metric("Portfolio Delay Rate", f"{payload['portfolio_summary']['delayed_pct']:.1f}%")
    dataset_cols[2].metric("Monitoring Accuracy", f"{payload['monitoring_summary']['accuracy_pct']:.1f}%")
    dataset_cols[3].metric("Open Monitoring Alerts", f"{int(payload['monitoring_summary']['alerts_open'])}")

    hist_col, graph_col = st.columns([1.05, 0.95], gap="large")
    with hist_col:
        histogram = completion_histogram(
            payload["completion"],
            payload["deadline_days"],
            metrics["p50"],
            metrics["p80"],
        )
        st.plotly_chart(_theme_chart(histogram), use_container_width=True)
    with graph_col:
        graph_figure = dependency_graph_figure(payload["graph"], critical_path)
        st.plotly_chart(_theme_chart(graph_figure), use_container_width=True)

    driver_col, scenario_col = st.columns([1.0, 1.0], gap="large")
    with driver_col:
        drivers_df = pd.DataFrame(drivers)
        st.plotly_chart(_theme_chart(risk_drivers_bar(drivers_df)), use_container_width=True)
    with scenario_col:
        st.plotly_chart(_theme_chart(scenario_comparison_chart(scenarios_df)), use_container_width=True)

    path_col, scenario_cards_col = st.columns([0.95, 1.05], gap="large")
    with path_col:
        critical_markup = "".join(
            f"""
            <div class="task-item">
              <div>
                <strong>{task_id}</strong>
                <span>{payload['graph'].nodes[task_id].get('name', task_id)}</span>
              </div>
              <div style="text-align:right;">
                <strong>{float(payload['graph'].nodes[task_id].get('mean_duration', 0.0)):.1f}d</strong>
                <span>mean duration</span>
              </div>
            </div>
            """
            for task_id in critical_path
        )
        st.markdown(
            f"""
            <div class="timeline-card">
              <span class="eyebrow">Critical path</span>
              <strong>{payload['critical_path_days']:.1f} mean days</strong>
              <p>The baseline longest path through the dependency graph.</p>
              <div class="task-list">{critical_markup}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
    with scenario_cards_col:
        scenario_cards = "".join(
            f"""
            <div class="scenario-mini-card">
              <span class="eyebrow">{row['Scenario']}</span>
              <h4>{float(row['Delay Prob (%)']):.1f}% risk</h4>
              <p>P80 {float(row['P80 (days)']):.1f} days · Mean {float(row['Mean (days)']):.1f} days</p>
              <p class="dashboard-note">{row['Notes']}</p>
            </div>
            """
            for _, row in scenarios_df.iterrows()
        )
        st.markdown(
            f"""
            <div class="timeline-card">
              <span class="eyebrow">Scenario cards</span>
              <p>Use these callouts when presenting the mitigation logic.</p>
              <div class="scenario-card-grid">{scenario_cards}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

    tabs = st.tabs(["Task plan", "ML advisory", "Export"])
    with tabs[0]:
        tasks_df = pd.DataFrame(payload["tasks"])
        st.dataframe(tasks_df, use_container_width=True, hide_index=True)
    with tabs[1]:
        if model_status.get("available", False):
            ml_cols = st.columns([0.85, 1.15], gap="large")
            with ml_cols[0]:
                st.dataframe(ml_summary_df, use_container_width=True, hide_index=True)
            with ml_cols[1]:
                st.dataframe(payload["ml_predictions_df"], use_container_width=True, hide_index=True)
        else:
            st.warning(
                "ML advisory scoring is unavailable in this environment. "
                f"Reason: {model_status.get('message', 'Unknown')}"
            )
            st.dataframe(payload["ml_features_df"], use_container_width=True, hide_index=True)
    with tabs[2]:
        st.download_button(
            "Download workspace summary",
            data=_serialize_payload(payload),
            file_name="certain_workspace_summary.json",
            mime="application/json",
            use_container_width=True,
        )
        st.markdown(
            """
            <div class="history-card">
              <strong style="font-family:Sora,sans-serif;">What is included</strong>
              <p>Project brief, metrics, critical path, delay drivers, scenarios, tasks, and ML scoring output.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )


def _render_user_guide() -> None:
    st.markdown(
        f"""
        <div class="page-shell">
          <div class="glass dashboard-banner">
            <span class="kicker">User Guide</span>
            <h1 class="page-title">Welcome to the {_brand_name_html()} workspace.</h1>
            <p class="page-copy">
              This landing page helps our users understand how to move through the platform from upload to executive reporting.
            </p>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="guide-how-section">
          <div class="guide-how-head">
            <svg viewBox="0 0 24 24" aria-hidden="true">
              <path d="M8 6l3-3"></path>
              <path d="M4 10l4-4 3 3-4 4z"></path>
              <path d="M13 11l3-3"></path>
              <path d="M9 15l4-4 3 3-4 4z"></path>
              <path d="M17 7l3-3"></path>
            </svg>
            <span>How It Works</span>
          </div>
          <div class="glass guide-how-shell">
            <div class="guide-how-grid">
              <div class="guide-how-card">
                <span class="guide-how-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M12 16V5"></path>
                    <path d="M8 9l4-4 4 4"></path>
                    <path d="M5 17v2h14v-2"></path>
                  </svg>
                </span>
                <strong>Upload</strong>
                <span>Import your project plan</span>
              </div>
              <div class="guide-how-card">
                <span class="guide-how-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M10 5a3 3 0 1 0-6 0c0 1.6 1.1 2.8 2.4 3.4"></path>
                    <path d="M14 5a3 3 0 1 1 6 0c0 1.6-1.1 2.8-2.4 3.4"></path>
                    <path d="M7 13a3 3 0 1 0 0 6"></path>
                    <path d="M17 13a3 3 0 1 1 0 6"></path>
                    <path d="M7 8h10"></path>
                    <path d="M12 8v8"></path>
                  </svg>
                </span>
                <strong>Analyze</strong>
                <span>AI scans for risks</span>
              </div>
              <div class="guide-how-card">
                <span class="guide-how-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M3 12h4l2-6 4 12 2-6h6"></path>
                  </svg>
                </span>
                <strong>Simulate</strong>
                <span>Run Monte Carlo forecasts</span>
              </div>
              <div class="guide-how-card">
                <span class="guide-how-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M12 3l7 4v5c0 5-3.5 8.5-7 9-3.5-.5-7-4-7-9V7l7-4z"></path>
                  </svg>
                </span>
                <strong>Review</strong>
                <span>Check AI insights</span>
              </div>
              <div class="guide-how-card">
                <span class="guide-how-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M7 4h7l5 5v11H7z"></path>
                    <path d="M14 4v5h5"></path>
                    <path d="M10 13h4"></path>
                    <path d="M10 17h4"></path>
                  </svg>
                </span>
                <strong>Share</strong>
                <span>Export executive report</span>
              </div>
            </div>
          </div>
        </div>
    """,
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <section id="modules" class="guide-modules-section">
          <div class="guide-modules-layout">
            <div class="guide-modules-main">
              <div class="guide-modules-head">
                <svg viewBox="0 0 24 24" aria-hidden="true">
                  <path d="M4 5h7v7H4z"></path>
                  <path d="M13 5h7v7h-7z"></path>
                  <path d="M4 13h7v7H4z"></path>
                  <path d="M13 13h7v7h-7z"></path>
                </svg>
                <span>Modules</span>
              </div>
              <div class="guide-modules-grid">
                <a class="guide-module-card guide-module-link" href="{_href("decision-hub")}">
                  <div class="guide-module-head">
                    <span class="guide-module-icon decision" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M5 5h5v5H5z"></path>
                        <path d="M14 5h5v5h-5z"></path>
                        <path d="M5 14h5v5H5z"></path>
                        <path d="M14 14h5v5h-5z"></path>
                      </svg>
                    </span>
                    <span class="guide-module-title">Decision Hub</span>
                  </div>
                  <p class="guide-module-copy">
                    Your command center for project intelligence. Real-time overview of project health, key metrics, and AI-powered predictive alerts.
                  </p>
                  <div class="guide-module-preview guide-sparkline" aria-hidden="true">
                    <svg viewBox="0 0 420 84" preserveAspectRatio="none">
                      <defs>
                        <linearGradient id="guideDecisionFill" x1="0%" y1="0%" x2="100%" y2="0%">
                          <stop offset="0%" stop-color="#386dff" stop-opacity="0.24" />
                          <stop offset="100%" stop-color="#19d4d3" stop-opacity="0.06" />
                        </linearGradient>
                        <linearGradient id="guideDecisionStroke" x1="0%" y1="0%" x2="100%" y2="0%">
                          <stop offset="0%" stop-color="#7b62ff" />
                          <stop offset="100%" stop-color="#26d6e7" />
                        </linearGradient>
                      </defs>
                      <path d="M0 18H420" stroke="rgba(255,196,92,0.45)" stroke-width="1.5" stroke-dasharray="7 7"></path>
                      <path d="M0 70C42 70 70 68 98 64C128 60 151 56 177 50C211 43 244 32 279 24C316 16 352 10 420 10V84H0Z" fill="url(#guideDecisionFill)"></path>
                      <path d="M0 70C42 70 70 68 98 64C128 60 151 56 177 50C211 43 244 32 279 24C316 16 352 10 420 10" stroke="url(#guideDecisionStroke)" stroke-width="3.2" fill="none" stroke-linecap="round"></path>
                    </svg>
                  </div>
                </a>
                <a class="guide-module-card guide-module-link" href="{_href("task-architect")}">
                  <div class="guide-module-head">
                    <span class="guide-module-icon task" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M10 5a3 3 0 1 0-6 0c0 1.6 1.1 2.8 2.4 3.4"></path>
                        <path d="M14 5a3 3 0 1 1 6 0c0 1.6-1.1 2.8-2.4 3.4"></path>
                        <path d="M7 13a3 3 0 1 0 0 6"></path>
                        <path d="M17 13a3 3 0 1 1 0 6"></path>
                        <path d="M7 8h10"></path>
                        <path d="M12 8v8"></path>
                      </svg>
                    </span>
                    <span class="guide-module-title">AI Task Architect</span>
                  </div>
                  <p class="guide-module-copy">
                    Let AI decompose your project goal into structured tasks with dependencies, durations, and risk assessments.
                  </p>
                  <div class="guide-module-preview">
                    <span class="guide-module-note-kicker">AI Analysis</span>
                    <p class="guide-module-note-copy">
                      Project health: 72/100. API Integration is the primary risk driver (78% delay probability). Recommend adding 5-day buffer to critical path. Budget performance strong at 1.08 CPI.
                    </p>
                  </div>
                </a>
                <a class="guide-module-card guide-module-link" href="{_href("simulator")}">
                  <div class="guide-module-head">
                    <span class="guide-module-icon monte" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M3 12h4l2-6 4 12 2-6h6"></path>
                      </svg>
                    </span>
                    <span class="guide-module-title">Monte Carlo Simulator</span>
                  </div>
                  <p class="guide-module-copy">
                    Run thousands of probabilistic simulations to forecast timelines with P50 and P80 confidence.
                  </p>
                  <div class="guide-module-preview guide-monte-preview" aria-hidden="true">
                    <svg viewBox="0 0 420 84" preserveAspectRatio="none">
                      <defs>
                        <linearGradient id="guideMonteFill" x1="0%" y1="0%" x2="0%" y2="100%">
                          <stop offset="0%" stop-color="#4f89ff" stop-opacity="0.95" />
                          <stop offset="100%" stop-color="#2e66c7" stop-opacity="0.85" />
                        </linearGradient>
                      </defs>
                      <rect x="8" y="72" width="16" height="4" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="28" y="68" width="16" height="8" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="48" y="62" width="16" height="14" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="68" y="54" width="16" height="22" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="88" y="44" width="16" height="32" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="108" y="34" width="16" height="42" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="128" y="25" width="16" height="51" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="148" y="18" width="16" height="58" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="168" y="12" width="16" height="64" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="188" y="9" width="16" height="67" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="208" y="7" width="16" height="69" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="228" y="8" width="16" height="68" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="248" y="13" width="16" height="63" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="268" y="24" width="16" height="52" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="288" y="36" width="16" height="40" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="308" y="45" width="16" height="31" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="328" y="54" width="16" height="22" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="348" y="62" width="16" height="14" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="368" y="68" width="16" height="8" rx="4" fill="url(#guideMonteFill)"></rect>
                      <rect x="388" y="72" width="16" height="4" rx="4" fill="url(#guideMonteFill)"></rect>
                    </svg>
                  </div>
                </a>
                <a class="guide-module-card guide-module-link" href="{_href("summary")}">
                  <div class="guide-module-head">
                    <span class="guide-module-icon summary" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M7 4h7l5 5v11H7z"></path>
                        <path d="M14 4v5h5"></path>
                        <path d="M10 13h4"></path>
                        <path d="M10 17h4"></path>
                      </svg>
                    </span>
                    <span class="guide-module-title">Executive Summary</span>
                  </div>
                  <p class="guide-module-copy">
                    Generate stakeholder-ready reports with health scores, SPI/CPI, and prioritized mitigation actions.
                  </p>
                  <div class="guide-module-preview">
                    <span class="guide-module-note-kicker">AI Analysis</span>
                    <p class="guide-module-note-copy">
                      Project health: 72/100. API Integration is the primary risk driver (78% delay probability). Recommend adding 5-day buffer to critical path. Budget performance strong at 1.08 CPI.
                    </p>
                  </div>
                </a>
              </div>
            </div>
            <aside class="guide-activity-panel">
              <h3 class="guide-activity-title">AI Activity</h3>
              <div class="guide-activity-block">
                <p class="guide-activity-subhead">Live AI Activity</p>
                <div class="guide-activity-list">
                  <div class="guide-activity-item">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M4 18l5-5 3 3 8-8"></path>
                        <path d="M15 8h5v5"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-copy">Dependency graph recalculated — critical path: 42 days</span>
                    <span class="guide-activity-time">Just now</span>
                  </div>
                  <div class="guide-activity-item">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M7 4h7l5 5v11H7z"></path>
                        <path d="M14 4v5h5"></path>
                        <path d="M10 13h4"></path>
                        <path d="M10 17h4"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-copy">Executive report generated with 4 mitigation actions</span>
                    <span class="guide-activity-time">Just now</span>
                  </div>
                  <div class="guide-activity-item">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M4 19h16"></path>
                        <path d="M7 15l3-4 3 2 4-6"></path>
                        <path d="M7 9h.01"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-copy">Schedule performance index updated: SPI 0.94</span>
                    <span class="guide-activity-time">Just now</span>
                  </div>
                  <div class="guide-activity-item">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M3 12h4l2-6 4 12 2-6h6"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-copy">Monte Carlo simulation finished — 10,000 runs processed</span>
                    <span class="guide-activity-time">Just now</span>
                  </div>
                </div>
              </div>
              <div class="guide-activity-block">
                <p class="guide-activity-subhead">Quick Actions</p>
                <div class="guide-activity-actions">
                  <a href="{_href("simulator")}" class="guide-activity-action">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M3 12h4l2-6 4 12 2-6h6"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-action-label">Start Simulation</span>
                    <span class="guide-activity-action-arrow" aria-hidden="true">→</span>
                  </a>
                  <a href="{_href("summary")}" class="guide-activity-action">
                    <span class="guide-activity-icon" aria-hidden="true">
                      <svg viewBox="0 0 24 24">
                        <path d="M7 4h7l5 5v11H7z"></path>
                        <path d="M14 4v5h5"></path>
                        <path d="M10 13h4"></path>
                        <path d="M10 17h4"></path>
                      </svg>
                    </span>
                    <span class="guide-activity-action-label">Generate Report</span>
                    <span class="guide-activity-action-arrow" aria-hidden="true">→</span>
                  </a>
                </div>
              </div>
            </aside>
          </div>
        </section>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="history-card" style="margin-top:1rem;">
          <strong style="font-family:Sora,sans-serif;">Spotlight Tip</strong>
          <p>
            To use <strong>Project Spotlight</strong>, press <strong>⌘ K</strong> on Mac or
            <strong>Ctrl K</strong> on Windows/Linux to instantly open the command search and run actions
            from anywhere in the platform.
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def _decision_hub_completion_label(days: float) -> str:
    projected_date = (datetime.now(tz=UTC) + timedelta(days=max(1, int(round(days))))).date()
    return projected_date.strftime("%b %d").replace(" 0", " ")


def _decision_hub_confidence_series(completion: Any) -> pd.DataFrame:
    values = sorted(float(value) for value in completion if pd.notna(value))
    if not values:
        return pd.DataFrame(
            [
                {"week": "W1", "confidence": 0.0, "threshold": 80.0},
                {"week": "W2", "confidence": 0.0, "threshold": 80.0},
            ]
        )

    max_weeks = max(8, min(30, int((max(values) + 4) // 5) + 2))
    rows: list[dict[str, float | str]] = []
    for week_index in range(1, max_weeks + 1):
        day_mark = float(week_index * 5)
        confidence = round(
            (sum(1 for value in values if value <= day_mark) / float(len(values))) * 100.0,
            1,
        )
        rows.append({"week": f"W{week_index}", "confidence": confidence, "threshold": 80.0})
    return pd.DataFrame(rows)


def _decision_hub_confidence_chart(series_df: pd.DataFrame) -> Any:
    figure = go.Figure()
    figure.add_trace(
        go.Scatter(
            x=series_df["week"],
            y=series_df["threshold"],
            mode="lines",
            name="80% target",
            line=dict(color="#f5c56b", width=2, dash="dash"),
            hovertemplate="%{x}<br>Target %{y:.0f}%<extra></extra>",
        )
    )
    figure.add_trace(
        go.Scatter(
            x=series_df["week"],
            y=series_df["confidence"],
            mode="lines",
            name="Confidence",
            line=dict(color="#5b6cff", width=3),
            fill="tozeroy",
            fillcolor="rgba(91, 108, 255, 0.26)",
            hovertemplate="%{x}<br>Confidence %{y:.1f}%<extra></extra>",
        )
    )
    figure.update_layout(
        height=260,
        margin=dict(l=18, r=16, t=20, b=20),
        showlegend=True,
    )
    figure.update_yaxes(range=[0, 100], ticksuffix="%", tick0=0, dtick=20)
    figure.update_xaxes(title_text="")
    return _theme_chart(figure)


def _decision_hub_gauge_figure(value: float, accent_color: str) -> Any:
    figure = go.Figure(
        go.Indicator(
            mode="gauge+number",
            value=max(0.0, min(100.0, value)),
            number={"suffix": "%", "font": {"size": 34, "family": "Sora, sans-serif"}},
            gauge={
                "axis": {
                    "range": [0, 100],
                    "tickvals": [0, 25, 50, 75, 100],
                    "tickfont": {"color": "#8fa5c5", "size": 10},
                },
                "bar": {"color": accent_color, "thickness": 0.34},
                "bgcolor": "rgba(255,255,255,0.03)",
                "borderwidth": 0,
                "steps": [
                    {"range": [0, 40], "color": "rgba(255,127,140,0.14)"},
                    {"range": [40, 70], "color": "rgba(245,197,107,0.13)"},
                    {"range": [70, 100], "color": "rgba(24,209,199,0.12)"},
                ],
            },
        )
    )
    figure.update_layout(
        height=205,
        margin=dict(l=16, r=16, t=10, b=4),
        paper_bgcolor="rgba(0,0,0,0)",
        font=dict(color="#f4f8ff", family="Manrope, sans-serif"),
    )
    return figure


def _render_decision_hub() -> None:
    _render_command_bar()
    payload = _current_payload()
    portfolio = payload["portfolio_summary"]
    monitoring = payload["monitoring_summary"]
    metrics = payload["metrics"]

    health_score = max(
        32.0,
        min(
            96.0,
            100.0
            - (metrics["delay_probability"] * 42.0)
            - (portfolio["delayed_pct"] * 0.22)
            - (monitoring["avg_drift"] * 18.0),
        ),
    )

    delay_probability_pct = metrics["delay_probability"] * 100.0
    confidence_score = _confidence_score(metrics)
    expected_completion = _decision_hub_completion_label(metrics["p50"])
    budget_overrun = float(portfolio["avg_budget_overrun"])
    budget_risk = "Low" if budget_overrun <= 5 else "Moderate" if budget_overrun <= 12 else "High"
    risk_driver_ranking = payload.get("drivers")
    if "risk_driver_ranking" not in locals() or risk_driver_ranking is None:
        risk_driver_ranking = []
    top_driver_name = (
        risk_driver_ranking[0].get("task_name", "Unknown driver")
        if isinstance(risk_driver_ranking, list)
        and len(risk_driver_ranking) > 0
        and isinstance(risk_driver_ranking[0], dict)
        else "Unknown driver"
    )
    alerts = [
        {
            "type": "critical",
            "msg": "Product scope and requirements leads the risk stack and appears on the critical path in 100% of simulations.",
        },
        {
            "type": "warning",
            "msg": "Monitoring drift averages 0.383; executive commitments should stay closer to the P80 forecast until variance narrows.",
        },
        {
            "type": "warning",
            "msg": "Portfolio delay rate is 66.1% with average forecast error of 15.2 days.",
        },
        {
            "type": "info",
            "msg": "Project health is 79/100 with a modeled deadline miss risk of 0.0%.",
        },
    ]

    insight_cards = [
        {
            "title": "Critical Path Analysis",
            "points": [
                f"{top_driver_name} is the top delay driver across this forecast run.",
                f"{len(payload['critical_path'])} tasks define the current longest path through the plan.",
                f"Baseline path duration is {payload['critical_path_days']:.1f} days before mitigation.",
            ],
        },
        {
            "title": "Resource Optimization",
            "points": [
                f"Portfolio delay rate sits at {portfolio['delayed_pct']:.1f}% across {int(portfolio['projects'])} historical projects.",
                f"Average budget overrun is {budget_overrun:.1f}% while forecast error averages {portfolio['avg_forecast_error']:.1f} days.",
                f"Monitoring accuracy is {monitoring['accuracy_pct']:.1f}% for recent risk snapshots.",
            ],
        },
        {
            "title": "Schedule Prediction",
            "points": [
                f"P50 completion lands around {expected_completion} ({metrics['p50']:.1f} days).",
                f"P80 commitment is {metrics['p80']:.1f} days with {delay_probability_pct:.1f}% miss risk.",
                f"Current project health holds at {health_score:.0f}/100 based on schedule, risk, and drift.",
            ],
        },
    ]

    st.markdown(
        """
        <style>
        .decision-hub-header {
          display: flex;
          align-items: flex-start;
          justify-content: space-between;
          gap: 1rem;
          margin-bottom: 1.5rem;
        }

        .decision-hub-header h1 {
          margin: 0;
          font-family: Sora, sans-serif;
          font-size: clamp(1.9rem, 2.6vw, 2.45rem);
          line-height: 1.02;
          color: #f4f8ff;
        }

        .decision-hub-header p {
          margin: 0.45rem 0 0;
          color: #92a4bf;
          font-size: 0.98rem;
        }

        .decision-status-chip {
          display: inline-flex;
          align-items: center;
          gap: 0.55rem;
          min-height: 44px;
          padding: 0.8rem 1rem;
          border-radius: 16px;
          border: 1px solid rgba(100, 162, 255, 0.2);
          background: rgba(6, 12, 24, 0.7);
          color: #8ef0af;
          font-size: 0.82rem;
          font-weight: 700;
          white-space: nowrap;
        }

        .decision-status-dot {
          width: 9px;
          height: 9px;
          border-radius: 999px;
          background: #27d37b;
          box-shadow: 0 0 0 0 rgba(39, 211, 123, 0.36);
          animation: decisionPulse 1.8s infinite;
        }

        .decision-hub-metric-grid {
          display: grid;
          grid-template-columns: repeat(4, minmax(0, 1fr));
          gap: 16px;
          margin-bottom: 1.1rem;
        }

        .decision-hub-metric-card {
          min-height: 138px;
          padding: 1.15rem 1.2rem;
          border-radius: 20px;
          border: 1px solid rgba(255, 255, 255, 0.08);
          background:
            linear-gradient(180deg, rgba(9, 18, 31, 0.95), rgba(7, 13, 24, 0.92)) padding-box,
            linear-gradient(135deg, rgba(155, 92, 255, 0.14), rgba(100, 162, 255, 0.14), rgba(24, 209, 199, 0.12)) border-box;
          box-shadow: var(--shadow);
          display: flex;
          flex-direction: column;
          justify-content: space-between;
        }

        .decision-hub-metric-card span {
          color: #95a7c4;
          font-size: 0.8rem;
          text-transform: uppercase;
          letter-spacing: 0.08em;
        }

        .decision-hub-metric-card strong {
          margin-top: 0.4rem;
          font-family: Sora, sans-serif;
          font-size: clamp(1.45rem, 1.9vw, 1.9rem);
          line-height: 1.05;
          color: #f4f8ff;
        }

        .decision-hub-metric-card small {
          color: #8ea2c0;
          font-size: 0.82rem;
        }

        .st-key-decision_hub_health,
        .st-key-decision_hub_main_chart,
        .st-key-decision_hub_gauge_success,
        .st-key-decision_hub_gauge_delay,
        .st-key-decision_hub_alerts,
        .st-key-decision_hub_insight_critical,
        .st-key-decision_hub_insight_resources,
        .st-key-decision_hub_insight_schedule {
          border: 1px solid rgba(255, 255, 255, 0.08);
          border-radius: 22px;
          padding: 1.2rem 1.25rem;
          background:
            linear-gradient(180deg, rgba(9, 18, 31, 0.95), rgba(7, 13, 24, 0.92)) padding-box,
            linear-gradient(135deg, rgba(155, 92, 255, 0.14), rgba(100, 162, 255, 0.14), rgba(24, 209, 199, 0.12)) border-box;
          box-shadow: var(--shadow);
        }

        .st-key-decision_hub_health {
          margin: 0.2rem 0 1.1rem;
        }

        .st-key-decision_hub_main_chart,
        .st-key-decision_hub_gauge_success,
        .st-key-decision_hub_gauge_delay {
          padding-bottom: 0.45rem;
        }

        .st-key-decision_hub_alerts,
        .st-key-decision_hub_insight_critical,
        .st-key-decision_hub_insight_resources,
        .st-key-decision_hub_insight_schedule {
          margin-top: 1rem;
        }

        .decision-health-card span,
        .decision-panel-label,
        .decision-insight-card span {
          color: #95a7c4;
          font-size: 0.8rem;
          text-transform: uppercase;
          letter-spacing: 0.08em;
        }

        .decision-health-card strong {
          display: block;
          margin: 0.55rem 0 0.35rem;
          font-family: Sora, sans-serif;
          font-size: clamp(2rem, 3vw, 2.6rem);
          color: #47e17b;
        }

        .decision-health-card p,
        .decision-insight-card p {
          margin: 0;
          color: #92a4bf;
          font-size: 0.92rem;
        }

        .decision-panel-title {
          margin: 0 0 0.35rem;
          font-family: Sora, sans-serif;
          font-size: 1.08rem;
          color: #f4f8ff;
        }

        .st-key-decision_alert_critical_0,
        .st-key-decision_alert_warning_1,
        .st-key-decision_alert_warning_2,
        .st-key-decision_alert_info_3 {
          border-radius: 16px;
          padding: 0.92rem 1rem;
          border: 1px solid rgba(255, 255, 255, 0.06);
          margin-top: 0.75rem;
        }

        .st-key-decision_alert_critical_0 {
          background: rgba(239, 68, 68, 0.05);
          border-color: rgba(239, 68, 68, 0.2);
        }

        .st-key-decision_alert_warning_1,
        .st-key-decision_alert_warning_2 {
          background: rgba(234, 179, 8, 0.05);
          border-color: rgba(234, 179, 8, 0.2);
        }

        .st-key-decision_alert_info_3 {
          background: rgba(59, 130, 246, 0.05);
          border-color: rgba(59, 130, 246, 0.2);
        }

        .st-key-decision_alert_critical_0 [data-testid="stMarkdownContainer"] p,
        .st-key-decision_alert_warning_1 [data-testid="stMarkdownContainer"] p,
        .st-key-decision_alert_warning_2 [data-testid="stMarkdownContainer"] p,
        .st-key-decision_alert_info_3 [data-testid="stMarkdownContainer"] p {
          margin: 0;
          color: #d9e5f7;
          font-size: 0.93rem;
          line-height: 1.65;
        }

        .st-key-decision_alert_critical_0 .stColumn:first-child [data-testid="stMarkdownContainer"] p {
          color: #f87171;
        }

        .st-key-decision_alert_warning_1 .stColumn:first-child [data-testid="stMarkdownContainer"] p,
        .st-key-decision_alert_warning_2 .stColumn:first-child [data-testid="stMarkdownContainer"] p {
          color: #facc15;
        }

        .st-key-decision_alert_info_3 .stColumn:first-child [data-testid="stMarkdownContainer"] p {
          color: #60a5fa;
        }

        .st-key-decision_alert_critical_0 .stColumn:first-child [data-testid="stMarkdownContainer"],
        .st-key-decision_alert_warning_1 .stColumn:first-child [data-testid="stMarkdownContainer"],
        .st-key-decision_alert_warning_2 .stColumn:first-child [data-testid="stMarkdownContainer"],
        .st-key-decision_alert_info_3 .stColumn:first-child [data-testid="stMarkdownContainer"] {
          text-align: center;
          font-size: 1rem;
          line-height: 1;
          padding-top: 0.15rem;
        }

        .decision-insights-title {
          margin: 1.25rem 0 0.15rem;
          font-family: Sora, sans-serif;
          font-size: 1.08rem;
          color: #f4f8ff;
        }

        .decision-insight-card h4 {
          margin: 0.45rem 0 0.75rem;
          font-family: Sora, sans-serif;
          font-size: 1rem;
          color: #f4f8ff;
        }

        .decision-insight-card ul {
          margin: 0;
          padding-left: 1rem;
          color: #d9e5f7;
          font-size: 0.93rem;
        }

        .decision-insight-card li + li {
          margin-top: 0.45rem;
        }

        .st-key-decision_hub_alerts div.stButton {
          display: flex;
          justify-content: flex-end;
        }

        .st-key-decision_hub_alerts div.stButton > button {
          width: auto !important;
          min-width: 164px;
          min-height: 50px;
          padding: 0 1.2rem !important;
          border: 1.5px solid transparent !important;
          border-radius: 17px !important;
          color: #f5fbff !important;
          font-weight: 800 !important;
          background:
            linear-gradient(180deg, rgba(5, 12, 24, 0.98), rgba(8, 18, 31, 0.96)) padding-box,
            linear-gradient(90deg, rgba(123, 78, 255, 0.98) 0%, rgba(79, 123, 255, 0.94) 54%, rgba(17, 225, 245, 0.98) 100%) border-box !important;
          box-shadow:
            0 14px 34px rgba(7, 15, 28, 0.34),
            0 0 0 1px rgba(151, 191, 255, 0.04) inset,
            0 0 22px rgba(74, 110, 255, 0.16) !important;
          transition: transform 160ms ease, box-shadow 160ms ease, filter 160ms ease !important;
        }

        .st-key-decision_hub_alerts div.stButton > button:hover {
          transform: translateY(-1px);
          box-shadow:
            0 16px 38px rgba(7, 15, 28, 0.38),
            0 0 0 1px rgba(173, 198, 255, 0.08) inset,
            0 0 28px rgba(74, 110, 255, 0.22) !important;
          filter: brightness(1.03);
        }

        .st-key-decision_hub_alerts .stCaptionContainer {
          padding-top: 0.2rem;
        }

        @keyframes decisionPulse {
          0% { box-shadow: 0 0 0 0 rgba(39, 211, 123, 0.36); }
          70% { box-shadow: 0 0 0 10px rgba(39, 211, 123, 0); }
          100% { box-shadow: 0 0 0 0 rgba(39, 211, 123, 0); }
        }

        @media (max-width: 1100px) {
          .decision-hub-metric-grid {
            grid-template-columns: repeat(2, minmax(0, 1fr));
          }
        }

        @media (max-width: 760px) {
          .decision-hub-header,
          .st-key-decision_hub_alerts .stHorizontalBlock {
            flex-direction: column;
            align-items: flex-start;
          }

          .decision-hub-metric-grid {
            grid-template-columns: minmax(0, 1fr);
          }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="page-shell">
          <div class="decision-hub-header">
            <div>
              <h1>Decision Hub</h1>
              <p>AI-powered project intelligence overview</p>
            </div>
            <div class="decision-status-chip">
              <span class="decision-status-dot"></span>
              <span>AI Monitoring Active</span>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <div class="page-shell">
          <div class="decision-hub-metric-grid">
            <div class="decision-hub-metric-card">
              <span>Expected Completion</span>
              <strong>{expected_completion}</strong>
              <small>P50 percentile</small>
            </div>
            <div class="decision-hub-metric-card">
              <span>Critical Path</span>
              <strong>{payload['critical_path_days']:.0f} days</strong>
              <small>{len(payload['critical_path'])} tasks on the path</small>
            </div>
            <div class="decision-hub-metric-card">
              <span>Delay Probability</span>
              <strong>{delay_probability_pct:.0f}%</strong>
              <small>Monte Carlo forecast exposure</small>
            </div>
            <div class="decision-hub-metric-card">
              <span>Budget Risk</span>
              <strong>{budget_risk}</strong>
              <small>{budget_overrun:.1f}% average budget variance</small>
            </div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.container(key="decision_hub_health"):
        st.markdown(
            f"""
            <div class="decision-health-card">
              <span>Project Health Score</span>
              <strong>{health_score:.0f} / 100</strong>
              <p>Based on schedule risk, portfolio drift, and resource stability.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    chart_config = {"displayModeBar": False, "responsive": True}
    chart_col, gauge_col = st.columns([1.9, 1.0], gap="large")
    with chart_col:
        with st.container(key="decision_hub_main_chart"):
            st.markdown(
                """
                <div class="decision-panel-label">Forecast Trend</div>
                <h3 class="decision-panel-title">Confidence Over Time</h3>
                """,
                unsafe_allow_html=True,
            )
            st.plotly_chart(
                _decision_hub_confidence_chart(_decision_hub_confidence_series(payload["completion"])),
                use_container_width=True,
                config=chart_config,
            )
    with gauge_col:
        with st.container(key="decision_hub_gauge_success"):
            st.markdown(
                """
                <div class="decision-panel-label">Confidence</div>
                <h3 class="decision-panel-title">Success Confidence</h3>
                """,
                unsafe_allow_html=True,
            )
            st.plotly_chart(
                _decision_hub_gauge_figure(confidence_score, "#5b6cff"),
                use_container_width=True,
                config=chart_config,
            )
        with st.container(key="decision_hub_gauge_delay"):
            st.markdown(
                """
                <div class="decision-panel-label">Risk</div>
                <h3 class="decision-panel-title">Delay Probability</h3>
                """,
                unsafe_allow_html=True,
            )
            st.plotly_chart(
                _decision_hub_gauge_figure(delay_probability_pct, "#22c7dd"),
                use_container_width=True,
                config=chart_config,
            )

    with st.container(key="decision_hub_alerts"):
        toolbar_col, action_col = st.columns([0.78, 0.22], gap="small")
        with toolbar_col:
            st.caption("Signals")
            st.markdown("### Predictive Alerts")
        with action_col:
            analysis_clicked = st.button(
                "Run AI Analysis",
                key="decision_hub_run_analysis",
                type="primary",
                use_container_width=True,
            )
        if analysis_clicked:
            with st.spinner("Analyzing live project signals..."):
                time.sleep(0.9)
            st.session_state.decision_hub_last_analysis = datetime.now(tz=UTC).strftime("%Y-%m-%d %H:%M UTC")
            st.rerun()
        if st.session_state.get("decision_hub_last_analysis"):
            st.caption(f"Last analysis refresh: {st.session_state.decision_hub_last_analysis}")

        alert_symbols = {"critical": "●", "warning": "●", "info": "●"}
        for index, alert in enumerate(alerts):
            with st.container(key=f"decision_alert_{alert['type']}_{index}"):
                marker_col, message_col = st.columns([0.06, 0.94], gap="small")
                with marker_col:
                    st.markdown(alert_symbols.get(alert["type"], "●"))
                with message_col:
                    st.markdown(alert["msg"])

    st.markdown(
        """
        <div class="page-shell">
          <h3 class="decision-insights-title">AI Insights</h3>
        </div>
        """,
        unsafe_allow_html=True,
    )

    insight_cols = st.columns(3, gap="large")
    insight_keys = [
        "decision_hub_insight_critical",
        "decision_hub_insight_resources",
        "decision_hub_insight_schedule",
    ]
    for column, card, container_key in zip(insight_cols, insight_cards, insight_keys, strict=False):
        with column:
            with st.container(key=container_key):
                points_markup = "".join(f"<li>{html.escape(point)}</li>" for point in card["points"])
                st.markdown(
                    f"""
                    <div class="decision-insight-card">
                      <span>AI Insight</span>
                      <h4>{html.escape(card['title'])}</h4>
                      <ul>{points_markup}</ul>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )


def _risk_intelligence_payload() -> dict[str, Any]:
    payload = st.session_state.get("task_architect_payload")
    if isinstance(payload, dict) and payload.get("tasks"):
        return payload
    return _current_payload()


def _risk_short_task_label(name: str) -> str:
    words = [word for word in re.split(r"[^A-Za-z0-9]+", str(name)) if word]
    filtered = [
        word for word in words if word.lower() not in {"and", "the", "for", "of", "to", "with", "a", "an"}
    ] or words
    if not filtered:
        return "Task"
    if len(filtered) == 1:
        return filtered[0][:12]
    short = " ".join(filtered[:2])
    if len(short) <= 14:
        return short
    acronym = "".join(word[0].upper() for word in filtered[:3])
    return acronym or short[:12]


def _risk_driver_rows(payload: dict[str, Any]) -> list[dict[str, Any]]:
    tasks = payload.get("tasks") or []
    task_map = {str(task.get("id")): task for task in tasks if isinstance(task, dict) and task.get("id")}
    rows: list[dict[str, Any]] = []

    drivers = payload.get("drivers") or []
    if isinstance(drivers, list) and drivers:
        for item in drivers:
            if not isinstance(item, dict):
                continue
            task_id = str(item.get("task_id", ""))
            task = task_map.get(task_id, {})
            frequency_pct = float(item.get("critical_path_frequency_pct", 0.0))
            risk_factor_pct = float(task.get("risk_factor", item.get("risk_factor", 0.0))) * 100.0
            score = round(min(100.0, (frequency_pct * 0.68) + (risk_factor_pct * 0.32)), 1)
            rows.append(
                {
                    "task_id": task_id or task.get("id", ""),
                    "task_name": str(item.get("task_name") or task.get("name") or "Unknown task"),
                    "critical_path_frequency_pct": frequency_pct,
                    "risk_factor_pct": round(risk_factor_pct, 1),
                    "score": score,
                    "mean_duration": float(task.get("mean_duration", 0.0)),
                    "std_dev": float(task.get("std_dev", 0.0)),
                }
            )

    if not rows:
        for task in tasks:
            if not isinstance(task, dict):
                continue
            risk_factor_pct = float(task.get("risk_factor", 0.0)) * 100.0
            rows.append(
                {
                    "task_id": str(task.get("id", "")),
                    "task_name": str(task.get("name", "Unknown task")),
                    "critical_path_frequency_pct": 0.0,
                    "risk_factor_pct": round(risk_factor_pct, 1),
                    "score": round(risk_factor_pct, 1),
                    "mean_duration": float(task.get("mean_duration", 0.0)),
                    "std_dev": float(task.get("std_dev", 0.0)),
                }
            )

    return sorted(
        rows,
        key=lambda row: (
            float(row.get("score", 0.0)),
            float(row.get("critical_path_frequency_pct", 0.0)),
            float(row.get("risk_factor_pct", 0.0)),
        ),
        reverse=True,
    )


def _risk_driver_ranking_figure(driver_rows: list[dict[str, Any]]) -> go.Figure:
    if not driver_rows:
        return go.Figure()

    chart_rows = list(reversed(driver_rows[:6]))
    colors = ["#1fd6dd", "#31bbe8", "#488df4", "#6d6fff", "#9358ff", "#b84cff"]

    figure = go.Figure(
        go.Bar(
            x=[row["score"] for row in chart_rows],
            y=[row["task_name"] for row in chart_rows],
            orientation="h",
            marker=dict(color=colors[: len(chart_rows)], line=dict(width=0)),
            customdata=[
                [row["critical_path_frequency_pct"], row["risk_factor_pct"], row["mean_duration"]]
                for row in chart_rows
            ],
            hovertemplate=(
                "<b>%{y}</b><br>"
                "Composite risk: %{x:.1f}<br>"
                "Critical-path frequency: %{customdata[0]:.1f}%<br>"
                "Risk factor: %{customdata[1]:.1f}%<br>"
                "Mean duration: %{customdata[2]:.1f} days<extra></extra>"
            ),
        )
    )
    max_score = max(float(row["score"]) for row in chart_rows)
    figure.update_layout(
        height=310,
        margin=dict(l=18, r=20, t=10, b=18),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(5,12,24,0.54)",
        font=dict(color="#dfe9f7", family="Manrope, sans-serif"),
        showlegend=False,
        bargap=0.22,
    )
    figure.update_xaxes(
        range=[0, max(100.0, max_score + 12.0)],
        showgrid=True,
        gridcolor="rgba(156, 176, 204, 0.10)",
        zeroline=False,
        color="#7f90ac",
        title_text="",
    )
    figure.update_yaxes(showgrid=False, color="#7f90ac", title_text="")
    return figure


def _critical_path_frequency_figure(driver_rows: list[dict[str, Any]]) -> go.Figure:
    if not driver_rows:
        return go.Figure()

    chart_rows = driver_rows[:5]
    colors = ["#b84cff", "#a859ff", "#7a6cff", "#4c93ff", "#23d1dd"]
    figure = go.Figure(
        go.Bar(
            x=[_risk_short_task_label(row["task_name"]) for row in chart_rows],
            y=[row["critical_path_frequency_pct"] for row in chart_rows],
            marker=dict(color=colors[: len(chart_rows)], line=dict(width=0)),
            hovertemplate="<b>%{x}</b><br>Critical-path frequency: %{y:.1f}%<extra></extra>",
        )
    )
    figure.update_layout(
        height=310,
        margin=dict(l=18, r=16, t=10, b=18),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(5,12,24,0.54)",
        font=dict(color="#dfe9f7", family="Manrope, sans-serif"),
        showlegend=False,
        bargap=0.2,
        barcornerradius=12,
    )
    figure.update_xaxes(showgrid=False, color="#7f90ac", title_text="")
    figure.update_yaxes(
        range=[0, 100],
        dtick=25,
        ticksuffix="%",
        showgrid=True,
        gridcolor="rgba(156, 176, 204, 0.10)",
        zeroline=False,
        color="#7f90ac",
        title_text="",
    )
    return figure


def _task_bottleneck_network_figure(payload: dict[str, Any], driver_rows: list[dict[str, Any]]) -> go.Figure:
    graph = payload.get("graph")
    if not isinstance(graph, nx.DiGraph) or graph.number_of_nodes() == 0:
        return go.Figure()

    levels: dict[str, int] = {}
    for node in nx.topological_sort(graph):
        predecessors = list(graph.predecessors(node))
        levels[node] = 0 if not predecessors else max(levels[pred] for pred in predecessors) + 1

    level_groups: dict[int, list[str]] = {}
    for node, level in levels.items():
        level_groups.setdefault(level, []).append(node)

    positions: dict[str, tuple[float, float]] = {}
    max_level = max(level_groups) if level_groups else 0
    for level, nodes in level_groups.items():
        total = len(nodes)
        if total == 1:
            y_positions = [0.0]
        else:
            step = 1.6 / max(total - 1, 1)
            y_positions = [0.8 - (index * step) for index in range(total)]
        x_position = -1.0 + ((2.0 * level) / max(max_level, 1))
        for node, y_position in zip(nodes, y_positions, strict=False):
            positions[node] = (x_position, y_position)

    critical_path = payload.get("critical_path") or []
    critical_edges = set(zip(critical_path, critical_path[1:], strict=False))
    driver_lookup = {row["task_id"]: row for row in driver_rows}

    normal_edge_x: list[float | None] = []
    normal_edge_y: list[float | None] = []
    critical_edge_x: list[float | None] = []
    critical_edge_y: list[float | None] = []

    for source, target in graph.edges():
        x0, y0 = positions[source]
        x1, y1 = positions[target]
        bucket_x = critical_edge_x if (source, target) in critical_edges else normal_edge_x
        bucket_y = critical_edge_y if (source, target) in critical_edges else normal_edge_y
        bucket_x.extend([x0, x1, None])
        bucket_y.extend([y0, y1, None])

    node_x = [positions[node][0] for node in graph.nodes()]
    node_y = [positions[node][1] for node in graph.nodes()]
    node_sizes: list[float] = []
    node_colors: list[str] = []
    node_line_colors: list[str] = []
    node_labels: list[str] = []
    hover_rows: list[list[Any]] = []

    for node in graph.nodes():
        task_name = str(graph.nodes[node].get("name", node))
        task_duration = float(graph.nodes[node].get("mean_duration", 0.0))
        risk_factor = float(graph.nodes[node].get("risk_factor", 0.0)) * 100.0
        frequency_pct = float(driver_lookup.get(node, {}).get("critical_path_frequency_pct", 0.0))
        node_sizes.append(26 + (risk_factor * 0.18) + (frequency_pct * 0.10))
        if node in critical_path:
            node_colors.append("rgba(157, 77, 255, 0.72)")
            node_line_colors.append("#ff4f97")
        else:
            node_colors.append("rgba(24, 138, 255, 0.40)")
            node_line_colors.append("#24d4ff")
        node_labels.append(_risk_short_task_label(task_name))
        hover_rows.append([task_name, risk_factor, frequency_pct, task_duration])

    figure = go.Figure()
    figure.add_trace(
        go.Scatter(
            x=normal_edge_x,
            y=normal_edge_y,
            mode="lines",
            line=dict(color="rgba(138, 154, 180, 0.20)", width=1.4, dash="dot"),
            hoverinfo="none",
            showlegend=False,
        )
    )
    figure.add_trace(
        go.Scatter(
            x=critical_edge_x,
            y=critical_edge_y,
            mode="lines",
            line=dict(color="rgba(187, 203, 226, 0.76)", width=2.6),
            hoverinfo="none",
            showlegend=False,
        )
    )
    figure.add_trace(
        go.Scatter(
            x=node_x,
            y=node_y,
            mode="markers",
            marker=dict(size=[size * 1.35 for size in node_sizes], color="rgba(135, 86, 255, 0.14)", line=dict(width=0)),
            hoverinfo="none",
            showlegend=False,
        )
    )
    figure.add_trace(
        go.Scatter(
            x=node_x,
            y=node_y,
            mode="markers+text",
            text=node_labels,
            textposition="middle center",
            textfont=dict(color="#eef5ff", size=11, family="Sora, sans-serif"),
            marker=dict(size=node_sizes, color=node_colors, line=dict(color=node_line_colors, width=1.8)),
            customdata=hover_rows,
            hovertemplate=(
                "<b>%{customdata[0]}</b><br>"
                "Risk factor: %{customdata[1]:.1f}%<br>"
                "Critical-path frequency: %{customdata[2]:.1f}%<br>"
                "Mean duration: %{customdata[3]:.1f} days<extra></extra>"
            ),
            showlegend=False,
        )
    )
    figure.update_layout(
        height=360,
        margin=dict(l=8, r=8, t=8, b=6),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        font=dict(color="#dfe9f7", family="Manrope, sans-serif"),
    )
    return figure


def _risk_reasoning_items(payload: dict[str, Any], driver_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    graph = payload.get("graph")
    tasks = payload.get("tasks") or []
    if not isinstance(graph, nx.DiGraph) or not tasks:
        return [
            {
                "title": "Risk drivers will appear after a plan is generated",
                "body": "Generate a task structure from AI Task Architect to see the leading bottlenecks, critical-path pressure, and mitigation guidance.",
            }
        ]

    task_map = {str(task.get("id")): task for task in tasks if isinstance(task, dict) and task.get("id")}
    reasoning: list[dict[str, str]] = []
    used_ids: set[str] = set()

    if driver_rows:
        top_row = driver_rows[0]
        top_id = str(top_row.get("task_id", ""))
        top_task = task_map.get(top_id, {})
        mean_duration = float(top_task.get("mean_duration", top_row.get("mean_duration", 0.0)))
        std_dev = float(top_task.get("std_dev", top_row.get("std_dev", 0.0)))
        variance_pct = (std_dev / mean_duration * 100.0) if mean_duration else 0.0
        downstream_tasks = len(nx.descendants(graph, top_id)) if top_id and graph.has_node(top_id) else 0
        reasoning.append(
            {
                "title": f"{top_row['task_name']} is the primary risk driver",
                "body": (
                    f"This task appears on the critical path in {top_row['critical_path_frequency_pct']:.0f}% of simulations. "
                    f"It fans into {downstream_tasks} downstream tasks and carries duration variance of {variance_pct:.0f}%."
                ),
            }
        )
        used_ids.add(top_id)

    connectivity_candidates: list[tuple[str, int, int]] = []
    for node in graph.nodes():
        if str(node) in used_ids:
            continue
        downstream_tasks = len(nx.descendants(graph, node))
        connectivity = int(graph.in_degree(node) + graph.out_degree(node))
        connectivity_candidates.append((str(node), connectivity, downstream_tasks))
    connectivity_candidates.sort(key=lambda item: (item[1], item[2]), reverse=True)
    if connectivity_candidates:
        node_id, _, downstream_tasks = connectivity_candidates[0]
        task = task_map.get(node_id, {})
        reasoning.append(
            {
                "title": f"{task.get('name', node_id)} creates the widest coordination bottleneck",
                "body": (
                    f"It connects {graph.in_degree(node_id)} upstream and {graph.out_degree(node_id)} downstream dependencies, "
                    f"so slippage here can cascade into {downstream_tasks} later tasks."
                ),
            }
        )
        used_ids.add(node_id)

    variance_candidates: list[tuple[str, float]] = []
    for task in tasks:
        task_id = str(task.get("id", ""))
        if task_id in used_ids:
            continue
        mean_duration = float(task.get("mean_duration", 0.0))
        std_dev = float(task.get("std_dev", 0.0))
        variance_pct = (std_dev / mean_duration * 100.0) if mean_duration else 0.0
        variance_candidates.append((task_id, variance_pct))
    variance_candidates.sort(key=lambda item: item[1], reverse=True)
    if variance_candidates:
        task_id, variance_pct = variance_candidates[0]
        task = task_map.get(task_id, {})
        reasoning.append(
            {
                "title": f"{task.get('name', task_id)} carries the widest duration uncertainty",
                "body": (
                    f"This step is estimated at {float(task.get('mean_duration', 0.0)):.1f} days with ±{float(task.get('std_dev', 0.0)):.1f} days "
                    f"of variability ({variance_pct:.0f}% spread), making it a major contributor to the P80 commitment."
                ),
            }
        )

    return reasoning[:3]


def _risk_prediction_items(payload: dict[str, Any], driver_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    tasks = payload.get("tasks") or []
    task_map = {str(task.get("id")): task for task in tasks if isinstance(task, dict) and task.get("id")}
    baseline_row = None
    boosted_row = None
    scenarios_df = payload.get("scenarios_df")
    if isinstance(scenarios_df, pd.DataFrame) and not scenarios_df.empty:
        baseline_match = scenarios_df[scenarios_df["Scenario"] == "Baseline"]
        boosted_match = scenarios_df[scenarios_df["Scenario"].str.contains("Increased Capacity", case=False, na=False)]
        baseline_row = baseline_match.iloc[0] if not baseline_match.empty else None
        boosted_row = boosted_match.iloc[0] if not boosted_match.empty else None

    if driver_rows:
        top_row = driver_rows[0]
        top_task = task_map.get(str(top_row.get("task_id", "")), {})
        emerging_body = (
            f"{top_row['task_name']} currently carries the highest combined risk score and shows up on the critical path in "
            f"{top_row['critical_path_frequency_pct']:.0f}% of simulations. "
            f"Its modeled variability is ±{float(top_task.get('std_dev', top_row.get('std_dev', 0.0))):.1f} days."
        )
        emerging_title = f"Emerging Risk: {top_row['task_name']}"
    else:
        emerging_title = "Emerging Risk"
        emerging_body = "Generate a project plan to see which tasks rise to the top of the risk stack."

    if baseline_row is not None and boosted_row is not None:
        mitigation_body = (
            f"If you implement the modeled capacity intervention, overall delay probability drops from "
            f"{float(baseline_row['Delay Prob (%)']):.1f}% to {float(boosted_row['Delay Prob (%)']):.1f}%, "
            f"and the P80 commitment improves from {float(baseline_row['P80 (days)']):.1f} to {float(boosted_row['P80 (days)']):.1f} days."
        )
    else:
        metrics = payload.get("metrics") or {}
        mitigation_body = (
            f"Current forecast confidence sits around P80 {float(metrics.get('p80', 0.0)):.1f} days. "
            "Run scenario analysis after generating a plan to compare mitigation options."
        )

    return [
        {"title": emerging_title, "body": emerging_body},
        {"title": "Mitigation Impact Forecast", "body": mitigation_body},
    ]


def _render_risk_intelligence() -> None:
    _render_command_bar()
    payload = _risk_intelligence_payload()
    driver_rows = _risk_driver_rows(payload)
    reasoning_items = _risk_reasoning_items(payload, driver_rows)
    prediction_items = _risk_prediction_items(payload, driver_rows)
    chart_config = {"displayModeBar": False, "responsive": True}

    st.markdown(
        """
        <style>
        .risk-intel-header {
          margin: 0 0 1.45rem;
        }

        .risk-intel-header h1 {
          margin: 0;
          font-family: Sora, sans-serif;
          font-size: clamp(2rem, 2.8vw, 2.7rem);
          line-height: 1.02;
          color: #f4f8ff;
        }

        .risk-intel-header p {
          margin: 0.45rem 0 0;
          color: #92a4bf;
          font-size: 0.98rem;
        }

        .st-key-risk_intel_driver_card,
        .st-key-risk_intel_frequency_card,
        .st-key-risk_intel_network_card,
        .st-key-risk_intel_reasoning_card,
        .st-key-risk_prediction_0,
        .st-key-risk_prediction_1 {
          border: 1px solid rgba(255, 255, 255, 0.08);
          border-radius: 22px;
          padding: 1.18rem 1.25rem;
          background:
            linear-gradient(180deg, rgba(9, 18, 31, 0.95), rgba(7, 13, 24, 0.92)) padding-box,
            linear-gradient(135deg, rgba(155, 92, 255, 0.14), rgba(100, 162, 255, 0.14), rgba(24, 209, 199, 0.12)) border-box;
          box-shadow: var(--shadow);
        }

        .st-key-risk_intel_network_card,
        .st-key-risk_intel_reasoning_card,
        .st-key-risk_prediction_0,
        .st-key-risk_prediction_1 {
          margin-top: 1rem;
        }

        .risk-intel-card-kicker,
        .risk-intel-section-kicker {
          color: #95a7c4;
          font-size: 0.8rem;
          text-transform: uppercase;
          letter-spacing: 0.14em;
          font-weight: 700;
        }

        .risk-intel-section-heading {
          margin: 1rem 0 0;
          padding-left: 1.25rem;
        }

        .risk-intel-card-title {
          margin: 0.3rem 0 0.1rem;
          font-family: Sora, sans-serif;
          font-size: 1.08rem;
          color: #f4f8ff;
        }

        .risk-intel-reasoning-item {
          padding: 0.95rem 1rem;
          border-radius: 16px;
          background: rgba(255, 255, 255, 0.03);
          border: 1px solid rgba(255, 255, 255, 0.04);
        }

        .risk-intel-reasoning-item + .risk-intel-reasoning-item {
          margin-top: 0.85rem;
        }

        .risk-intel-reasoning-item strong {
          display: block;
          margin: 0 0 0.45rem;
          color: #f4f8ff;
          font-size: 0.98rem;
          font-family: Sora, sans-serif;
        }

        .risk-intel-reasoning-item p,
        .risk-intel-prediction-body {
          margin: 0;
          color: #9db0ca;
          line-height: 1.72;
          font-size: 0.95rem;
        }

        .risk-intel-prediction-head {
          display: flex;
          align-items: center;
          gap: 0.6rem;
          color: #95a7c4;
          text-transform: uppercase;
          letter-spacing: 0.12em;
          font-size: 0.78rem;
          font-weight: 700;
        }

        .risk-intel-prediction-dot {
          width: 9px;
          height: 9px;
          border-radius: 999px;
          background: linear-gradient(135deg, #9b5cff 0%, #18d1c7 100%);
          box-shadow: 0 0 16px rgba(91, 108, 255, 0.28);
          flex: 0 0 auto;
        }

        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        """
        <div class="page-shell">
          <div class="risk-intel-header">
            <h1>Risk Intelligence</h1>
            <p>Explainable AI insights into project risks</p>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    ranking_col, frequency_col = st.columns(2, gap="large")
    with ranking_col:
        with st.container(key="risk_intel_driver_card"):
            st.markdown(
                """
                <div class="risk-intel-card-kicker">Risk Driver Ranking</div>
                <h3 class="risk-intel-card-title">Tasks with the highest modeled pressure</h3>
                """,
                unsafe_allow_html=True,
            )
            st.plotly_chart(_risk_driver_ranking_figure(driver_rows), use_container_width=True, config=chart_config)
    with frequency_col:
        with st.container(key="risk_intel_frequency_card"):
            st.markdown(
                """
                <div class="risk-intel-card-kicker">Critical Path Frequency</div>
                <h3 class="risk-intel-card-title">Which tasks dominate the delivery path</h3>
                """,
                unsafe_allow_html=True,
            )
            st.plotly_chart(_critical_path_frequency_figure(driver_rows), use_container_width=True, config=chart_config)

    with st.container(key="risk_intel_network_card"):
        st.markdown(
            """
            <div class="risk-intel-card-kicker">Task Bottleneck Network</div>
            <h3 class="risk-intel-card-title">How risk propagates through dependencies</h3>
            """,
            unsafe_allow_html=True,
        )
        st.plotly_chart(_task_bottleneck_network_figure(payload, driver_rows), use_container_width=True, config=chart_config)

    with st.container(key="risk_intel_reasoning_card"):
        st.markdown(
            """
            <div class="risk-intel-card-kicker">AI Reasoning — Why Delays May Occur</div>
            <h3 class="risk-intel-card-title">What the model sees in the generated plan</h3>
            """,
            unsafe_allow_html=True,
        )
        for item in reasoning_items:
            st.markdown(
                f"""
                <div class="risk-intel-reasoning-item">
                  <strong>{html.escape(item['title'])}</strong>
                  <p>{html.escape(item['body'])}</p>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown(
        """
        <div class="risk-intel-section-heading">
          <div class="risk-intel-section-kicker">AI Risk Predictions</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    for index, item in enumerate(prediction_items):
        with st.container(key=f"risk_prediction_{index}"):
            st.markdown(
                f"""
                <div class="risk-intel-prediction-head">
                  <span class="risk-intel-prediction-dot" aria-hidden="true"></span>
                  <span>{html.escape(item['title'].upper())}</span>
                </div>
                <p class="risk-intel-prediction-body">{html.escape(item['body'])}</p>
                """,
                unsafe_allow_html=True,
            )


def _render_structured_workspace(kicker: str) -> None:
    _render_command_bar()
    payload = _current_payload()
    st.markdown(
        f"""
        <div class="page-shell">
          <div class="glass dashboard-banner">
            <span class="kicker">{kicker}</span>
            <h1 class="page-title">Inspect the generated work structure and dependency graph.</h1>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    _render_structured_workspace_content(payload)


def _render_structured_workspace_content(
    payload: dict[str, Any],
    *,
    notes_title: str = "Architect notes",
    notes_copy: str | None = None,
) -> None:
    graph_col, task_col = st.columns([1.0, 1.0], gap="large")
    with graph_col:
        st.plotly_chart(_theme_chart(dependency_graph_figure(payload["graph"], payload["critical_path"])), use_container_width=True)
    with task_col:
        tasks_df = pd.DataFrame(payload["tasks"])
        st.dataframe(tasks_df, use_container_width=True, hide_index=True)
        default_copy = (
            f"{len(payload['tasks'])} tasks were generated for this workspace. "
            f"The baseline critical path spans {payload['critical_path_days']:.1f} days."
        )
        st.markdown(
            f"""
            <div class="history-card">
              <strong style="font-family:Sora,sans-serif;">{notes_title}</strong>
              <p>{notes_copy or default_copy}</p>
            </div>
            """,
            unsafe_allow_html=True,
        )


def _render_task_architect() -> None:
    st.markdown(
        """
        <style>
        div[data-testid="stForm"] {
          border: 1px solid transparent;
          border-radius: 22px;
          padding: 1.35rem 1.5rem 1.5rem;
          background:
            linear-gradient(180deg, rgba(10, 18, 30, 0.95), rgba(6, 11, 22, 0.92)) padding-box,
            linear-gradient(135deg, rgba(102, 176, 255, 0.4), rgba(38, 214, 231, 0.22), rgba(155, 92, 255, 0.24)) border-box;
          box-shadow: var(--shadow);
        }

        div[data-testid="stForm"] div[data-testid="stFormSubmitButton"] {
          margin-top: 0.35rem;
        }

        div[data-testid="stForm"] div[data-testid="stFormSubmitButton"] button {
          width: auto !important;
          min-width: 178px;
          padding-left: 1.2rem !important;
          padding-right: 1.2rem !important;
        }

        div[data-testid="stForm"] [data-testid="stTextInputRootElement"],
        div[data-testid="stForm"] div[data-testid="stNumberInput"] > div,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] > div > div,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"],
        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"] > div {
          border: 1.5px solid transparent !important;
          border-radius: 15px !important;
          background-color: transparent !important;
          background:
            linear-gradient(180deg, rgba(7, 14, 28, 0.99), rgba(6, 12, 24, 0.97)) padding-box,
            linear-gradient(90deg, rgba(123, 78, 255, 0.98) 0%, rgba(79, 123, 255, 0.94) 54%, rgba(17, 225, 245, 0.98) 100%) border-box !important;
          box-shadow:
            0 0 0 1px rgba(68, 157, 255, 0.08),
            0 0 26px rgba(17, 225, 245, 0.08),
            0 10px 24px rgba(0, 0, 0, 0.14) !important;
        }

        div[data-testid="stForm"] [data-testid="stTextInputRootElement"]:focus-within,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] > div:focus-within,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] > div > div:focus-within,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"]:focus-within,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"] > div:focus-within {
          box-shadow:
            0 0 0 1px rgba(17, 225, 245, 0.24),
            0 0 26px rgba(79, 123, 255, 0.14) !important;
        }

        div[data-testid="stForm"] [data-testid="stTextInputRootElement"] input,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] input {
          min-height: 50px !important;
          color: #f4f8ff !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"],
        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-baseweb="input"] > div {
<<<<<<< ours
<<<<<<< ours
          overflow: hidden !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] button {
          min-width: 29px !important;
=======
=======
>>>>>>> theirs
          flex: 1 1 auto !important;
          min-width: 0 !important;
          overflow: hidden !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-testid="stNumberInputContainer"] {
          display: flex !important;
          align-items: stretch !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-testid="stNumberInputContainer"] > div:first-child {
          flex: 1 1 auto !important;
          min-width: 0 !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] [data-testid="stNumberInputContainer"] > div:last-child {
          display: flex !important;
          flex-direction: row !important;
          flex: 0 0 auto !important;
          width: 64px !important;
          min-width: 64px !important;
          height: 50px !important;
          align-self: stretch !important;
          justify-content: stretch !important;
          position: relative !important;
          z-index: 2 !important;
          visibility: visible !important;
          opacity: 1 !important;
          pointer-events: auto !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"],
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"] {
          width: 32px !important;
          min-width: 32px !important;
          height: 50px !important;
>>>>>>> theirs
          min-height: 50px !important;
          padding: 0 !important;
          border: none !important;
          border-left: 1px solid rgba(82, 215, 241, 0.22) !important;
          border-radius: 0 !important;
          background: transparent !important;
          color: rgba(118, 204, 235, 0.9) !important;
          box-shadow: none !important;
        }

<<<<<<< ours
<<<<<<< ours
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button:hover {
=======
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"]:hover,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"]:hover {
>>>>>>> theirs
=======
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"]:hover,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"]:hover {
>>>>>>> theirs
          background: rgba(83, 215, 241, 0.06) !important;
          color: rgba(162, 232, 255, 0.96) !important;
        }

<<<<<<< ours
<<<<<<< ours
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button svg {
          fill: currentColor !important;
=======
=======
>>>>>>> theirs
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"] > *,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"] > * {
          width: 100% !important;
          height: 100% !important;
          display: flex !important;
          align-items: center !important;
          justify-content: center !important;
          margin: 0 !important;
          padding: 0 !important;
          line-height: 1 !important;
          transform: none !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"] > * > *,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"] > * > * {
          display: flex !important;
          align-items: center !important;
          justify-content: center !important;
          margin: 0 !important;
          padding: 0 !important;
          line-height: 1 !important;
          transform: none !important;
        }

        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepDown"] svg,
        div[data-testid="stForm"] div[data-testid="stNumberInput"] button[data-testid="stNumberInputStepUp"] svg {
          display: block !important;
          width: 14px !important;
          height: 14px !important;
          margin: 0 !important;
          padding: 0 !important;
          line-height: 1 !important;
          transform: none !important;
          position: static !important;
>>>>>>> theirs
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    _render_command_bar()

    st.markdown(
        """
        <div class="page-shell task-architect-shell">
          <div class="task-architect-hero">
            <span class="task-architect-hero-kicker">AI Task Architect</span>
            <h1 class="task-architect-hero-title">Inspect the generated work structure and dependency graph.</h1>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    uploaded_project_file, upload_generate_clicked = _render_upload_intake_section(
        upload_key="task_architect_upload_file",
        show_generate_button=True,
        generate_button_key="task_architect_upload_generate",
    )

    submission_error: str | None = None
    with st.form("task_architect_form", clear_on_submit=False):
        st.markdown(
            """
            <div class="task-architect-panel-head">
              <svg viewBox="0 0 24 24" aria-hidden="true">
                <path d="M10 5a3 3 0 1 0-6 0c0 1.6 1.1 2.8 2.4 3.4"></path>
                <path d="M14 5a3 3 0 1 1 6 0c0 1.6-1.1 2.8-2.4 3.4"></path>
                <path d="M7 13a3 3 0 1 0 0 6"></path>
                <path d="M17 13a3 3 0 1 1 0 6"></path>
                <path d="M7 8h10"></path>
                <path d="M12 8v8"></path>
              </svg>
              <span>Project Parameters</span>
            </div>
            """,
            unsafe_allow_html=True,
        )
        goal_col, timeline_col, team_col = st.columns(3, gap="large")
        with goal_col:
            project_goal = st.text_input(
                "Project Goal",
                key="task_architect_goal",
                placeholder="e.g., Build a customer portal",
            )
        with timeline_col:
            timeline_text = st.number_input(
                "Timeline (week)",
                min_value=1,
                max_value=104,
                step=1,
                key="task_architect_timeline",
            )
        with team_col:
            team_size = st.number_input(
                "Team Size",
                min_value=1,
                max_value=50,
                step=1,
                key="task_architect_team_size",
            )
        submitted = st.form_submit_button("Generate AI Plan")

    if submitted:
        submission_error = _run_task_architect_generation(project_goal, timeline_text, int(team_size))
    elif upload_generate_clicked:
        submission_error = _run_task_architect_upload_generation(uploaded_project_file)

    if submission_error:
        st.error(submission_error)

    generated_payload = st.session_state.get("task_architect_payload")
    if generated_payload:
        meta = generated_payload.get("architect_meta", {})
        goal_label = html.escape(meta.get("project_goal", "your project"))
        timeline_label = html.escape(meta.get("timeline_text", ""))
        team_size_value = int(meta.get("team_size", 0))
        notes_copy = (
            f"Generated an AI plan for <strong>{goal_label}</strong> with a target of "
            f"{timeline_label} and a team of {team_size_value}. "
            f"The baseline critical path spans {generated_payload['critical_path_days']:.1f} days."
        )
        _render_structured_workspace_content(
            generated_payload,
            notes_title="Generated plan",
            notes_copy=notes_copy,
        )
    else:
        st.markdown(
            """
            <div class="glass task-architect-empty">
              <div class="task-architect-empty-inner">
                <span class="task-architect-empty-icon" aria-hidden="true">
                  <svg viewBox="0 0 24 24">
                    <path d="M9 7a3 3 0 0 1 6 0c1.8 0 3 1.3 3 3 0 1.3-.7 2.4-1.9 2.9"></path>
                    <path d="M9 7c-1.8 0-3 1.3-3 3 0 1.3.7 2.4 1.9 2.9"></path>
                    <path d="M8.5 12.9c0 1.8 1.2 3.1 3 3.1"></path>
                    <path d="M15.5 12.9c0 1.8-1.2 3.1-3 3.1"></path>
                    <path d="M12 16v3"></path>
                    <path d="M10 21h4"></path>
                  </svg>
                </span>
                <h2 class="task-architect-empty-title">Ready to Architect ?</h2>
              </div>
            </div>
            """,
            unsafe_allow_html=True,
        )

def _render_simulator() -> None:
    _render_command_bar()
    st.markdown(
        """
        <style>
        .st-key-sim_run_button div.stButton > button,
        .st-key-sim_run_button button {
          width: 100% !important;
          justify-content: center !important;
          min-height: 52px !important;
          padding: 0 1rem !important;
          border: 1.5px solid transparent !important;
          border-radius: 16px !important;
          color: #f5fbff !important;
          background-color: transparent !important;
          background:
            linear-gradient(180deg, #030811, #02060d) padding-box,
            linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
          box-shadow: none !important;
          font-weight: 700 !important;
        }

        .st-key-sim_run_button div.stButton > button:hover,
        .st-key-sim_run_button button:hover {
          transform: translateY(-1px);
          filter: brightness(1.03);
          box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.16) !important;
        }

        .st-key-sim_run_button div.stButton > button[kind="primary"],
        .st-key-sim_run_button button[kind="primary"] {
          background-color: transparent !important;
          background:
            linear-gradient(180deg, #030811, #02060d) padding-box,
            linear-gradient(90deg, var(--purple-strong) 0%, var(--blue-strong) 52%, var(--teal) 100%) border-box !important;
          box-shadow: 0 0 0 1px rgba(117, 231, 255, 0.22) !important;
        }

        div[data-testid="stSlider"] {
          margin-bottom: 0.4rem;
        }

        div[data-testid="stSlider"] label,
        div[data-testid="stSlider"] label p {
          color: #f5fbff !important;
          font-family: "Sora", sans-serif !important;
          font-weight: 700 !important;
          letter-spacing: -0.02em !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    _render_monte_carlo_explainer()
    _render_monte_carlo_lens()
    payload = _current_payload()
    variance = st.slider(
        "Schedule uncertainty level",
        min_value=0.5,
        max_value=3.0,
        value=1.0,
        step=0.1,
        key="sim_variance",
        help="Higher values widen task-duration variability and create a broader forecast range.",
    )
    runs = st.slider(
        "Simulation runs",
        min_value=1000,
        max_value=10000,
        value=min(3000, int(payload["iterations"])),
        step=500,
        key="sim_runs",
        help="More runs create a smoother Monte Carlo distribution and more stable forecast metrics.",
    )
    if st.button("Run AI Simulation", key="sim_run_button", type="primary", use_container_width=True):
        tasks = copy.deepcopy(payload["tasks"])
        for task in tasks:
            task["std_dev"] = float(task["std_dev"]) * float(variance)
        graph = build_project_graph(tasks)
        completion = run_monte_carlo(graph, iterations=int(runs), seed=payload["seed"])
        sim_metrics = compute_metrics(completion, payload["deadline_days"])
        hist = completion_histogram(completion, payload["deadline_days"], sim_metrics["p50"], sim_metrics["p80"])
        st.session_state["simulator_result"] = {"completion": completion, "metrics": sim_metrics, "figure": hist}

    result = st.session_state.get("simulator_result")
    if result:
        cols = st.columns(3)
        cols[0].metric("P50 Timeline", f"{result['metrics']['p50']:.1f}d")
        cols[1].metric("P80 Timeline", f"{result['metrics']['p80']:.1f}d")
        cols[2].metric("Delay Probability", f"{result['metrics']['delay_probability'] * 100:.1f}%")
        st.plotly_chart(_theme_chart(result["figure"]), use_container_width=True)
        st.progress(min(1.0, runs / 10000))
        st.caption("Sampling distributions → Computing paths → Scoring risks")


def _render_what_if() -> None:
    _render_command_bar()
    payload = _current_payload()
    capacity_factor = st.slider("Capacity factor", min_value=0.65, max_value=1.0, value=0.85, step=0.05, key="scenario_capacity")
    deadline_factor = st.slider("Aggressive deadline factor", min_value=0.65, max_value=1.0, value=0.85, step=0.05, key="scenario_deadline")
    scenarios_df = scenario_comparison(
        build_project_graph,
        payload["tasks"],
        deadline_days=payload["deadline_days"],
        iterations=payload["iterations"],
        seed=payload["seed"],
        capacity_factor=float(capacity_factor),
        aggressive_deadline_factor=float(deadline_factor),
    )
    st.plotly_chart(_theme_chart(scenario_comparison_chart(scenarios_df)), use_container_width=True)
    st.dataframe(scenarios_df, use_container_width=True, hide_index=True)


def _render_history() -> None:
    _render_command_bar()
    _initialize_database()
    recent_runs = list_recent_runs(limit=12)
    st.markdown(
        """
        <div class="page-shell">
          <div class="glass dashboard-banner">
            <span class="kicker">Version History</span>
            <h1 class="page-title">Compare recent planning runs and retrieve saved workspaces.</h1>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if not recent_runs:
        st.info("No saved runs yet. Generate a forecast from the dashboard first.")
        return

    labels = [
        f"Run {run['simulation_run_id']} · {run['created_at']} · {run['delay_probability'] * 100:.1f}% risk"
        for run in recent_runs
    ]
    selected_label = st.selectbox("Saved workspace", labels, key="history_selected")
    selected_index = labels.index(selected_label)
    selected_run = recent_runs[selected_index]
    details = get_run_details(selected_run["simulation_run_id"])

    if details:
        cols = st.columns(4)
        cols[0].metric("Delay Probability", f"{details['metrics'].get('delay_probability', 0.0) * 100:.1f}%")
        cols[1].metric("P80", f"{details['metrics'].get('p80', 0.0):.1f}d")
        cols[2].metric("Iterations", f"{details['iterations']}")
        cols[3].metric("Mode", details["mode"])
        st.dataframe(pd.DataFrame(details["tasks"]), use_container_width=True, hide_index=True)


def _render_summary() -> None:
    _render_command_bar()
    payload = _current_payload()
    model_metrics = payload["model_metrics"]
    summary_cols = st.columns(4)
    summary_cols[0].metric("Projects in history", f"{int(payload['portfolio_summary']['projects'])}")
    summary_cols[1].metric("Portfolio delay rate", f"{payload['portfolio_summary']['delayed_pct']:.1f}%")
    summary_cols[2].metric("Monitoring accuracy", f"{payload['monitoring_summary']['accuracy_pct']:.1f}%")
    summary_cols[3].metric("Selected model", model_metrics.get("selected_model_name", "unknown"))

    st.markdown(
        f"""
        <div class="history-card">
          <strong style="font-family:Sora,sans-serif;">AI narrative summary</strong>
          <p>
            The current workspace combines live Monte Carlo forecasting with a deployed advisory classifier.
            Portfolio history shows a {payload['portfolio_summary']['delayed_pct']:.1f}% delayed-project rate,
            while monitoring accuracy is {payload['monitoring_summary']['accuracy_pct']:.1f}% across recent snapshots.
            For the active plan, the recommended commitment is {payload['metrics']['p80']:.1f} days and the main driver is
            <strong>{payload['drivers'][0]['task_name'] if payload['drivers'] else 'not available'}</strong>.
          </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    charts = st.columns(2, gap="large")
    with charts[0]:
        cost_df = payload["portfolio_df"].groupby("Portfolio_Segment", as_index=False)["Budget_Overrun_Pct"].mean()
        burn = px.bar(cost_df, x="Portfolio_Segment", y="Budget_Overrun_Pct", title="Budget burn profile by segment")
        st.plotly_chart(_theme_chart(burn), use_container_width=True)
    with charts[1]:
        milestone_df = pd.DataFrame(payload["tasks"])
        milestone_df["Cumulative"] = milestone_df["mean_duration"].cumsum()
        milestone = px.line(milestone_df, x="name", y="Cumulative", markers=True, title="Milestone tracker")
        st.plotly_chart(_theme_chart(milestone), use_container_width=True)

    st.download_button(
        "Export summary to JSON",
        data=_serialize_payload(payload),
        file_name="certain_executive_summary.json",
        mime="application/json",
        use_container_width=True,
    )


def _uploaded_file_extension(file_name: str) -> str:
    return Path(file_name).suffix.lower()


def _extract_docx_preview(file_bytes: bytes) -> tuple[str, dict[str, Any]]:
    document = DocxDocument(BytesIO(file_bytes))
    paragraph_lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
    preview_lines = paragraph_lines + table_lines
    preview_text = "\n".join(preview_lines).strip()
    meta = {
        "Paragraphs": len(paragraph_lines),
        "Tables": len(document.tables),
        "Characters": len(preview_text),
    }
    if not preview_text:
        preview_text = "No readable text content was found in this Word document."
    return preview_text, meta


def _extract_pdf_preview(file_bytes: bytes) -> tuple[str, dict[str, Any]]:
    reader = PdfReader(BytesIO(file_bytes))
    extracted_pages: list[str] = []
    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        if page_text:
            extracted_pages.append(page_text)
    preview_text = "\n\n".join(extracted_pages).strip()
    meta = {
        "Pages": len(reader.pages),
        "Extracted pages": len(extracted_pages),
        "Characters": len(preview_text),
    }
    if not preview_text:
        preview_text = "No readable text content was found in this PDF preview."
    return preview_text, meta


def _extract_spreadsheet_preview(file_name: str, file_bytes: bytes) -> tuple[pd.DataFrame, dict[str, Any]]:
    extension = _uploaded_file_extension(file_name)
    if extension == ".csv":
        preview_df = pd.read_csv(BytesIO(file_bytes))
        return preview_df, {"Rows": len(preview_df), "Columns": len(preview_df.columns), "Preview": "CSV"}

    workbook = pd.ExcelFile(BytesIO(file_bytes), engine="openpyxl")
    sheet_name = workbook.sheet_names[0]
    preview_df = pd.read_excel(workbook, sheet_name=sheet_name, engine="openpyxl")
    return preview_df, {
        "Rows": len(preview_df),
        "Columns": len(preview_df.columns),
        "Sheets": len(workbook.sheet_names),
        "Preview sheet": sheet_name,
    }


def _render_upload_meta(meta: dict[str, Any]) -> None:
    if not meta:
        return
    columns = st.columns(len(meta))
    for column, (label, value) in zip(columns, meta.items(), strict=False):
        column.metric(label, str(value))


def _render_uploaded_file_preview(file_name: str, file_bytes: bytes) -> None:
    extension = _uploaded_file_extension(file_name)
    try:
        if extension in {".csv", ".xlsx"}:
            preview_df, meta = _extract_spreadsheet_preview(file_name, file_bytes)
            _render_upload_meta(meta)
            st.dataframe(preview_df.head(20), use_container_width=True, hide_index=True)
            if len(preview_df) > 20:
                st.caption("Showing the first 20 rows from the uploaded spreadsheet.")
            return

        if extension == ".pdf":
            preview_text, meta = _extract_pdf_preview(file_bytes)
            _render_upload_meta(meta)
            st.text_area("PDF preview", value=preview_text[:6000], height=320, disabled=True)
            return

        if extension == ".docx":
            preview_text, meta = _extract_docx_preview(file_bytes)
            _render_upload_meta(meta)
            st.text_area("Word preview", value=preview_text[:6000], height=320, disabled=True)
            return

        st.warning("This file type is not yet supported in the preview workspace.")
    except Exception as exc:  # noqa: BLE001
        st.error(f"Unable to parse `{file_name}`: {exc}")


def _clean_upload_line(value: Any) -> str:
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except TypeError:
        pass
    text = str(value).replace("\xa0", " ")
    text = re.sub(r"[\t\r\n]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.strip(" -–—:;,.")


def _normalize_upload_token(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def _find_matching_column(columns: list[Any], keywords: list[str]) -> Any | None:
    for column in columns:
        normalized = str(column).lower().replace("_", " ").replace("-", " ")
        if any(keyword in normalized for keyword in keywords):
            return column
    return None


def _extract_dependency_tokens(value: Any) -> list[str]:
    clean = _clean_upload_line(value)
    if not clean or clean.lower() in {"none", "nan", "n/a", "na"}:
        return []
    parts = re.split(r"[,;/|\n]+|\band\b", clean, flags=re.IGNORECASE)
    return [token for token in (_clean_upload_line(part) for part in parts) if token]


def _parse_duration_days(value: Any) -> float | None:
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except TypeError:
        pass

    if isinstance(value, (int, float)):
        numeric_value = float(value)
        return numeric_value if numeric_value > 0 else None

    clean = _clean_upload_line(value).lower()
    if not clean:
        return None
    match = re.search(r"(\d+(?:\.\d+)?)", clean)
    if not match:
        return None

    numeric_value = float(match.group(1))
    if "month" in clean:
        return numeric_value * 20.0
    if "week" in clean or "wk" in clean:
        return numeric_value * 5.0
    if "hour" in clean:
        return max(1.0, numeric_value / 8.0)
    return numeric_value if numeric_value > 0 else None


def _fallback_upload_candidates(file_name: str) -> list[dict[str, Any]]:
    stem = Path(file_name).stem.replace("_", " ").replace("-", " ").strip() or "uploaded project"
    focus = re.sub(r"\s+", " ", stem)
    return [
        {"name": f"Review {focus}", "source_id": "1", "dependency_tokens": [], "duration_days": 4.0},
        {"name": "Define delivery workstreams", "source_id": "2", "dependency_tokens": ["1"], "duration_days": 5.0},
        {"name": "Design workflow architecture", "source_id": "3", "dependency_tokens": ["2"], "duration_days": 6.0},
        {"name": "Build execution plan", "source_id": "4", "dependency_tokens": ["3"], "duration_days": 8.0},
        {"name": "Validate launch readiness", "source_id": "5", "dependency_tokens": ["4"], "duration_days": 5.0},
    ]


def _estimate_upload_team_size(task_count: int, owner_count: int = 0) -> int:
    if owner_count > 0:
        return max(2, min(14, owner_count))
    return max(3, min(12, round(max(task_count, 1) / 2.5) + 1))


def _estimate_upload_timeline_weeks(task_count: int, dependency_count: int, duration_samples: list[float]) -> int:
    if duration_samples:
        total_days = sum(duration_samples[: min(12, len(duration_samples))])
        heuristic_days = max(total_days * 0.42, max(duration_samples) * 2.2)
    else:
        heuristic_days = (task_count * 2.6) + (dependency_count * 0.55)
    return max(3, int(round(max(15.0, heuristic_days) / 5.0)))


def _derive_upload_project_goal(file_name: str, task_candidates: list[dict[str, Any]]) -> str:
    stem = Path(file_name).stem.replace("_", " ").replace("-", " ").strip()
    if stem:
        return stem.title()
    if task_candidates:
        candidate = task_candidates[0]["name"]
        return candidate if len(candidate) <= 48 else f"{candidate[:45].rstrip()}..."
    return "Uploaded Project"


def _extract_spreadsheet_task_context(file_name: str, file_bytes: bytes) -> dict[str, Any]:
    preview_df, _meta = _extract_spreadsheet_preview(file_name, file_bytes)
    working_df = preview_df.dropna(how="all").copy()
    if working_df.empty:
        fallback_candidates = _fallback_upload_candidates(file_name)
        timeline_weeks = _estimate_upload_timeline_weeks(len(fallback_candidates), 4, [4.0, 5.0, 6.0])
        return {
            "project_goal": _derive_upload_project_goal(file_name, fallback_candidates),
            "timeline_weeks": timeline_weeks,
            "team_size": _estimate_upload_team_size(len(fallback_candidates)),
            "task_candidates": fallback_candidates,
            "detected_tasks": len(fallback_candidates),
            "dependency_count": 4,
        }

    columns = list(working_df.columns)
    name_column = _find_matching_column(
        columns,
        ["task name", "task", "activity", "work package", "deliverable", "milestone", "name", "summary"],
    )
    if name_column is None:
        name_column = columns[0]

    id_column = _find_matching_column(columns, ["task id", "activity id", "wbs", "id", "code"])
    dependency_column = _find_matching_column(columns, ["dependency", "depends on", "predecessor", "blocked by", "after"])
    duration_column = _find_matching_column(columns, ["duration", "estimate", "effort", "days", "weeks", "timeline", "eta"])
    owner_column = _find_matching_column(columns, ["owner", "assignee", "resource", "team member", "lead"])

    task_candidates: list[dict[str, Any]] = []
    seen_names: set[str] = set()
    duration_samples: list[float] = []
    dependency_count = 0
    owner_names: set[str] = set()
    detected_tasks = 0

    for _, row in working_df.iterrows():
        name = _clean_upload_line(row.get(name_column))
        if not name or re.fullmatch(r"[\W\d_]+", name):
            continue

        detected_tasks += 1
        dependency_tokens = _extract_dependency_tokens(row.get(dependency_column)) if dependency_column else []
        dependency_count += len(dependency_tokens)
        duration_days = _parse_duration_days(row.get(duration_column)) if duration_column else None
        if duration_days is not None:
            duration_samples.append(duration_days)

        if owner_column:
            owner_names.update(token for token in _extract_dependency_tokens(row.get(owner_column)) if token)

        normalized_name = _normalize_upload_token(name)
        if normalized_name in seen_names:
            continue
        seen_names.add(normalized_name)

        source_id = _clean_upload_line(row.get(id_column)) if id_column else str(detected_tasks)
        task_candidates.append(
            {
                "name": name,
                "source_id": source_id or str(detected_tasks),
                "dependency_tokens": dependency_tokens,
                "duration_days": duration_days,
            }
        )
        if len(task_candidates) >= 8:
            break

    if not task_candidates:
        task_candidates = _fallback_upload_candidates(file_name)

    timeline_weeks = _estimate_upload_timeline_weeks(max(detected_tasks, len(task_candidates)), dependency_count, duration_samples)
    return {
        "project_goal": _derive_upload_project_goal(file_name, task_candidates),
        "timeline_weeks": timeline_weeks,
        "team_size": _estimate_upload_team_size(max(detected_tasks, len(task_candidates)), len(owner_names)),
        "task_candidates": task_candidates,
        "detected_tasks": max(detected_tasks, len(task_candidates)),
        "dependency_count": max(dependency_count, len(task_candidates) - 1),
    }


def _extract_text_task_context(file_name: str, file_bytes: bytes, extension: str) -> dict[str, Any]:
    if extension == ".pdf":
        preview_text, _meta = _extract_pdf_preview(file_bytes)
    else:
        preview_text, _meta = _extract_docx_preview(file_bytes)

    lines: list[str] = []
    for raw_line in re.split(r"[\r\n]+", preview_text):
        candidate = _clean_upload_line(raw_line)
        candidate = re.sub(r"^(?:[-•*]|\d+[\.\)])\s*", "", candidate)
        if len(candidate) < 4:
            continue
        if len(candidate.split()) > 12:
            candidate = re.split(r"[.;:]", candidate, maxsplit=1)[0].strip()
        if candidate:
            lines.append(candidate)

    task_candidates: list[dict[str, Any]] = []
    seen_names: set[str] = set()
    for line in lines:
        normalized = _normalize_upload_token(line)
        if not normalized or normalized in seen_names:
            continue
        seen_names.add(normalized)
        task_candidates.append(
            {
                "name": line,
                "source_id": str(len(task_candidates) + 1),
                "dependency_tokens": [],
                "duration_days": None,
            }
        )
        if len(task_candidates) >= 8:
            break

    if not task_candidates:
        task_candidates = _fallback_upload_candidates(file_name)

    detected_tasks = max(len(task_candidates), min(18, max(4, len(lines))))
    dependency_count = max(0, len(task_candidates) - 1)
    timeline_weeks = _estimate_upload_timeline_weeks(detected_tasks, dependency_count, [])
    return {
        "project_goal": _derive_upload_project_goal(file_name, task_candidates),
        "timeline_weeks": timeline_weeks,
        "team_size": _estimate_upload_team_size(detected_tasks),
        "task_candidates": task_candidates,
        "detected_tasks": detected_tasks,
        "dependency_count": dependency_count,
    }


def _build_upload_task_architect_tasks(
    project_goal: str,
    timeline_weeks: int,
    team_size: int,
    task_candidates: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    base_plan = generate_task_plan("Create a delivery plan", max_tasks=8, mode="mock")
    base_tasks = [task.model_dump() for task in base_plan.tasks]
    candidates = copy.deepcopy(task_candidates[:8]) or _fallback_upload_candidates(project_goal)
    selected_candidates = candidates[: min(len(candidates), len(base_tasks))]
    if len(selected_candidates) < 4:
        selected_candidates = _fallback_upload_candidates(project_goal)[: len(base_tasks)]

    target_days = max(20.0, timeline_weeks * 5.0)
    base_total = sum(float(task["mean_duration"]) for task in base_tasks[: len(selected_candidates)])
    schedule_scale = max(0.7, min(1.5, target_days / max(base_total, 1.0)))
    team_factor = max(0.82, min(1.14, 1.08 - ((min(team_size, 12) - 5) * 0.03)))

    normalized_lookup: dict[str, str] = {}
    for index, candidate in enumerate(selected_candidates, start=1):
        task_id = f"T{index}"
        source_id = candidate.get("source_id") or str(index)
        for token in {
            _normalize_upload_token(task_id),
            _normalize_upload_token(str(index)),
            _normalize_upload_token(str(source_id)),
            _normalize_upload_token(str(candidate["name"])),
        }:
            if token:
                normalized_lookup[token] = task_id

    tasks: list[dict[str, Any]] = []
    for index, candidate in enumerate(selected_candidates, start=1):
        task_id = f"T{index}"
        base_task = base_tasks[min(index - 1, len(base_tasks) - 1)]
        dependencies: list[str] = []
        for token in candidate.get("dependency_tokens", []):
            dependency_id = normalized_lookup.get(_normalize_upload_token(token))
            if dependency_id and dependency_id != task_id and int(dependency_id[1:]) < index and dependency_id not in dependencies:
                dependencies.append(dependency_id)

        if not dependencies and index > 1:
            dependencies = [f"T{index - 1}"]
            if index > 3 and index % 3 == 1:
                dependencies.append(f"T{index - 2}")

        duration_days = candidate.get("duration_days")
        if duration_days is None:
            duration_days = float(base_task["mean_duration"]) * schedule_scale * team_factor
        mean_duration = max(2, round(float(duration_days)))
        risk_factor = min(
            0.82,
            max(0.18, round(float(base_task["risk_factor"]) + (len(dependencies) * 0.05), 2)),
        )
        std_dev = round(max(1.0, mean_duration * (0.17 + risk_factor * 0.18)), 1)
        tasks.append(
            {
                **base_task,
                "id": task_id,
                "name": str(candidate["name"]),
                "mean_duration": mean_duration,
                "std_dev": std_dev,
                "dependencies": dependencies,
                "risk_factor": risk_factor,
            }
        )
    return tasks


def _generate_task_architect_payload_from_upload(file_name: str, file_bytes: bytes) -> dict[str, Any]:
    extension = _uploaded_file_extension(file_name)
    if extension in {".csv", ".xlsx"}:
        upload_context = _extract_spreadsheet_task_context(file_name, file_bytes)
    elif extension in {".pdf", ".docx"}:
        upload_context = _extract_text_task_context(file_name, file_bytes, extension)
    else:
        raise TaskGenerationError("Upload a supported project file first.")

    project_goal = upload_context["project_goal"]
    timeline_weeks = int(upload_context["timeline_weeks"])
    deadline_days = float(max(15, timeline_weeks * 5))
    timeline_label = f"{timeline_weeks} weeks"
    team_size = int(upload_context["team_size"])
    tasks = _build_upload_task_architect_tasks(project_goal, timeline_weeks, team_size, upload_context["task_candidates"])
    iterations = 2000
    seed = int(st.session_state.get("saas_seed", 17))
    graph = build_project_graph(tasks)
    critical_path, critical_days = critical_path_by_mean(graph)
    completion = run_monte_carlo(graph, iterations=iterations, seed=seed)
    metrics = compute_metrics(completion, deadline_days)
    drivers = rank_delay_drivers(graph, iterations=min(500, iterations), seed=seed)
    scenarios_df = scenario_comparison(
        build_project_graph,
        tasks,
        deadline_days=deadline_days,
        iterations=iterations,
        seed=seed,
    )
    ml_features_df, ml_predictions_df, ml_summary_df, model_status = _run_ml_scoring(tasks)
    portfolio_df = _load_portfolio_history()
    monitoring_df = _load_monitoring_snapshot()
    model_metrics = _load_model_metrics()
    portfolio_summary = _portfolio_summary(portfolio_df)
    monitoring_summary = _monitoring_summary(monitoring_df)

    return {
        "project_text": project_goal,
        "sample": "AI Task Architect Upload",
        "mode": "mock",
        "deadline_days": deadline_days,
        "iterations": iterations,
        "max_tasks": len(tasks),
        "seed": seed,
        "tasks": tasks,
        "graph": graph,
        "critical_path": critical_path,
        "critical_path_days": critical_days,
        "completion": completion,
        "metrics": metrics,
        "drivers": drivers,
        "scenarios_df": scenarios_df,
        "ml_features_df": ml_features_df,
        "ml_predictions_df": ml_predictions_df,
        "ml_summary_df": ml_summary_df,
        "model_status": asdict(model_status),
        "portfolio_df": portfolio_df,
        "monitoring_df": monitoring_df,
        "model_metrics": model_metrics,
        "portfolio_summary": portfolio_summary,
        "monitoring_summary": monitoring_summary,
        "architect_meta": {
            "project_goal": project_goal,
            "timeline_text": timeline_label,
            "timeline_weeks": timeline_weeks,
            "team_size": team_size,
            "source_file": file_name,
            "detected_tasks": int(upload_context["detected_tasks"]),
            "dependency_count": int(upload_context["dependency_count"]),
            "generation_source": "upload",
        },
    }


def _render_upload_intake_section(
    *,
    upload_key: str,
    show_generate_button: bool = False,
    generate_button_key: str | None = None,
) -> tuple[Any | None, bool]:
    st.markdown(
        f"""
        <div class="page-shell">
          <div class="glass dashboard-banner">
            <h1 class="page-title">Upload files or define project parameters. {_brand_name_html()} transforms them into optimized, risk-aware plans.</h1>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    generate_clicked = False
    if show_generate_button and generate_button_key:
        upload_col, action_col = st.columns([5.2, 1.35], gap="small")
        with upload_col:
            uploaded = st.file_uploader(
                "Upload .xlsx, .csv, .pdf, or .docx",
                type=["xlsx", "csv", "pdf", "docx"],
                key=upload_key,
                label_visibility="collapsed",
            )
        with action_col:
            generate_clicked = st.button("Generate AI Plan", key=generate_button_key, use_container_width=True)
    else:
        uploaded = st.file_uploader(
            "Upload .xlsx, .csv, .pdf, or .docx",
            type=["xlsx", "csv", "pdf", "docx"],
            key=upload_key,
            label_visibility="collapsed",
        )
    if uploaded is not None:
        st.success(f"Uploaded `{uploaded.name}`")
        _render_uploaded_file_preview(uploaded.name, uploaded.getvalue())

    st.markdown(
        """
        <div class="history-card">
          <strong style="font-family:Sora,sans-serif;">Supported templates</strong>
          <p>Import: Excel, CSV, PDF, Word • Enterprise integrations (MPP, XML) coming soon</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    return uploaded, generate_clicked

def _render_integrations() -> None:
    _render_command_bar()
    integrations = [
        "Jira", "Asana", "Monday.com", "MS Project", "Trello", "Wrike",
        "Microsoft Planner", "Figma", "Google Workspace", "Bing", "Microsoft",
    ]
    cards = "".join(
        f"""
        <div class="feature-card">
          <span class="eyebrow">Available</span>
          <h3>{tool}</h3>
          <p>Connect planning data and sync project status into {_brand_name_html()}.</p>
        </div>
        """
        for tool in integrations
    )
    st.markdown(f"<div class='feature-grid'>{cards}</div>", unsafe_allow_html=True)


def _render_settings_page() -> None:
    _render_command_bar()
    left, right = st.columns(2, gap="large")
    with left:
        st.text_input("Profile name", value="certAIn Team", key="settings_name")
        st.text_input("Email", value=st.session_state.saas_user_email, key="settings_email")
        st.slider("AI sensitivity", min_value=0.1, max_value=1.0, value=float(st.session_state.saas_ai_sensitivity), key="saas_ai_sensitivity")
        st.slider("Confidence threshold", min_value=0.5, max_value=0.99, value=float(st.session_state.saas_confidence_threshold), key="saas_confidence_threshold")
    with right:
        st.toggle("Email alerts", value=bool(st.session_state.saas_email_alerts), key="saas_email_alerts")
        st.toggle("In-app alerts", value=bool(st.session_state.saas_in_app_alerts), key="saas_in_app_alerts")
        st.toggle("AI alerts", value=bool(st.session_state.saas_ai_alerts), key="saas_ai_alerts")
        language_labels = [label for _, label, _ in LANGUAGE_OPTIONS]
        current_language_label = st.session_state.get("saas_language", "English")
        selected_language = st.selectbox(
            "Language",
            language_labels,
            index=language_labels.index(current_language_label) if current_language_label in language_labels else 0,
            key="settings_language_selector",
        )
        if selected_language != current_language_label:
            selected_code = next(code for code, label, _ in LANGUAGE_OPTIONS if label == selected_language)
            st.session_state.saas_language = selected_language
            st.session_state.saas_language_code = selected_code
            try:
                st.query_params["lang"] = selected_code
            except Exception:
                st.experimental_set_query_params(page="settings-page", lang=selected_code)
            st.rerun()
        st.selectbox("Region", ["Global", "EU", "US"], key="saas_region")


def main() -> None:
    _init_state()
    _sync_auth_email_from_query()
    _sync_contact_email_from_query()
    _sync_registration_email_from_query()
    _sync_language_from_query()
    _inject_styles()
    page = _resolve_page()
    _sync_newsletter_subscription_flow(page)
    _sync_registration_flow(page)
    _inject_layout_state(page)
    _render_sidebar(page)
    _render_nav(page)
    _render_floating_shortcut(page)
    _render_command_palette(page)
    _mount_command_palette_shortcuts()
    _mount_page_translation()

    if page == "home":
        _render_home()
    elif page == "product":
        _render_product()
    elif page == "forecasting":
        _render_forecasting()
    elif page == "monte-carlo":
        _render_monte_carlo()
    elif page == "pricing":
        _render_pricing()
    elif page == "contact":
        _render_contact()
    elif page == "faq":
        _render_faq()
    elif page in {"privacy-policy", "terms-of-service", "cookie-settings", "security-data-protection"}:
        _render_legal_page(page)
    elif page == "subscription-confirmation":
        _render_subscription_confirmation()
    elif page == "registration-confirmation":
        _render_registration_confirmation()
    elif page == "about":
        _render_about()
    elif page == "login":
        _render_login()
    elif page == "guide":
        _render_user_guide()
    elif page == "decision-hub":
        _render_decision_hub()
    elif page == "task-architect":
        _render_task_architect()
    elif page == "risk-intelligence":
        _render_risk_intelligence()
    elif page == "simulator":
        _render_simulator()
    elif page == "what-if":
        _render_what_if()
    elif page == "history":
        _render_history()
    elif page == "summary":
        _render_summary()
    elif page == "integrations":
        _render_integrations()
    elif page == "settings-page":
        _render_settings_page()
    else:
        _render_dashboard()

    _render_footer(page)
    _render_ai_assistant(page)
    _mount_ai_assistant(page)


if __name__ == "__main__":
    main()
